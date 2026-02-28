#!/usr/bin/env python3
"""Generate the River Valley Vikings Defensive Playbook WITH diagrams."""

import os
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# COLORS
# =============================================================================
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD_RGB = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_RGB = RGBColor(0x88, 0x88, 0x88)
NAVY_HEX = "002D62"

# Matplotlib colors
C_NAVY = "#002D62"
C_COLUMBIA = "#6CACE4"
C_GOLD = "#CFA700"
C_GRAY = "#888888"
C_OL = "#AAAAAA"
C_WR = "#666666"
C_WHITE = "#FFFFFF"
C_FIELD = "#4a7c3f"

DIAGRAM_DIR = "/tmp/rv_playbook_diagrams"
os.makedirs(DIAGRAM_DIR, exist_ok=True)

# =============================================================================
# FIELD / DIAGRAM HELPERS
# =============================================================================
# Offense positions (x coords): C=0, G=+-2, T=+-4
# Field = right (positive x), Boundary = left (negative x)
# y: LOS=0, defense is positive y (upfield)

# Technique -> x offset from center of the OL man
# For guards at +-2 and tackles at +-4
TECH_X = {
    # Relative to center
    "0":  0.0,    # head up C
    "1f": 0.5,    # shade C to field
    "1b": -0.5,   # shade C to boundary
    "2i_f": 1.5,  # inside shade field G
    "2i_b": -1.5, # inside shade boundary G
    "2f":  2.0,   # head up field G
    "2b": -2.0,   # head up boundary G
    "3f":  2.5,   # outside shade field G
    "3b": -2.5,   # outside shade boundary G
    "4i_f": 3.5,  # inside shade field T
    "4i_b": -3.5, # inside shade boundary T
    "4f":  4.0,   # head up field T
    "4b": -4.0,   # head up boundary T
    "5f":  4.7,   # outside shade field T
    "5b": -4.7,   # outside shade boundary T
    "6i_f": 5.3,  # inside shade field TE
    "7f":  6.7,   # outside shade field TE
    "7b": -6.7,   # outside shade boundary TE
    "9f":  8.0,   # wide field
    "9b": -8.0,   # wide boundary
    "10f": 5.5,   # 10 tech field (off ball, between T and TE)
}

# Standard y-depths
Y_LOS = 0
Y_DL = 1.2
Y_LB = 3.5
Y_OLB = 2.5
Y_APEX = 3.0
Y_CB = 5.5
Y_S = 8.0

# Standard WR/slot positions for 2x2
WR_FIELD = (13, 0)
SLOT_FIELD = (7, 0)
WR_BND = (-13, 0)
SLOT_BND = (-7, 0)

# Standard OL
OL_POSITIONS = [
    (0, 0, "C"),
    (2, 0, "G"),
    (-2, 0, "G"),
    (4, 0, "T"),
    (-4, 0, "T"),
]

QB_POS = (0, -1.5)
RB_POS = (0, -3)


def new_fig(title="", figsize=(10, 6)):
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-16, 16)
    ax.set_ylim(-5, 11)
    ax.set_aspect("equal")
    ax.axis("off")
    # LOS line
    ax.axhline(y=0.5, color=C_GRAY, linewidth=1, linestyle="--", alpha=0.4)
    # Field/Boundary labels
    ax.text(14, 10, "FIELD \u2192", fontsize=8, color=C_GRAY, ha="right", va="top", style="italic")
    ax.text(-14, 10, "\u2190 BOUNDARY", fontsize=8, color=C_GRAY, ha="left", va="top", style="italic")
    if title:
        ax.set_title(title, fontsize=14, fontweight="bold", color=C_NAVY, pad=10)
    return fig, ax


def draw_ol(ax, te_side=None):
    """Draw offensive line. te_side='field' or 'boundary' or None."""
    for x, y, lbl in OL_POSITIONS:
        ax.add_patch(plt.Rectangle((x - 0.4, y - 0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", linewidth=1.5))
    if te_side == "field":
        ax.add_patch(plt.Rectangle((6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", linewidth=1.5))
        ax.text(6, -1.0, "TE", fontsize=7, ha="center", color="#555")
    elif te_side == "boundary":
        ax.add_patch(plt.Rectangle((-6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", linewidth=1.5))
        ax.text(-6, -1.0, "TE", fontsize=7, ha="center", color="#555")


def draw_qb_rb(ax, show_rb=True):
    ax.plot(*QB_POS, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(QB_POS[0], QB_POS[1] - 0.7, "QB", fontsize=7, ha="center", color="#555")
    if show_rb:
        ax.plot(*RB_POS, "s", color=C_WR, markersize=8, markeredgecolor="#333")
        ax.text(RB_POS[0], RB_POS[1] - 0.7, "RB", fontsize=7, ha="center", color="#555")


def draw_wr_2x2(ax):
    for pos, lbl in [(WR_FIELD, "#1"), (SLOT_FIELD, "#2"), (WR_BND, "#1"), (SLOT_BND, "#2")]:
        ax.plot(*pos, "s", color=C_WR, markersize=7, markeredgecolor="#333")
        ax.text(pos[0], pos[1] - 0.7, lbl, fontsize=6, ha="center", color="#555")


def draw_wr_3x1_field(ax):
    """3x1 trips to field."""
    trips = [(13, 0, "#1"), (9, 0, "#2"), (7, 0, "#3")]
    for x, y, lbl in trips:
        ax.plot(x, y, "s", color=C_WR, markersize=7, markeredgecolor="#333")
        ax.text(x, y - 0.7, lbl, fontsize=6, ha="center", color="#555")
    # Single receiver boundary
    ax.plot(-13, 0, "s", color=C_WR, markersize=7, markeredgecolor="#333")
    ax.text(-13, -0.7, "#1", fontsize=6, ha="center", color="#555")


def def_player(ax, x, y, label, color=C_NAVY, fontsize=8):
    """Draw a defensive player."""
    ax.plot(x, y, "o", color=color, markersize=14, markeredgecolor=C_NAVY, markeredgewidth=1.2)
    ax.text(x, y, label, fontsize=fontsize, ha="center", va="center",
            color=C_WHITE, fontweight="bold")


def def_player_cb(ax, x, y, label):
    def_player(ax, x, y, label, color=C_COLUMBIA)


def def_player_safety(ax, x, y, label):
    def_player(ax, x, y, label, color=C_GOLD)


def arrow(ax, x1, y1, x2, y2, color=C_NAVY):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=2))


def save_fig(fig, name):
    path = os.path.join(DIAGRAM_DIR, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# =============================================================================
# FRONT DIAGRAMS
# =============================================================================

def draw_front(title, a_x, t_x, n_x, e_x, m_x=1.5, w_x=-1.5,
               b_x=7.5, b_y=Y_APEX, d_x=-6, fc_x=12, bc_x=-12,
               fs_x=5, ds_x=-5, te_side=None,
               extra_labels=None, m_y=Y_LB, w_y=Y_LB,
               grizzly=False, boss_m_label=None):
    fig, ax = new_fig(title)
    draw_ol(ax, te_side=te_side)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL
    def_player(ax, a_x, Y_DL, "A")
    def_player(ax, t_x, Y_DL, "T")
    def_player(ax, n_x, Y_DL, "N")
    def_player(ax, e_x, Y_DL, "E")

    # LBs
    if grizzly:
        # Grizzly: M in a 10, B and W are OLBs
        def_player(ax, 5.5, Y_OLB, "M", color=C_COLUMBIA)
        def_player(ax, 6.5, Y_OLB, "B", color=C_COLUMBIA)
        def_player(ax, -6.5, Y_OLB, "W", color=C_COLUMBIA)
    else:
        def_player(ax, m_x, m_y, "M", color=C_COLUMBIA)
        def_player(ax, w_x, w_y, "W", color=C_COLUMBIA)
        def_player(ax, b_x, b_y, "B", color=C_COLUMBIA)

    # DBs
    def_player_cb(ax, fc_x, Y_CB, "FC")
    def_player_cb(ax, bc_x, Y_CB, "BC")
    def_player_safety(ax, fs_x, Y_S, "FS")
    def_player_safety(ax, ds_x, Y_S, "D")

    if extra_labels:
        for (x, y, txt) in extra_labels:
            ax.text(x, y, txt, fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_3down(title, a_x, t_x, n_x, b_x=7.5, e_x=-7.5,
               m_x=1.5, w_x=-1.5, fc_x=12, bc_x=-12,
               fs_x=5, ds_x=-5):
    fig, ax = new_fig(title)
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Down 3
    def_player(ax, a_x, Y_DL, "A")
    def_player(ax, t_x, Y_DL, "T")
    def_player(ax, n_x, Y_DL, "N")

    # OLBs
    def_player(ax, b_x, Y_OLB, "B", color=C_COLUMBIA)
    def_player(ax, e_x, Y_OLB, "E", color=C_COLUMBIA)

    # ILBs
    def_player(ax, m_x, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, w_x, Y_LB, "W", color=C_COLUMBIA)

    # DBs
    def_player_cb(ax, fc_x, Y_CB, "FC")
    def_player_cb(ax, bc_x, Y_CB, "BC")
    def_player_safety(ax, fs_x, Y_S, "FS")
    def_player_safety(ax, ds_x, Y_S, "D")

    return fig


# =============================================================================
# COVERAGE DIAGRAMS
# =============================================================================

def draw_ninja_2x2():
    fig, ax = new_fig("NINJA vs 2x2 \u2014 MOD (Field) / CLAMP (Boundary)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL (Shade base)
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # LBs
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # DBs -- NINJA alignment
    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 6, Y_S, "FS")
    def_player_safety(ax, -6, Y_S, "D")

    # Zone shading
    ax.axvline(x=0, color=C_GRAY, linewidth=1, linestyle=":", alpha=0.3)

    # Labels
    ax.text(7, 10.2, "MOD", fontsize=11, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor=C_COLUMBIA, alpha=0.3))
    ax.text(-7, 10.2, "CLAMP", fontsize=11, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor=C_GOLD, alpha=0.3))

    return fig


def draw_ninja_3x1():
    fig, ax = new_fig("NINJA vs 3x1 \u2014 POACH")
    draw_ol(ax)
    draw_wr_3x1_field(ax)
    draw_qb_rb(ax)

    # DL (Shade base)
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # LBs
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # DBs
    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    # D is poach player for trips to field (boundary safety poaches toward trips)
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -3, 7.5, "D")

    ax.text(-3, 10.2, "POACH (D)", fontsize=11, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor=C_GOLD, alpha=0.3))
    ax.text(-3, 8.5, "#3 vertical?\nD takes it", fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_cover1(title, post_label, post_x, post_y,
                fs_assignment, d_assignment, b_assignment):
    fig, ax = new_fig(title)
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL (Shade base)
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # ILBs
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    # Corners (man on #1)
    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")

    # Draw the three variable positions based on which variant
    if post_label == "FS":
        def_player_safety(ax, 0, 9, "FS")  # post/MOF
        ax.text(0, 10, "POST", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
        def_player_safety(ax, -6, 6, "D")
        ax.text(-6, 5, d_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")
        def_player(ax, 7, Y_APEX, "B", color=C_COLUMBIA)
        ax.text(7, Y_APEX - 1.2, b_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")
    elif post_label == "D":
        def_player_safety(ax, 0, 9, "D")  # post/MOF
        ax.text(0, 10, "POST", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
        def_player_safety(ax, 5, 6, "FS")
        ax.text(5, 5, fs_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")
        def_player(ax, 7, Y_APEX, "B", color=C_COLUMBIA)
        ax.text(7, Y_APEX - 1.2, b_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")
    elif post_label == "B":
        def_player(ax, 0, 9, "B", color=C_COLUMBIA)  # post/MOF
        ax.text(0, 10, "POST", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
        def_player_safety(ax, 5, 6, "FS")
        ax.text(5, 5, fs_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")
        def_player_safety(ax, -6, 6, "D")
        ax.text(-6, 5, d_assignment, fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_zeus():
    fig, ax = new_fig("ZEUS \u2014 Cover 0 Delayed Pressure")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL -- always rush on pass
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # M -- primary add-on rusher
    def_player(ax, 0, Y_LB, "M", color=C_COLUMBIA)

    # B and W -- funnel RB
    def_player(ax, 4, Y_LB, "B", color=C_COLUMBIA)
    def_player(ax, -3, Y_LB, "W", color=C_COLUMBIA)

    # DBs -- C0 man
    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 6, 6, "FS")
    def_player_safety(ax, -6, 6, "D")

    # Rush arrows -- toward the QB (downfield)
    arrow(ax, 4.7, Y_DL, 4.2, Y_DL - 1.5)   # A cage/contain (slight inside angle)
    arrow(ax, 2.5, Y_DL, 1.5, Y_DL - 1.5)    # T vertical push
    arrow(ax, -1.5, Y_DL, -0.8, Y_DL - 1.5)  # N vertical push
    arrow(ax, -4.7, Y_DL, -4.2, Y_DL - 1.5)  # E cage/contain (slight inside angle)
    arrow(ax, 0, Y_LB, 0, Y_LB - 2.0, color=C_COLUMBIA)  # M rush

    # Labels
    ax.text(5.3, Y_DL - 0.5, "cage", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-5.3, Y_DL - 0.5, "cage", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0.7, Y_LB - 0.5, "rush", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(4, Y_LB + 0.7, "RB funnel", fontsize=6, ha="center", color=C_NAVY, style="italic")
    ax.text(-3, Y_LB + 0.7, "RB funnel", fontsize=6, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_viking():
    fig, ax = new_fig("VIKING \u2014 Cover 3 (Spot Drop)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # Underneath
    def_player(ax, 7, Y_APEX, "B", color=C_COLUMBIA)
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player_safety(ax, -6, 5.5, "D")

    # Deep 3
    def_player_cb(ax, 11, 8.5, "FC")
    def_player_safety(ax, 0, 9.5, "FS")
    def_player_cb(ax, -11, 8.5, "BC")

    # Zone labels
    ax.text(11, 10, "Deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0, 10.5, "Deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-11, 10, "Deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(7, Y_APEX + 0.8, "curl/flat", fontsize=6, ha="center", color=C_NAVY, style="italic")
    ax.text(1.5, Y_LB + 0.8, "hook", fontsize=6, ha="center", color=C_NAVY, style="italic")
    ax.text(-1.5, Y_LB + 0.8, "hook", fontsize=6, ha="center", color=C_NAVY, style="italic")
    ax.text(-6, 6.3, "seam-curl-flat", fontsize=6, ha="center", color=C_NAVY, style="italic")

    return fig


# =============================================================================
# STUNT DIAGRAMS
# =============================================================================

def draw_stunt_slant():
    fig, ax = new_fig("SLANT (to Field)")
    draw_ol(ax)
    draw_qb_rb(ax, show_rb=False)

    # Under front (best pairing): A=5, T=2i, N=3, E=5
    # T=2i on field guard inside = x=1.5, N=3 on boundary guard outside = x=-2.5
    def_player(ax, 4.7, Y_DL, "A")
    t_start = 1.5
    n_start = -2.5
    def_player(ax, t_start, Y_DL, "T")
    def_player(ax, n_start, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # Slant arrows to field (right) and toward LOS (down)
    arrow(ax, t_start, Y_DL, t_start + 1.5, Y_DL - 0.8)
    arrow(ax, n_start, Y_DL, n_start + 1.5, Y_DL - 0.8)

    # LBs
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    ax.text(0, -4, "T + N slant to field. Edges play normal contain.", fontsize=9, ha="center", color=C_NAVY)

    return fig


def draw_stunt_pinch():
    fig, ax = new_fig("PINCH")
    draw_ol(ax)
    draw_qb_rb(ax, show_rb=False)

    # Shade base
    def_player(ax, 4.7, Y_DL, "A")
    t_x, n_x = 2.5, -1.5
    def_player(ax, t_x, Y_DL, "T")
    def_player(ax, n_x, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # Pinch arrows into A gaps (inward and toward LOS)
    arrow(ax, t_x, Y_DL, 0.7, Y_DL - 0.8)
    arrow(ax, n_x, Y_DL, -0.7, Y_DL - 0.8)

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    ax.text(0, -4, "T + N pinch into A gaps. A and E play normal.", fontsize=9, ha="center", color=C_NAVY)

    return fig


def draw_stunt_jacks():
    fig, ax = new_fig("JACKS")
    draw_ol(ax)
    draw_qb_rb(ax, show_rb=False)

    # Shade base
    def_player(ax, 4.7, Y_DL, "A")
    t_x, n_x = 2.5, -1.5
    def_player(ax, t_x, Y_DL, "T")
    def_player(ax, n_x, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # Jacks arrows into B gaps (outward and toward LOS)
    arrow(ax, t_x, Y_DL, 3.5, Y_DL - 0.8)
    arrow(ax, n_x, Y_DL, -3.0, Y_DL - 0.8)

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    ax.text(0, -4, "T + N shoot B gaps. A and E contain.", fontsize=9, ha="center", color=C_NAVY)

    return fig


def draw_stunt_crash():
    fig, ax = new_fig("CRASH")
    draw_ol(ax)
    draw_qb_rb(ax, show_rb=False)

    # Wide front (best pairing)
    a_x, t_x, n_x, e_x = 4.7, 2.5, -2.5, -4.7
    def_player(ax, a_x, Y_DL, "A")
    def_player(ax, t_x, Y_DL, "T")
    def_player(ax, n_x, Y_DL, "N")
    def_player(ax, e_x, Y_DL, "E")

    # Crash: A+E shoot B, T+N shoot A (all toward LOS)
    arrow(ax, a_x, Y_DL, 3.3, Y_DL - 0.8)   # A to B gap
    arrow(ax, e_x, Y_DL, -3.3, Y_DL - 0.8)  # E to B gap
    arrow(ax, t_x, Y_DL, 0.7, Y_DL - 0.8)   # T to A gap
    arrow(ax, n_x, Y_DL, -0.7, Y_DL - 0.8)  # N to A gap

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7, Y_APEX, "B", color=C_COLUMBIA)

    ax.text(0, -4, "A+E shoot B gaps. T+N shoot A gaps. B/M/W contain.", fontsize=9, ha="center", color=C_NAVY)

    return fig


# =============================================================================
# PRESSURE DIAGRAMS
# =============================================================================

def draw_pressure_swarm():
    fig, ax = new_fig("sWarM (M + W Blitz)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Shade base: M=open A (field), W=open B (boundary)
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # M blitzes A gap (field) -- between C and field G (~x=1.0)
    # W blitzes B gap (boundary) -- between bnd G and bnd T (~x=-3.0)
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    arrow(ax, 1.5, Y_LB, 1.0, Y_DL - 0.5, color=C_COLUMBIA)   # M to A gap field
    arrow(ax, -1.5, Y_LB, -3.0, Y_DL - 0.5, color=C_COLUMBIA)  # W to B gap boundary

    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -5, Y_S, "D")

    return fig


def draw_pressure_hammer():
    fig, ax = new_fig("HAMMER (B Edge + Anchor Attack)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Shade base
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # B edge blitz
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)
    arrow(ax, 7.5, Y_APEX, 5.5, Y_DL, color=C_COLUMBIA)

    # A anchor attack -- drives/washes OL inside to create the edge for B
    arrow(ax, 4.7, Y_DL, 3.0, Y_DL - 0.3)

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -5, Y_S, "D")

    ax.text(3.5, Y_DL + 0.7, "A washes\nOL inside", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(6.5, Y_APEX - 0.5, "B off edge", fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_pressure_eat():
    fig, ax = new_fig("EAT (M + W + B Blitz)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Shade: M=A gap field, W=B gap boundary, B=field edge
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    arrow(ax, 1.5, Y_LB, 1.0, Y_DL - 0.5, color=C_COLUMBIA)   # M to A gap field
    arrow(ax, -1.5, Y_LB, -3.0, Y_DL - 0.5, color=C_COLUMBIA)  # W to B gap boundary
    arrow(ax, 7.5, Y_APEX, 5.5, Y_DL - 0.3, color=C_COLUMBIA)  # B off the edge

    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -5, Y_S, "D")

    return fig


def draw_pressure_boom():
    fig, ax = new_fig("BooM (B + M Blitz)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Shade: M=A gap field, B=field edge
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    arrow(ax, 1.5, Y_LB, 1.0, Y_DL - 0.5, color=C_COLUMBIA)   # M to A gap field
    arrow(ax, 7.5, Y_APEX, 5.5, Y_DL - 0.3, color=C_COLUMBIA)  # B off the edge

    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -5, Y_S, "D")

    return fig


def draw_pressure_bow():
    fig, ax = new_fig("BoW (B + W Blitz)")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # Shade: W=B gap boundary, B=field edge
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    arrow(ax, -1.5, Y_LB, -3.0, Y_DL - 0.5, color=C_COLUMBIA)  # W to B gap boundary
    arrow(ax, 7.5, Y_APEX, 5.5, Y_DL - 0.3, color=C_COLUMBIA)  # B off the edge

    def_player_cb(ax, 12, Y_CB, "FC")
    def_player_cb(ax, -12, Y_CB, "BC")
    def_player_safety(ax, 5, Y_S, "FS")
    def_player_safety(ax, -5, Y_S, "D")

    return fig


# =============================================================================
# GENERATE ALL DIAGRAMS
# =============================================================================

def generate_all_diagrams():
    paths = {}

    # 4-Down Fronts
    fronts = {
        "shade": ("SHADE (Base \u2014 Set Field)", 4.7, 2.5, -1.5, -4.7),
        "under": ("UNDER (Set Boundary)", 4.7, 1.5, -2.5, -4.7),
        "eyes":  ("EYES (Balanced Interior)", 4.7, 1.5, -1.5, -4.7),
        "wide":  ("WIDE", 4.7, 2.5, -2.5, -4.7),
        "deuces": ("DEUCES", 4.7, 2.0, -2.0, -4.7),
        "boss":  ("BOSS (Bigs to Field)", 4.7, 2.5, 0.5, -4.7),
        "boss_under": ("BOSS UNDER (Bigs to Boundary)", 4.7, -0.5, -2.5, -4.7),
    }
    for key, (title, a, t, n, e) in fronts.items():
        fig = draw_front(f"4-DOWN: {title}", a, t, n, e)
        paths[key] = save_fig(fig, key)

    # Grizzly (special)
    fig = draw_front("4-DOWN: GRIZZLY", 3.5, 1.5, -1.5, -3.5, grizzly=True)
    paths["grizzly"] = save_fig(fig, "grizzly")

    # TE SET: TE to field -> A goes to 7-tech on TE, E stays 5 boundary
    # B inserts into box at LB depth over C gap (~x=5.0)
    fig = draw_front("TE SET Adjustment (Shade + TE Field)", 6.7, 2.5, -1.5, -4.7,
                     te_side="field", b_x=5.0, b_y=Y_LB,
                     extra_labels=[(6.7, Y_DL + 0.8, "A = 7-tech\ncontain"),
                                   (5.0, Y_LB + 0.8, "B inserts\nC gap")])
    paths["te_set"] = save_fig(fig, "te_set")

    # 3-Down Packages
    fig = draw_3down("3-DOWN: MINT (4i / 0 / 4i)", 3.5, 0, -3.5)
    paths["mint"] = save_fig(fig, "mint")
    fig = draw_3down("3-DOWN: ACE (4 / 0 / 4)", 4.0, 0, -4.0)
    paths["ace"] = save_fig(fig, "ace")
    fig = draw_3down("3-DOWN: JET (5 / 0 / 5)", 4.7, 0, -4.7)
    paths["jet"] = save_fig(fig, "jet")
    fig = draw_3down("3-DOWN: SLIP (5 / 0 / 4i)", 4.7, 0, -3.5)
    paths["slip"] = save_fig(fig, "slip")

    # Coverage diagrams
    fig = draw_ninja_2x2()
    paths["ninja_2x2"] = save_fig(fig, "ninja_2x2")
    fig = draw_ninja_3x1()
    paths["ninja_3x1"] = save_fig(fig, "ninja_3x1")

    fig = draw_cover1("COVER 1: OREGON (Post = FS)", "FS", 0, 9,
                      "", "man #2 bnd", "man #2 field")
    paths["oregon"] = save_fig(fig, "oregon")
    fig = draw_cover1("COVER 1: OKLAHOMA (Post = D)", "D", 0, 9,
                      "man #2 bnd", "", "man #2 field")
    paths["oklahoma"] = save_fig(fig, "oklahoma")
    fig = draw_cover1("COVER 1: OHIO (Post = B)", "B", 0, 9,
                      "man #2 field", "man #2 bnd", "")
    paths["ohio"] = save_fig(fig, "ohio")

    fig = draw_zeus()
    paths["zeus"] = save_fig(fig, "zeus")

    fig = draw_viking()
    paths["viking"] = save_fig(fig, "viking")

    # Stunts
    fig = draw_stunt_slant()
    paths["slant"] = save_fig(fig, "slant")
    fig = draw_stunt_pinch()
    paths["pinch"] = save_fig(fig, "pinch")
    fig = draw_stunt_jacks()
    paths["jacks"] = save_fig(fig, "jacks")
    fig = draw_stunt_crash()
    paths["crash"] = save_fig(fig, "crash")

    # Pressures
    fig = draw_pressure_swarm()
    paths["swarm"] = save_fig(fig, "swarm")
    fig = draw_pressure_hammer()
    paths["hammer"] = save_fig(fig, "hammer")
    fig = draw_pressure_eat()
    paths["eat"] = save_fig(fig, "eat")
    fig = draw_pressure_boom()
    paths["boom"] = save_fig(fig, "boom")
    fig = draw_pressure_bow()
    paths["bow"] = save_fig(fig, "bow")

    return paths


# =============================================================================
# DOCX HELPERS
# =============================================================================

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def make_header_row(table, headers):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, NAVY_HEX)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True


# =============================================================================
# BUILD THE DOCUMENT
# =============================================================================

def build_document(diagram_paths):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.font.color.rgb = DARK

    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def add_diagram(key, width=5.5):
        if key in diagram_paths:
            doc.add_picture(diagram_paths[key], width=Inches(width))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("DEFENSIVE PLAYBOOK")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = y.add_run("DIAGRAM EDITION \u2014 2026 Season")
    r.font.size = Pt(18); r.font.color.rgb = NAVY
    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("CONFIDENTIAL \u2014 COACHING STAFF ONLY")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    doc.add_page_break()

    # =========================================================================
    # 4-DOWN FRONTS
    # =========================================================================
    heading1("4-DOWN FRONT CATALOG")

    heading2("SHADE (Base \u2014 Set Field)")
    body("A=5, T=3, N=2i, E=5. M=open A (field), W=open B (boundary). Edges contain.")
    add_diagram("shade")
    doc.add_page_break()

    heading2("UNDER (Set Boundary)")
    body("A=5, T=2i, N=3, E=5. M=open B (field), W=open A (boundary). \"Same as Shade just to the boundary.\"")
    add_diagram("under")
    doc.add_page_break()

    heading2("EYES (Balanced Interior)")
    body("A=5, T=2i, N=2i, E=5. M=open B, W=open B. Square interior vs zone/duo.")
    add_diagram("eyes")
    doc.add_page_break()

    heading2("WIDE")
    body("A=5, T=3, N=3, E=5. M=open A, W=open A. Force bounce; vs B-gap heavy / gap schemes.")
    add_diagram("wide")
    doc.add_page_break()

    heading2("DEUCES")
    body("A=5, T=2, N=2, E=5. M/W make T and N right; gap only if stunt creates it. Vs Wing-T / pullers.")
    add_diagram("deuces")
    doc.add_page_break()

    heading2("GRIZZLY")
    body("A=4i, T=2i, N=2i, E=4i. Special structure: B and W are OLBs. M in a 10-tech, "
         "no gap by default (C to TE if TE surface). TE SET does NOT override.")
    add_diagram("grizzly")
    doc.add_page_break()

    heading2("BOSS (Bigs to Field)")
    body("A=5, T=3 (field), N=1 (field), E=5. M=A gap boundary, W=B gap boundary. TE SET overrides if TE present.")
    add_diagram("boss")
    doc.add_page_break()

    heading2("BOSS UNDER (Bigs to Boundary)")
    body("A=5, T=1 (boundary), N=3 (boundary), E=5. M=A gap field, W=B gap field. TE SET overrides if TE present.")
    add_diagram("boss_under")
    doc.add_page_break()

    heading2("TE SET \u2014 Default 4-Down Adjustment")
    body("Overrides ALL fronts except Grizzly when TE/Y-off surface. Set front to TE: 3-tech to TE, "
         "2i away. B inserts. End to TE aligns 7-tech and plays contain.")
    add_diagram("te_set")
    doc.add_page_break()

    # =========================================================================
    # 3-DOWN PACKAGES
    # =========================================================================
    heading1("3-DOWN PACKAGES")
    body("Field OLB = B. Boundary OLB = E. Slice: backside OLB has slicer. "
         "ILB fits: play-to STACK B\u2192D; play-away SLOW PLAY A.")

    heading2("MINT (4i / 0 / 4i)")
    body("A=4i, T=0, N=4i. Contain: man = B; NINJA/VIKING = M is QB contain.")
    add_diagram("mint")
    doc.add_page_break()

    heading2("ACE (4 / 0 / 4)")
    body("A=4, T=0, N=4. A/T/N are ALL 2-gapping (no stunt). Contains = OLBs (B and E).")
    add_diagram("ace")
    doc.add_page_break()

    heading2("JET (5 / 0 / 5)")
    body("A=5, T=0, N=5. T is 2-gapping. A and N are contain (C-gap edges).")
    add_diagram("jet")
    doc.add_page_break()

    heading2("SLIP (5 / 0 / 4i)")
    body("A=5, T=0, N=4i. 1-gap penetrating. A and E are contain. "
         "Use: 3-down vs spread with NINJA or VIKING. Takes B out of conflict \u2014 B plays coverage clean.")
    add_diagram("slip")
    doc.add_page_break()

    # =========================================================================
    # COVERAGES
    # =========================================================================
    heading1("COVERAGE: NINJA (COVER 7 FAMILY)")
    body("Coach calls NINJA. DBs auto-check: 2x2 = MOD/CLAMP, 3x1 = POACH. "
         "Corners: inside leverage, 6 yards. Safeties: 10-12 deep, align off #2.")

    heading2("NINJA vs 2x2 \u2014 MOD / CLAMP")
    bullet("MOD (Field): FC man-match #1. FS top-down on #2. B apex to #2/#3.", bold_prefix="Field: ")
    bullet("CLAMP (Boundary): BC clamp technique on #1. D control #2 and help corner.", bold_prefix="Boundary: ")
    add_diagram("ninja_2x2")
    doc.add_page_break()

    heading2("NINJA vs 3x1 \u2014 POACH")
    bullet("Trips to field: D is poach safety. Trips to boundary: FS is poach safety.")
    bullet("#3 vertical: poach player takes it. #3 shallow: overlap crossers. Backside stays sound.")
    add_diagram("ninja_3x1")
    doc.add_page_break()

    heading1("COVERAGE: COVER 1 FAMILY")
    body("CAMP: Corners = #1. Non-post safeties = #2 to their side or #3 away. M/W funnel RB.")

    heading2("OREGON (Post = FS)")
    bullet("FS = Post (MOF). D = man #2 boundary. B = man #2 field. M/W = RB funnel.")
    add_diagram("oregon")
    doc.add_page_break()

    heading2("OKLAHOMA (Post = D)")
    bullet("D = Post (MOF). FS = man #2 boundary. B = man #2 field. M/W = RB funnel.")
    add_diagram("oklahoma")
    doc.add_page_break()

    heading2("OHIO (Post = B)")
    bullet("B = Post (MOF). FS = man #2 field. D = man #2 boundary. M/W = RB funnel.")
    add_diagram("ohio")
    doc.add_page_break()

    heading1("COVERAGE: COVER 0 / Z-FAMILY")
    heading2("ZEUS \u2014 Delayed Pressure")
    body("Run-first; on pass: A/T/N/E/M rush. A/E cage contain. B/W funnel RB (to your side = take; away = rush; "
         "middle = W). Call-off: bail to NINJA or VIKING.")
    add_diagram("zeus")

    doc.add_paragraph()
    heading2("Z-Family RB Reference")
    ztable = doc.add_table(rows=6, cols=2)
    ztable.style = "Light Grid Accent 1"
    ztable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ztable, ["Z-Call", "RB Player"])
    for i, (c, p) in enumerate([("Zeus", "B or W (funnel)"), ("Zorro", "B"),
                                 ("Zunnel", "M or W (funnel)"), ("Zill", "W"), ("Zike", "M")]):
        ztable.rows[i + 1].cells[0].text = c
        ztable.rows[i + 1].cells[1].text = p
    doc.add_page_break()

    heading1("COVERAGE: VIKING (COVER 3)")
    body("Deep 3: FC=1/3 field, FS=1/3 middle, BC=1/3 boundary. "
         "Under: B=curl/flat, D=seam-curl-flat, M=hook, W=hook. "
         "Tags: SEAM, PUSH, CROSS, SCREEN. RIP/LIZ = true match 3.")
    add_diagram("viking")
    doc.add_page_break()

    # =========================================================================
    # STUNTS
    # =========================================================================
    heading1("STUNT CATALOG")

    heading2("SLANT (to Field)")
    body("4-down: T+N slant to field. 3-down: A+T+N slant. Best fronts: Under, BOSS Under. "
         "Do NOT call with Shade.")
    add_diagram("slant")
    doc.add_page_break()

    heading2("PINCH")
    body("4-down: T+N pinch A gaps. 3-down: A+N pinch B gaps. Best fronts: Shade, Under, Wide, Deuces.")
    add_diagram("pinch")
    doc.add_page_break()

    heading2("JACKS")
    body("4-down: T+N shoot B gaps. 3-down: A+N expand to C gaps (NOT pinch). "
         "Do NOT call with Wide. Best pressure: sWarM.")
    add_diagram("jacks")
    doc.add_page_break()

    heading2("CRASH")
    body("A+E shoot B gaps. T+N shoot A gaps. B/M/W are contain. "
         "Do NOT call with Grizzly or Freebird. Best pressure: BoW.")
    add_diagram("crash")
    doc.add_page_break()

    heading2("Other Stunts (Text Only)")
    bullet("Same as Slant but to boundary. Do NOT call with Under/BOSS Under. Best fronts: Shade, BOSS.", bold_prefix="ANGLE: ")
    bullet("With BOSS: 1-tech crosses C face to opposite A. With GRIZZLY: field-side gap-out.", bold_prefix="SPLIT: ")
    bullet("A attacks OL closing one inside gap. NOT paired with Cobra.", bold_prefix="ANCHOR ATTACK: ")
    bullet("E attacks OL closing one inside gap. CAN pair with Cobra.", bold_prefix="EDGE ATTACK: ")
    bullet("A shoots B gap from 5. Signal: point field \u2192 flap flap. If + Bandit = BANDIT RAVEN.", bold_prefix="ANCHOR RAVEN: ")
    bullet("E shoots B gap from 5. Signal: point boundary \u2192 flap flap.", bold_prefix="EDGE RAVEN: ")
    doc.add_page_break()

    # =========================================================================
    # PRESSURES
    # =========================================================================
    heading1("PRESSURE DIAGRAMS")

    heading2("sWarM (M + W)")
    add_diagram("swarm")
    doc.add_page_break()

    heading2("BooM (B + M)")
    add_diagram("boom")
    doc.add_page_break()

    heading2("BoW (B + W)")
    add_diagram("bow")
    doc.add_page_break()

    heading2("HAMMER (B Edge + Anchor Attack)")
    body("B blitzes off the edge. A attacks OL closing one inside gap.")
    add_diagram("hammer")
    doc.add_page_break()

    heading2("EAT (M + W + B)")
    add_diagram("eat")
    doc.add_page_break()

    heading2("Other Pressures (Text Only)")
    bullet("M blitzes to his gap.", bold_prefix="Mike: ")
    bullet("W blitzes to his gap.", bold_prefix="Will: ")
    bullet("B blitzes to his gap.", bold_prefix="Bandit: ")
    bullet("D blitzes to his gap.", bold_prefix="Dawg: ")
    bullet("B one gap inside EOL.", bold_prefix="staB: ")
    bullet("MOF safety blitzes.", bold_prefix="Freebird: ")
    bullet("Boundary corner blitz. D takes #1 bnd, FS assumes D rules, B assumes FS rules.", bold_prefix="Cobra: ")
    bullet("W edge + Edge Attack.", bold_prefix="Shave: ")
    bullet("M + D.", bold_prefix="MaD: ")

    doc.add_paragraph()
    heading2("Weekly Pressure Menu (Locked)")
    body("Mike, Will, Bandit, sWarM, BooM, BoW, Hammer, Eat.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run("\u2014 END OF DIAGRAM EDITION \u2014")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    doc.add_paragraph()
    g = doc.add_paragraph()
    g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = g.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    r.font.size = Pt(9); r.font.color.rgb = GRAY_RGB

    output = "/home/ksc4130/src/defensive_playbook/River_Valley_Vikings_Defensive_Playbook_DIAGRAMS.docx"
    doc.save(output)
    print(f"Diagram playbook saved to: {output}")
    return output


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("Generating diagrams...")
    paths = generate_all_diagrams()
    print(f"Generated {len(paths)} diagrams.")
    print("Building document...")
    build_document(paths)
