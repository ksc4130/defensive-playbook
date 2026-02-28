#!/usr/bin/env python3
"""Generate the River Valley Vikings Coverage Rules & Pre-Snap Disguise document."""

import os
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD_RGB = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_RGB = RGBColor(0x88, 0x88, 0x88)
NAVY_HEX = "002D62"

C_NAVY = "#002D62"
C_COLUMBIA = "#6CACE4"
C_GOLD = "#CFA700"
C_GRAY = "#888888"
C_OL = "#AAAAAA"
C_WR = "#666666"
C_WHITE = "#FFFFFF"

DIAGRAM_DIR = "/tmp/rv_coverage_diagrams"
os.makedirs(DIAGRAM_DIR, exist_ok=True)

# Positions
Y_LOS = 0
Y_DL = 1.2
Y_LB = 3.5
Y_OLB = 2.5
Y_APEX = 3.0
Y_CB = 5.5
Y_S = 8.0

OL_POSITIONS = [(0, 0), (2, 0), (-2, 0), (4, 0), (-4, 0)]


def new_fig(title="", figsize=(10, 7)):
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-16, 16)
    ax.set_ylim(-5, 12)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.axhline(y=0.5, color=C_GRAY, linewidth=1, linestyle="--", alpha=0.3)
    ax.text(14, 11, "FIELD \u2192", fontsize=8, color=C_GRAY, ha="right", va="top", style="italic")
    ax.text(-14, 11, "\u2190 BOUNDARY", fontsize=8, color=C_GRAY, ha="left", va="top", style="italic")
    if title:
        ax.set_title(title, fontsize=13, fontweight="bold", color=C_NAVY, pad=10)
    return fig, ax


def draw_ol(ax):
    for x, y in OL_POSITIONS:
        ax.add_patch(plt.Rectangle((x - 0.4, y - 0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", linewidth=1.5))


def draw_qb_rb(ax):
    ax.plot(0, -1.5, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(0, -2.2, "QB", fontsize=7, ha="center", color="#555")
    ax.plot(0, -3, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(0, -3.7, "RB", fontsize=7, ha="center", color="#555")


def draw_wr_2x2(ax):
    for pos, lbl in [((13, 0), "#1"), ((7, 0), "#2"), ((-13, 0), "#1"), ((-7, 0), "#2")]:
        ax.plot(*pos, "s", color=C_WR, markersize=7, markeredgecolor="#333")
        ax.text(pos[0], pos[1] - 0.7, lbl, fontsize=6, ha="center", color="#555")


def draw_wr_3x1(ax):
    for x, lbl in [(13, "#1"), (9, "#2"), (7, "#3")]:
        ax.plot(x, 0, "s", color=C_WR, markersize=7, markeredgecolor="#333")
        ax.text(x, -0.7, lbl, fontsize=6, ha="center", color="#555")
    ax.plot(-13, 0, "s", color=C_WR, markersize=7, markeredgecolor="#333")
    ax.text(-13, -0.7, "#1", fontsize=6, ha="center", color="#555")


def def_player(ax, x, y, label, color=C_NAVY, fontsize=8):
    ax.plot(x, y, "o", color=color, markersize=14, markeredgecolor=C_NAVY, markeredgewidth=1.2)
    ax.text(x, y, label, fontsize=fontsize, ha="center", va="center",
            color=C_WHITE, fontweight="bold")


def ghost_player(ax, x, y, label, color=C_NAVY):
    """Draw a ghost/destination position (dashed circle)."""
    circle = plt.Circle((x, y), 0.5, fill=False, edgecolor=color,
                         linewidth=1.5, linestyle="--")
    ax.add_patch(circle)
    ax.text(x, y, label, fontsize=7, ha="center", va="center", color=color, fontweight="bold")


def arrow(ax, x1, y1, x2, y2, color=C_NAVY, style="-|>", lw=1.8):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))


def curved_arrow(ax, x1, y1, x2, y2, color=C_NAVY):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=1.8,
                                connectionstyle="arc3,rad=0.3"))


def save_fig(fig, name):
    path = os.path.join(DIAGRAM_DIR, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def draw_pre_snap_shell():
    """The universal two-high pre-snap look."""
    fig, ax = new_fig("UNIVERSAL PRE-SNAP LOOK \u2014 Two-High Shell")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    # DL (Shade)
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    # LBs
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # DBs -- standard two-high
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    # Alignment notes
    ax.text(12, Y_CB + 1.0, "inside lev\n6 yds", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-12, Y_CB + 1.0, "inside lev\n6 yds", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(5, Y_S + 1.0, "10-12 yds\noff #2", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-5, Y_S + 1.0, "10-12 yds\noff #2", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, Y_APEX + 1.0, "apex\nfield", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(0, 11.5, "SAME LOOK \u2014 EVERY COVERAGE (except Zorro & Ohio)",
            fontsize=10, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.4", facecolor=C_COLUMBIA, alpha=0.2))

    return fig


def draw_ninja_post_snap():
    fig, ax = new_fig("NINJA (Cover 7) \u2014 Post-Snap vs 2x2")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # DBs stay in two-high -- no rotation needed
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    ax.axvline(x=0, color=C_GRAY, linewidth=1, linestyle=":", alpha=0.3)
    ax.text(7, 11, "MOD", fontsize=11, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor=C_COLUMBIA, alpha=0.25))
    ax.text(-7, 11, "CLAMP", fontsize=11, ha="center", color=C_NAVY, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor=C_GOLD, alpha=0.25))

    # Annotations
    ax.text(12, Y_CB - 1.2, "man-match #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(5, Y_S - 1.2, "top-down #2", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, Y_APEX - 1.2, "#2/#3 threats", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-12, Y_CB - 1.2, "clamp #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-5, Y_S - 1.2, "control #2", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(0, -4.5, "No rotation needed \u2014 pre-snap look IS the coverage", fontsize=9,
            ha="center", color=C_NAVY, fontweight="bold")

    return fig


def draw_ninja_3x1_post_snap():
    fig, ax = new_fig("NINJA (Cover 7) \u2014 Post-Snap vs 3x1 (POACH)")
    draw_ol(ax)
    draw_wr_3x1(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # Pre-snap: same two-high. Post-snap: D poaches toward trips
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    # D poach arrow toward trips
    ghost_player(ax, -1, 7.5, "D", color=C_GOLD)
    curved_arrow(ax, -5, Y_S, -1, 7.5, color=C_GOLD)

    ax.text(-1, 8.5, "POACH\n#3 vertical?", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0, -4.5, "D poaches toward trips to field. FS holds field half.",
            fontsize=9, ha="center", color=C_NAVY, fontweight="bold")

    return fig


def draw_oregon_post_snap():
    fig, ax = new_fig("OREGON (Cover 1, Post = FS) \u2014 Post-Snap Drop")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    # Pre-snap positions (solid)
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # Post-snap: FS stays deep/rotates to MOF post
    ghost_player(ax, 0, 10, "FS", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10, color=C_GOLD)

    # D drops to man #2 boundary
    ghost_player(ax, -7, 6, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    # B drops to man #2 field
    ghost_player(ax, 7, 4.5, "B", color=C_COLUMBIA)
    arrow(ax, 7.5, Y_APEX, 7, 4.5, color=C_COLUMBIA)

    ax.text(0, 11, "POST (MOF)", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
    ax.text(-7.5, 5.2, "man #2 bnd", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, 3.7, "man #2 field", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(1.5, Y_LB - 1.2, "RB funnel", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-1.5, Y_LB - 1.2, "RB funnel", fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_oklahoma_post_snap():
    fig, ax = new_fig("OKLAHOMA (Cover 1, Post = D) \u2014 Post-Snap Drop")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    # Oklahoma: FS and D switch sides pre-snap
    def_player(ax, 5, Y_S, "D", color=C_GOLD)     # D to field
    def_player(ax, -5, Y_S, "FS", color=C_GOLD)    # FS to boundary
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # Post-snap: D rotates from field to MOF post
    ghost_player(ax, 0, 10, "D", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10, color=C_GOLD)

    # FS drops from boundary to man #2 boundary
    ghost_player(ax, -7, 6, "FS", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    # B drops to man #2 field
    ghost_player(ax, 7, 4.5, "B", color=C_COLUMBIA)
    arrow(ax, 7.5, Y_APEX, 7, 4.5, color=C_COLUMBIA)

    ax.text(0, 11, "POST (MOF)", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
    ax.text(-7.5, 5.2, "FS: man #2 bnd", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, 3.7, "man #2 field", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(5, Y_S + 1.0, "D starts\nfield", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-5, Y_S + 1.0, "FS starts\nboundary", fontsize=7, ha="center", color=C_NAVY, style="italic")

    return fig


def draw_ohio_exception():
    """Ohio is an exception -- B leaves apex to become post."""
    fig, ax = new_fig("OHIO (Cover 1, Post = B) \u2014 EXCEPTION")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)
    # B pre-snap: ~8 yds deep in the middle of the formation (exception alignment)
    def_player(ax, 0, 8, "B", color=C_COLUMBIA)

    # B rotates to MOF post
    ghost_player(ax, 0, 10, "B", color=C_COLUMBIA)
    arrow(ax, 0, 8, 0, 10, color=C_COLUMBIA)

    # FS drops to man #2 field
    ghost_player(ax, 7, 6, "FS", color=C_GOLD)
    arrow(ax, 5, Y_S, 7, 6, color=C_GOLD)

    # D drops to man #2 boundary
    ghost_player(ax, -7, 6, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    ax.text(0, 11, "POST (MOF)", fontsize=8, ha="center", color=C_NAVY, fontweight="bold")
    ax.text(7.5, 5.2, "FS: man #2 field", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-7.5, 5.2, "D: man #2 bnd", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0, 6.8, "B: ~8 yds\nmiddle", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(0, -4.5, "EXCEPTION: B at ~8 yds in middle, becomes MOF \u2014 changes pre-snap picture",
            fontsize=9, ha="center", color="#990000", fontweight="bold")

    return fig


def draw_zeus_post_snap():
    fig, ax = new_fig("ZEUS (Cover 0) \u2014 Post-Snap Drop")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # Post-snap: safeties drop to man #2
    ghost_player(ax, 7, 6, "FS", color=C_GOLD)
    arrow(ax, 5, Y_S, 7, 6, color=C_GOLD)
    ghost_player(ax, -7, 6, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    ax.text(7.5, 5.2, "man #2 field", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-7.5, 5.2, "man #2 bnd", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-12, Y_CB - 1.2, "man #1", fontsize=7, ha="center", color=C_NAVY, style="italic")

    # M rushes, B/W funnel
    arrow(ax, 1.5, Y_LB, 0, Y_LB - 2, color=C_COLUMBIA)
    ax.text(0.5, Y_LB - 0.8, "rush", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, Y_APEX - 1.2, "RB funnel", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-1.5, Y_LB - 1.2, "RB funnel", fontsize=7, ha="center", color=C_NAVY, style="italic")

    # DL rush
    arrow(ax, 4.7, Y_DL, 4.2, Y_DL - 1.5)
    arrow(ax, 2.5, Y_DL, 1.5, Y_DL - 1.5)
    arrow(ax, -1.5, Y_DL, -0.8, Y_DL - 1.5)
    arrow(ax, -4.7, Y_DL, -4.2, Y_DL - 1.5)

    ax.text(0, -4.5, "Same two-high pre-snap. Post-snap: safeties drop to man, no deep help.",
            fontsize=9, ha="center", color=C_NAVY, fontweight="bold")

    return fig


def draw_zorro_exception():
    """Zorro is an exception -- B has RB, changes his alignment."""
    fig, ax = new_fig("ZORRO (Cover 0) \u2014 EXCEPTION")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)

    # B pre-snap: ~8 yds deep in the middle of the formation (exception alignment)
    def_player(ax, 0, 8, "B", color=C_COLUMBIA)

    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    # Safeties drop to man
    ghost_player(ax, 7, 6, "FS", color=C_GOLD)
    arrow(ax, 5, Y_S, 7, 6, color=C_GOLD)
    ghost_player(ax, -7, 6, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    # B keys RB from middle
    arrow(ax, 0, 8, 0, -2.5, color=C_COLUMBIA, lw=1.5)
    ax.text(0, 6.8, "B: ~8 yds middle\nhas RB", fontsize=7, ha="center",
            color=C_NAVY, style="italic")
    ax.text(7.5, 5.2, "man #2 field", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-7.5, 5.2, "man #2 bnd", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(0, -4.5, "EXCEPTION: B at ~8 yds in middle, keys RB \u2014 changes pre-snap picture",
            fontsize=9, ha="center", color="#990000", fontweight="bold")

    return fig


def draw_zunnel_post_snap():
    fig, ax = new_fig("ZUNNEL (Cover 0) \u2014 Post-Snap Drop")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    ghost_player(ax, 7, 6, "FS", color=C_GOLD)
    arrow(ax, 5, Y_S, 7, 6, color=C_GOLD)
    ghost_player(ax, -7, 6, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -7, 6, color=C_GOLD)

    ax.text(1.5, Y_LB - 1.2, "M/W funnel RB", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, Y_APEX - 1.2, "B = man #2\nor edge", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0, -4.5, "Same pre-snap look. M/W read RB: to you = take, away = rush.",
            fontsize=9, ha="center", color=C_NAVY, fontweight="bold")

    return fig


def draw_viking_post_snap():
    fig, ax = new_fig("VIKING (Cover 3) \u2014 Post-Snap Drop")
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)

    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")

    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)

    # Pre-snap
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)

    # Post-snap: FC bails to deep 1/3, FS rotates to deep middle 1/3, BC bails to deep 1/3
    ghost_player(ax, 11, 10.5, "FC", color=C_COLUMBIA)
    arrow(ax, 12, Y_CB, 11, 10.5, color=C_COLUMBIA)
    ghost_player(ax, 0, 10.5, "FS", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10.5, color=C_GOLD)
    ghost_player(ax, -11, 10.5, "BC", color=C_COLUMBIA)
    arrow(ax, -12, Y_CB, -11, 10.5, color=C_COLUMBIA)

    # D drops to seam-curl-flat
    ghost_player(ax, -6, 5.5, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -6, 5.5, color=C_GOLD)

    ax.text(11, 11.3, "deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(0, 11.3, "deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-11, 11.3, "deep 1/3", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-6.5, 4.7, "seam-curl-flat", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(7.5, Y_APEX - 1.2, "curl/flat", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(1.5, Y_LB - 1.2, "hook", fontsize=7, ha="center", color=C_NAVY, style="italic")
    ax.text(-1.5, Y_LB - 1.2, "hook", fontsize=7, ha="center", color=C_NAVY, style="italic")

    ax.text(0, -4.5, "Same pre-snap look. Post-snap: corners bail deep, FS rotates to middle 1/3.",
            fontsize=9, ha="center", color=C_NAVY, fontweight="bold")

    return fig


# =============================================================================
# CONCEPT DIAGRAM HELPERS
# =============================================================================

C_ROUTE = "#CC3333"


def draw_route_path(ax, points):
    """Draw an offensive route as dashed line with arrowhead."""
    for i in range(len(points) - 1):
        if i < len(points) - 2:
            ax.plot([points[i][0], points[i + 1][0]],
                    [points[i][1], points[i + 1][1]],
                    "--", color=C_ROUTE, lw=1.3, alpha=0.7)
        else:
            ax.annotate("", xy=points[i + 1], xytext=points[i],
                        arrowprops=dict(arrowstyle="-|>", color=C_ROUTE,
                                        lw=1.3, linestyle="dashed"))


def rlbl(ax, x, y, text):
    """Route label."""
    ax.text(x, y, text, fontsize=6, ha="center", color=C_ROUTE, style="italic")


def nt(ax, x, y, text):
    """Defender annotation note."""
    ax.text(x, y, text, fontsize=6.5, ha="center", color=C_NAVY, style="italic")


def bn(ax, text):
    """Bottom summary note."""
    ax.text(0, -4.5, text, fontsize=8, ha="center", color=C_NAVY, fontweight="bold")


def draw_dl_only(ax):
    def_player(ax, 4.7, Y_DL, "A")
    def_player(ax, 2.5, Y_DL, "T")
    def_player(ax, -1.5, Y_DL, "N")
    def_player(ax, -4.7, Y_DL, "E")


def std_def(ax):
    """Standard pre-snap defensive look."""
    draw_dl_only(ax)
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "FS", color=C_GOLD)
    def_player(ax, -5, Y_S, "D", color=C_GOLD)


def ok_def(ax):
    """Oklahoma pre-snap: FS and D switch sides."""
    draw_dl_only(ax)
    def_player(ax, 1.5, Y_LB, "M", color=C_COLUMBIA)
    def_player(ax, -1.5, Y_LB, "W", color=C_COLUMBIA)
    def_player(ax, 7.5, Y_APEX, "B", color=C_COLUMBIA)
    def_player(ax, 12, Y_CB, "FC", color=C_COLUMBIA)
    def_player(ax, -12, Y_CB, "BC", color=C_COLUMBIA)
    def_player(ax, 5, Y_S, "D", color=C_GOLD)
    def_player(ax, -5, Y_S, "FS", color=C_GOLD)


def or_drops(ax, d_tgt, b_tgt):
    """Oregon post-snap: FS->MOF, D->target, B->target."""
    ghost_player(ax, 0, 10, "FS", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10, color=C_GOLD)
    ghost_player(ax, d_tgt[0], d_tgt[1], "D", color=C_GOLD)
    arrow(ax, -5, Y_S, d_tgt[0], d_tgt[1], color=C_GOLD)
    ghost_player(ax, b_tgt[0], b_tgt[1], "B", color=C_COLUMBIA)
    arrow(ax, 7.5, Y_APEX, b_tgt[0], b_tgt[1], color=C_COLUMBIA)


def ok_drops(ax, fs_tgt, b_tgt):
    """Oklahoma post-snap: D(field)->MOF, FS(bnd)->target, B->target."""
    ghost_player(ax, 0, 10, "D", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10, color=C_GOLD)
    ghost_player(ax, fs_tgt[0], fs_tgt[1], "FS", color=C_GOLD)
    arrow(ax, -5, Y_S, fs_tgt[0], fs_tgt[1], color=C_GOLD)
    ghost_player(ax, b_tgt[0], b_tgt[1], "B", color=C_COLUMBIA)
    arrow(ax, 7.5, Y_APEX, b_tgt[0], b_tgt[1], color=C_COLUMBIA)


def z_drops(ax, fs_tgt, d_tgt):
    """Zeus post-snap: FS->target, D->target, M rushes."""
    ghost_player(ax, fs_tgt[0], fs_tgt[1], "FS", color=C_GOLD)
    arrow(ax, 5, Y_S, fs_tgt[0], fs_tgt[1], color=C_GOLD)
    ghost_player(ax, d_tgt[0], d_tgt[1], "D", color=C_GOLD)
    arrow(ax, -5, Y_S, d_tgt[0], d_tgt[1], color=C_GOLD)
    arrow(ax, 1.5, Y_LB, 0.5, Y_LB - 2, color=C_COLUMBIA)


def vik_drops(ax):
    """Viking post-snap: FC/BC deep 1/3, FS deep mid, D seam-curl-flat."""
    ghost_player(ax, 11, 10.5, "FC", color=C_COLUMBIA)
    arrow(ax, 12, Y_CB, 11, 10.5, color=C_COLUMBIA)
    ghost_player(ax, -11, 10.5, "BC", color=C_COLUMBIA)
    arrow(ax, -12, Y_CB, -11, 10.5, color=C_COLUMBIA)
    ghost_player(ax, 0, 10.5, "FS", color=C_GOLD)
    curved_arrow(ax, 5, Y_S, 0, 10.5, color=C_GOLD)
    ghost_player(ax, -6, 5.5, "D", color=C_GOLD)
    arrow(ax, -5, Y_S, -6, 5.5, color=C_GOLD)


def concept_base(ax):
    """Draw OL, QB, RB, WRs for 2x2."""
    draw_ol(ax)
    draw_wr_2x2(ax)
    draw_qb_rb(ax)


# =============================================================================
# CONCEPT DIAGRAM GENERATORS (10 concepts x 5 coverages = 50 diagrams)
# =============================================================================

def gen_4verts():
    """Four Verticals vs all coverages."""
    R = {}

    def routes(ax):
        for pts in [[(13, 0), (13, 10)], [(7, 0), (7, 10)],
                    [(-13, 0), (-13, 10)], [(-7, 0), (-7, 10)]]:
            draw_route_path(ax, pts)
        for x, y, t in [(13.5, 10.5, "Go"), (7.5, 10.5, "Seam"),
                        (-12.5, 10.5, "Go"), (-6.5, 10.5, "Seam")]:
            rlbl(ax, x, y, t)

    # NINJA
    fig, ax = new_fig("Four Verticals vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "matches\n#1 vert")
    nt(ax, -12, Y_CB + 1, "clamps\n#1 vert")
    nt(ax, 5, Y_S + 1, "top-down\n#2 seam")
    nt(ax, -5, Y_S + 1, "controls\n#2 seam")
    nt(ax, 7.5, Y_APEX + 1, "carries\n#2/#3")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Two-high caps all seams. Best zone look vs 4 verts.")
    R["4v_ninja"] = save_fig(fig, "4v_ninja")

    # OREGON
    fig, ax = new_fig("Four Verticals vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 8), (7, 7))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, -8, 7.2, "D: man\n#2B seam")
    nt(ax, 8, 6.2, "B: man\n#2F seam")
    nt(ax, 1.5, Y_LB + 1, "RB funnel")
    nt(ax, -1.5, Y_LB + 1, "RB funnel")
    bn(ax, "FS as post helps on seams but must choose a side.")
    R["4v_oregon"] = save_fig(fig, "4v_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Four Verticals vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 8), (7, 7))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, -8, 7.2, "FS: man\n#2B seam")
    nt(ax, 8, 6.2, "B: man\n#2F seam")
    nt(ax, 1.5, Y_LB + 1, "RB funnel")
    nt(ax, -1.5, Y_LB + 1, "RB funnel")
    bn(ax, "D (field) rotates to post. FS (bnd) stays with #2B seam.")
    R["4v_oklahoma"] = save_fig(fig, "4v_oklahoma")

    # ZEUS
    fig, ax = new_fig("Four Verticals vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (7, 8), (-7, 8))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 8, 7.2, "FS: man\n#2F seam")
    nt(ax, -8, 7.2, "D: man\n#2B seam")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    nt(ax, 7.5, Y_APEX + 1, "B: RB\nfunnel")
    nt(ax, -1.5, Y_LB + 1, "W: RB\nfunnel")
    bn(ax, "No help deep. DL must get pressure before routes develop.")
    R["4v_zeus"] = save_fig(fig, "4v_zeus")

    # VIKING
    fig, ax = new_fig("Four Verticals vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid\n1/3")
    nt(ax, -7, 4.7, "D: wall\n#2B seam")
    nt(ax, 7.5, Y_APEX + 1, "B: wall\n#2F seam")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "3 deep zones cap all verts. FS splits seams in middle 1/3.")
    R["4v_viking"] = save_fig(fig, "4v_viking")

    return R


def gen_mesh():
    """Mesh / Crossers vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 3), (-3, 5.5)])
        draw_route_path(ax, [(-13, 0), (-13, 3), (3, 5)])
        draw_route_path(ax, [(7, 0), (7, 8)])
        draw_route_path(ax, [(-7, 0), (-7, 8)])
        rlbl(ax, -4, 6.2, "#1F cross")
        rlbl(ax, 4, 5.7, "#1B cross")
        rlbl(ax, 7.5, 8.5, "#2F vert")
        rlbl(ax, -7.5, 8.5, "#2B vert")

    # NINJA
    fig, ax = new_fig("Mesh / Crossers vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "passes off\ncrosser")
    nt(ax, -12, Y_CB + 1, "passes off\ncrosser")
    nt(ax, 5, Y_S + 1, "top-down\n#2 vert")
    nt(ax, -5, Y_S + 1, "controls\n#2 vert")
    nt(ax, 7.5, Y_APEX + 1, "walls\ncrosser")
    nt(ax, 1.5, Y_LB + 1, "collision\ncrosser")
    nt(ax, -1.5, Y_LB + 1, "collision\ncrosser")
    bn(ax, "Zone handles picks naturally. Crossers run into zone drops.")
    R["mesh_ninja"] = save_fig(fig, "mesh_ninja")

    # OREGON
    fig, ax = new_fig("Mesh / Crossers vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 7), (7, 6))
    ghost_player(ax, -3, 5.5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, -3, 5.5, color=C_COLUMBIA)
    ghost_player(ax, 3, 5, "BC", color=C_COLUMBIA)
    curved_arrow(ax, -12, Y_CB, 3, 5, color=C_COLUMBIA)
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, -8, 6.2, "D: man\n#2B vert")
    nt(ax, 8, 5.2, "B: man\n#2F vert")
    nt(ax, -4, 4.5, "FC follows\n#1F cross")
    nt(ax, 4, 4, "BC follows\n#1B cross")
    bn(ax, "Picks/rubs trouble man coverage. BANJO helps at LOS.")
    R["mesh_oregon"] = save_fig(fig, "mesh_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Mesh / Crossers vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 7), (7, 6))
    ghost_player(ax, -3, 5.5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, -3, 5.5, color=C_COLUMBIA)
    ghost_player(ax, 3, 5, "BC", color=C_COLUMBIA)
    curved_arrow(ax, -12, Y_CB, 3, 5, color=C_COLUMBIA)
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, -8, 6.2, "FS: man\n#2B vert")
    nt(ax, 8, 5.2, "B: man\n#2F vert")
    bn(ax, "D (field) to post. CBs must fight through mesh picks.")
    R["mesh_oklahoma"] = save_fig(fig, "mesh_oklahoma")

    # ZEUS
    fig, ax = new_fig("Mesh / Crossers vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (7, 7), (-7, 7))
    ghost_player(ax, -3, 5.5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, -3, 5.5, color=C_COLUMBIA)
    ghost_player(ax, 3, 5, "BC", color=C_COLUMBIA)
    curved_arrow(ax, -12, Y_CB, 3, 5, color=C_COLUMBIA)
    nt(ax, 8, 6.2, "FS: man\n#2F vert")
    nt(ax, -8, 6.2, "D: man\n#2B vert")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Most vulnerable vs mesh. Picks + no help = need instant pressure.")
    R["mesh_zeus"] = save_fig(fig, "mesh_zeus")

    # VIKING
    fig, ax = new_fig("Mesh / Crossers vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: walls\ncrosser")
    nt(ax, 7.5, Y_APEX + 1, "B: walls\ncrosser")
    nt(ax, 1.5, Y_LB + 1, "collision\ncrosser")
    nt(ax, -1.5, Y_LB + 1, "collision\ncrosser")
    bn(ax, "Zone handles picks. 3 deep prevents shots over mesh bait.")
    R["mesh_viking"] = save_fig(fig, "mesh_viking")

    return R


def gen_smash():
    """Smash (Corner/Hitch) vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 5)])
        draw_route_path(ax, [(7, 0), (7, 6), (10, 9.5)])
        draw_route_path(ax, [(-13, 0), (-13, 5)])
        draw_route_path(ax, [(-7, 0), (-7, 6), (-10, 9.5)])
        rlbl(ax, 14, 5, "Hitch")
        rlbl(ax, 11, 9.5, "Corner")
        rlbl(ax, -14, 5, "Hitch")
        rlbl(ax, -11, 9.5, "Corner")

    # NINJA
    fig, ax = new_fig("Smash vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "sits on\nhitch")
    nt(ax, -12, Y_CB + 1, "clamps\nhitch")
    nt(ax, 5, Y_S + 1, "over top\nof corner")
    nt(ax, -5, Y_S + 1, "controls\n#2 corner")
    nt(ax, 7.5, Y_APEX + 1, "under\ncorner")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "FS already over #2 \u2014 corner route runs into his coverage.")
    R["smash_ninja"] = save_fig(fig, "smash_ninja")

    # OREGON
    fig, ax = new_fig("Smash vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-10, 8), (10, 8))
    nt(ax, 12, Y_CB + 1, "man #1\nhitch")
    nt(ax, -12, Y_CB + 1, "man #1\nhitch")
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, -11, 7.2, "D: man\n#2B corner")
    nt(ax, 11, 7.2, "B: man\n#2F corner")
    bn(ax, "Man handles smash well. Post safety helps on corner if needed.")
    R["smash_oregon"] = save_fig(fig, "smash_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Smash vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-10, 8), (10, 8))
    nt(ax, 12, Y_CB + 1, "man #1\nhitch")
    nt(ax, -12, Y_CB + 1, "man #1\nhitch")
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, -11, 7.2, "FS: man\n#2B corner")
    nt(ax, 11, 7.2, "B: man\n#2F corner")
    bn(ax, "D (field) to post. Man matches both levels cleanly.")
    R["smash_oklahoma"] = save_fig(fig, "smash_oklahoma")

    # ZEUS
    fig, ax = new_fig("Smash vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (10, 8), (-10, 8))
    nt(ax, 12, Y_CB + 1, "man #1\nhitch")
    nt(ax, -12, Y_CB + 1, "man #1\nhitch")
    nt(ax, 11, 7.2, "FS: man\n#2F corner")
    nt(ax, -11, 7.2, "D: man\n#2B corner")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Man on hitch easy. Corner route has no help over top.")
    R["smash_zeus"] = save_fig(fig, "smash_zeus")

    # VIKING
    fig, ax = new_fig("Smash vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3\ncaps corner")
    nt(ax, -11.5, 11.3, "deep 1/3\ncaps corner")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: sits on\nhitch")
    nt(ax, 7.5, Y_APEX + 1, "B: sits on\nhitch")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Best zone look vs smash. Deep 1/3 caps corner, flat covers hitch.")
    R["smash_viking"] = save_fig(fig, "smash_viking")

    return R


def gen_flood():
    """Flood / Sail vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 10)])
        draw_route_path(ax, [(7, 0), (7, 5), (10, 7.5)])
        draw_route_path(ax, [(0, -3), (5, 0), (10, 2)])
        draw_route_path(ax, [(-13, 0), (-13, 8)])
        draw_route_path(ax, [(-7, 0), (-7, 5)])
        rlbl(ax, 14, 10, "Go")
        rlbl(ax, 11, 7.5, "Out")
        rlbl(ax, 11, 2, "RB flat")
        rlbl(ax, -14, 8, "Clear")
        rlbl(ax, -7.5, 5.5, "Sit")

    # NINJA
    fig, ax = new_fig("Flood / Sail vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "matches\n#1 go")
    nt(ax, -12, Y_CB + 1, "clamps\n#1 clear")
    nt(ax, 5, Y_S + 1, "takes\ndeepest")
    nt(ax, -5, Y_S + 1, "controls\n#2 sit")
    nt(ax, 7.5, Y_APEX + 1, "walls\nintermediate")
    nt(ax, 1.5, Y_LB + 1, "trigger\nto flat")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "3rd level (flat) can stress B if he's slow. M must trigger.")
    R["flood_ninja"] = save_fig(fig, "flood_ninja")

    # OREGON
    fig, ax = new_fig("Flood / Sail vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 5), (10, 6))
    nt(ax, 12, Y_CB + 1, "man #1\ngo")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "post/MOF\nleans flood")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 11, 5.2, "B: man\n#2F out")
    nt(ax, 1.5, Y_LB + 1, "RB funnel\nor take")
    bn(ax, "Post safety leans toward flood. Man stays with each route.")
    R["flood_oregon"] = save_fig(fig, "flood_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Flood / Sail vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 5), (10, 6))
    nt(ax, 12, Y_CB + 1, "man #1\ngo")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "D: post/MOF\nleans flood")
    nt(ax, -8, 4.2, "FS: man\n#2B sit")
    nt(ax, 11, 5.2, "B: man\n#2F out")
    bn(ax, "D (field) to post, leans to flood side. FS covers #2B.")
    R["flood_oklahoma"] = save_fig(fig, "flood_oklahoma")

    # ZEUS
    fig, ax = new_fig("Flood / Sail vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (10, 6), (-7, 5))
    nt(ax, 12, Y_CB + 1, "man #1\ngo")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 11, 5.2, "FS: man\n#2F out")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    nt(ax, 7.5, Y_APEX + 1, "B: RB\nfunnel/take")
    bn(ax, "Man runs with each level. No conflict for flat defender.")
    R["flood_zeus"] = save_fig(fig, "flood_zeus")

    # VIKING
    fig, ax = new_fig("Flood / Sail vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3\ncaps go")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: seam\ncurl-flat")
    nt(ax, 7.5, Y_APEX + 1, "B: drives\nto out/flat")
    nt(ax, 1.5, Y_LB + 1, "trigger\nto flat")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Deep 1/3 caps go. Flat can be vulnerable if B is out-leveraged.")
    R["flood_viking"] = save_fig(fig, "flood_viking")

    return R


def gen_levels():
    """Levels (Hi-Lo Crossers) vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 3), (0, 5)])
        draw_route_path(ax, [(7, 0), (7, 6), (0, 9)])
        draw_route_path(ax, [(-13, 0), (-13, 9)])
        draw_route_path(ax, [(-7, 0), (-7, 5)])
        rlbl(ax, 0.5, 5.7, "Drag")
        rlbl(ax, 0.5, 9.7, "Dig")
        rlbl(ax, -14, 9, "Clear")
        rlbl(ax, -7.5, 5.5, "Sit")

    # NINJA
    fig, ax = new_fig("Levels vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "passes off\ndrag")
    nt(ax, -12, Y_CB + 1, "clamps\n#1 clear")
    nt(ax, 5, Y_S + 1, "relates to\ndig crosser")
    nt(ax, -5, Y_S + 1, "controls\n#2 sit")
    nt(ax, 7.5, Y_APEX + 1, "walls\ndrag")
    nt(ax, 1.5, Y_LB + 1, "sits in\nhi-lo")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Zone reads hi-lo. Defenders play in between. No picks.")
    R["levels_ninja"] = save_fig(fig, "levels_ninja")

    # OREGON
    fig, ax = new_fig("Levels vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 5), (0, 8))
    ghost_player(ax, 0, 5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, 0, 5, color=C_COLUMBIA)
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "post/MOF\nhelps dig")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 1, 7.2, "B: man\n#2F dig")
    nt(ax, 0.5, 4, "FC follows\n#1F drag")
    bn(ax, "FC chases drag across field. Long chase = potential big play.")
    R["levels_oregon"] = save_fig(fig, "levels_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Levels vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 5), (0, 8))
    ghost_player(ax, 0, 5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, 0, 5, color=C_COLUMBIA)
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "D: post/MOF\nhelps dig")
    nt(ax, -8, 4.2, "FS: man\n#2B sit")
    nt(ax, 1, 7.2, "B: man\n#2F dig")
    bn(ax, "D (field) to post. FC chases drag \u2014 long man trail.")
    R["levels_oklahoma"] = save_fig(fig, "levels_oklahoma")

    # ZEUS
    fig, ax = new_fig("Levels vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (0, 8), (-7, 5))
    ghost_player(ax, 0, 5, "FC", color=C_COLUMBIA)
    curved_arrow(ax, 12, Y_CB, 0, 5, color=C_COLUMBIA)
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 1, 7.2, "FS: man\n#2F dig")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    nt(ax, 0.5, 4, "FC trails\ndrag")
    bn(ax, "Long crosses + no help = vulnerable. Need pressure now.")
    R["levels_zeus"] = save_fig(fig, "levels_zeus")

    # VIKING
    fig, ax = new_fig("Levels vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: walls\ndig crosser")
    nt(ax, 7.5, Y_APEX + 1, "B: walls\ndrag")
    nt(ax, 1.5, Y_LB + 1, "collision\ndrag")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Zone layers handle hi-lo naturally. Crossers run into drops.")
    R["levels_viking"] = save_fig(fig, "levels_viking")

    return R


def gen_curl_flat():
    """Curl / Flat vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 10), (12, 9.5)])
        draw_route_path(ax, [(7, 0), (10, 3)])
        draw_route_path(ax, [(-13, 0), (-13, 10), (-12, 9.5)])
        draw_route_path(ax, [(-7, 0), (-10, 3)])
        rlbl(ax, 13.5, 10.5, "Curl")
        rlbl(ax, 11, 3.5, "Flat")
        rlbl(ax, -13.5, 10.5, "Curl")
        rlbl(ax, -11, 3.5, "Flat")

    # NINJA
    fig, ax = new_fig("Curl / Flat vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "sits on\ncurl")
    nt(ax, -12, Y_CB + 1, "clamps\ncurl")
    nt(ax, 5, Y_S + 1, "over top\nof #2")
    nt(ax, -5, Y_S + 1, "controls\n#2")
    nt(ax, 7.5, Y_APEX + 1, "drives\nto flat")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Two-high covers both levels. No conflict for flat defender.")
    R["cf_ninja"] = save_fig(fig, "cf_ninja")

    # OREGON
    fig, ax = new_fig("Curl / Flat vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-10, 3.5), (10, 3.5))
    nt(ax, 12, Y_CB + 1, "man #1\ncurl")
    nt(ax, -12, Y_CB + 1, "man #1\ncurl")
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, -11, 2.7, "D: man\n#2B flat")
    nt(ax, 11, 2.7, "B: man\n#2F flat")
    bn(ax, "Man handles curl/flat cleanly. CB on curl, B/D on flat.")
    R["cf_oregon"] = save_fig(fig, "cf_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Curl / Flat vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-10, 3.5), (10, 3.5))
    nt(ax, 12, Y_CB + 1, "man #1\ncurl")
    nt(ax, -12, Y_CB + 1, "man #1\ncurl")
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, -11, 2.7, "FS: man\n#2B flat")
    nt(ax, 11, 2.7, "B: man\n#2F flat")
    bn(ax, "D (field) to post. FS covers #2B flat from boundary side.")
    R["cf_oklahoma"] = save_fig(fig, "cf_oklahoma")

    # ZEUS
    fig, ax = new_fig("Curl / Flat vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (10, 3.5), (-10, 3.5))
    nt(ax, 12, Y_CB + 1, "man #1\ncurl")
    nt(ax, -12, Y_CB + 1, "man #1\ncurl")
    nt(ax, 11, 2.7, "FS: man\n#2F flat")
    nt(ax, -11, 2.7, "D: man\n#2B flat")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Man on each level. Pressure forces quick throw.")
    R["cf_zeus"] = save_fig(fig, "cf_zeus")

    # VIKING
    fig, ax = new_fig("Curl / Flat vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3\ndrives curl")
    nt(ax, -11.5, 11.3, "deep 1/3\ndrives curl")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: curl or\nflat?")
    nt(ax, 7.5, Y_APEX + 1, "B: curl or\nflat?")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Classic Cover 3 conflict. B/D must choose curl or flat.")
    R["cf_viking"] = save_fig(fig, "cf_viking")

    return R


def gen_slant_flat():
    """Slant / Flat vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (10, 4.5)])
        draw_route_path(ax, [(7, 0), (10, 2)])
        draw_route_path(ax, [(-13, 0), (-10, 4.5)])
        draw_route_path(ax, [(-7, 0), (-10, 2)])
        rlbl(ax, 10.5, 5, "Slant")
        rlbl(ax, 11, 2.5, "Flat")
        rlbl(ax, -10.5, 5, "Slant")
        rlbl(ax, -11, 2.5, "Flat")

    # NINJA
    fig, ax = new_fig("Slant / Flat vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "inside lev\ncollision slant")
    nt(ax, -12, Y_CB + 1, "clamp\ncollision slant")
    nt(ax, 5, Y_S + 1, "triggers\nto flat")
    nt(ax, -5, Y_S + 1, "triggers\nto flat")
    nt(ax, 7.5, Y_APEX + 1, "reads\n#2 flat")
    nt(ax, 1.5, Y_LB + 1, "hook/wall\nslant")
    nt(ax, -1.5, Y_LB + 1, "hook/wall\nslant")
    bn(ax, "Inside leverage by CBs = perfect counter to slants.")
    R["sf_ninja"] = save_fig(fig, "sf_ninja")

    # OREGON
    fig, ax = new_fig("Slant / Flat vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-10, 2.5), (10, 2.5))
    nt(ax, 12, Y_CB + 1, "man #1\nslant")
    nt(ax, -12, Y_CB + 1, "man #1\nslant")
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, -11, 1.7, "D: man\n#2B flat")
    nt(ax, 11, 1.7, "B: man\n#2F flat")
    bn(ax, "CBs physical at LOS on slant. Man handles it cleanly.")
    R["sf_oregon"] = save_fig(fig, "sf_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Slant / Flat vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-10, 2.5), (10, 2.5))
    nt(ax, 12, Y_CB + 1, "man #1\nslant")
    nt(ax, -12, Y_CB + 1, "man #1\nslant")
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, -11, 1.7, "FS: man\n#2B flat")
    nt(ax, 11, 1.7, "B: man\n#2F flat")
    bn(ax, "D (field) to post. CBs collision slant at line.")
    R["sf_oklahoma"] = save_fig(fig, "sf_oklahoma")

    # ZEUS
    fig, ax = new_fig("Slant / Flat vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (10, 2.5), (-10, 2.5))
    nt(ax, 12, Y_CB + 1, "man #1\nslant")
    nt(ax, -12, Y_CB + 1, "man #1\nslant")
    nt(ax, 11, 1.7, "FS: man\n#2F flat")
    nt(ax, -11, 1.7, "D: man\n#2B flat")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Physical CB at LOS kills slant. Pressure forces quick throw.")
    R["sf_zeus"] = save_fig(fig, "sf_zeus")

    # VIKING
    fig, ax = new_fig("Slant / Flat vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "bails deep\nslant under")
    nt(ax, -11.5, 11.3, "bails deep\nslant under")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: triggers\nto flat")
    nt(ax, 7.5, Y_APEX + 1, "B: triggers\nto flat")
    nt(ax, 1.5, Y_LB + 1, "wall\nslant")
    nt(ax, -1.5, Y_LB + 1, "wall\nslant")
    bn(ax, "CBs bail deep = slant underneath. LBs must wall it.")
    R["sf_viking"] = save_fig(fig, "sf_viking")

    return R


def gen_rpo():
    """RPO Concepts vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(7, 0), (11, 1)])
        draw_route_path(ax, [(13, 0), (13, 2)])
        draw_route_path(ax, [(-13, 0), (-13, 2)])
        draw_route_path(ax, [(-7, 0), (-7, 2)])
        ax.annotate("", xy=(0, -1), xytext=(0, -3),
                    arrowprops=dict(arrowstyle="-|>", color=C_WR, lw=2))
        rlbl(ax, 11, 1.7, "Bubble")
        rlbl(ax, 14, 2.5, "Stalk")
        rlbl(ax, 0.5, -2, "RB run")
        rlbl(ax, -7.5, 2.5, "Block")
        rlbl(ax, -14, 2.5, "Block")

    # NINJA
    fig, ax = new_fig("RPO vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "reads\nRPO key")
    nt(ax, -12, Y_CB + 1, "reads\nRPO key")
    nt(ax, 5, Y_S + 1, "triggers if\n#2 screens")
    nt(ax, -5, Y_S + 1, "run fit")
    nt(ax, 7.5, Y_APEX + 1, "KEY: don't\nbite on run")
    nt(ax, 1.5, Y_LB + 1, "run fit\nread mesh")
    nt(ax, -1.5, Y_LB + 1, "run fit")
    bn(ax, "B must not over-commit to run. FS drives on bubble if thrown.")
    R["rpo_ninja"] = save_fig(fig, "rpo_ninja")

    # OREGON
    fig, ax = new_fig("RPO vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 3), (11, 2))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 0.5, 10.8, "post/MOF")
    nt(ax, 12, 1.2, "B: on\nbubble")
    nt(ax, -8, 2.2, "D: man\n#2B")
    bn(ax, "Man accounts for every receiver. B has #2F bubble.")
    R["rpo_oregon"] = save_fig(fig, "rpo_oregon")

    # OKLAHOMA
    fig, ax = new_fig("RPO vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 3), (11, 2))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 0.5, 10.8, "D: post/MOF")
    nt(ax, 12, 1.2, "B: on\nbubble")
    nt(ax, -8, 2.2, "FS: man\n#2B")
    bn(ax, "D (field) to post. Man accounts for bubble receiver.")
    R["rpo_oklahoma"] = save_fig(fig, "rpo_oklahoma")

    # ZEUS
    fig, ax = new_fig("RPO vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (11, 2), (-7, 3))
    nt(ax, 12, Y_CB + 1, "man #1")
    nt(ax, -12, Y_CB + 1, "man #1")
    nt(ax, 12, 1.2, "FS: on\nbubble")
    nt(ax, -8, 2.2, "D: man\n#2B")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Man covers all receivers. Pressure forces quick decision.")
    R["rpo_zeus"] = save_fig(fig, "rpo_zeus")

    # VIKING
    fig, ax = new_fig("RPO vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid 1/3")
    nt(ax, -7, 4.7, "D: run fit\ndiscipline")
    nt(ax, 7.5, Y_APEX + 1, "B: KEY\ndon't vacate")
    nt(ax, 1.5, Y_LB + 1, "run fit\nread mesh")
    nt(ax, -1.5, Y_LB + 1, "run fit")
    bn(ax, "If B vacates flat on run fake, bubble is open behind him.")
    R["rpo_viking"] = save_fig(fig, "rpo_viking")

    return R


def gen_double_post():
    """Double Post vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 7), (8, 10)])
        draw_route_path(ax, [(7, 0), (7, 6), (3, 9)])
        draw_route_path(ax, [(-13, 0), (-13, 8)])
        draw_route_path(ax, [(-7, 0), (-7, 5)])
        rlbl(ax, 8.5, 10.5, "#1 Post")
        rlbl(ax, 3.5, 9.5, "#2 Post")
        rlbl(ax, -14, 8, "Clear")
        rlbl(ax, -7.5, 5.5, "Sit")

    # NINJA
    fig, ax = new_fig("Double Post vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "matches\n#1 post")
    nt(ax, -12, Y_CB + 1, "clamps\n#1 clear")
    nt(ax, 5, Y_S + 1, "takes\n#2 post")
    nt(ax, -5, Y_S + 1, "controls\n#2 sit")
    nt(ax, 7.5, Y_APEX + 1, "carries\n#2 in")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Two-high splits both posts. Best coverage vs double post.")
    R["dp_ninja"] = save_fig(fig, "dp_ninja")

    # OREGON
    fig, ax = new_fig("Double Post vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 5), (3, 8))
    nt(ax, 12, Y_CB + 1, "man #1\npost")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "post/MOF\n2 posts at him")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 4, 7.2, "B: man\n#2F post")
    bn(ax, "Single post vs 2 post routes = losing matchup. Need pressure.")
    R["dp_oregon"] = save_fig(fig, "dp_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Double Post vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 5), (3, 8))
    nt(ax, 12, Y_CB + 1, "man #1\npost")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "D: post/MOF\n2 posts at him")
    nt(ax, -8, 4.2, "FS: man\n#2B sit")
    nt(ax, 4, 7.2, "B: man\n#2F post")
    bn(ax, "D (field) to post. Two posts converge on him \u2014 tough matchup.")
    R["dp_oklahoma"] = save_fig(fig, "dp_oklahoma")

    # ZEUS
    fig, ax = new_fig("Double Post vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (3, 8), (-7, 5))
    nt(ax, 12, Y_CB + 1, "man #1\npost")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 4, 7.2, "FS: man\n#2F post")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    bn(ax, "Extremely vulnerable. Both posts free if defenders trail.")
    R["dp_zeus"] = save_fig(fig, "dp_zeus")

    # VIKING
    fig, ax = new_fig("Double Post vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3\nleans on post")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "FS: splits\nboth posts")
    nt(ax, -7, 4.7, "D: seam\ncurl-flat")
    nt(ax, 7.5, Y_APEX + 1, "B: wall\n#2 post")
    nt(ax, 1.5, Y_LB + 1, "hook")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "FS splits both posts in mid 1/3. Better than single-high.")
    R["dp_viking"] = save_fig(fig, "dp_viking")

    return R


def gen_stick():
    """Stick (Quick Game) vs all coverages."""
    R = {}

    def routes(ax):
        draw_route_path(ax, [(13, 0), (13, 5), (12, 5)])
        draw_route_path(ax, [(7, 0), (10, 2)])
        draw_route_path(ax, [(0, -3), (0, 7)])
        draw_route_path(ax, [(-13, 0), (-13, 8)])
        draw_route_path(ax, [(-7, 0), (-7, 5)])
        rlbl(ax, 12.5, 5.7, "Stick")
        rlbl(ax, 11, 2.5, "Flat")
        rlbl(ax, 1, 7, "RB vert")
        rlbl(ax, -14, 8, "Clear")
        rlbl(ax, -7.5, 5.5, "Sit")

    # NINJA
    fig, ax = new_fig("Stick vs NINJA (Cover 7)")
    concept_base(ax); routes(ax); std_def(ax)
    nt(ax, 12, Y_CB + 1, "matches\n#1 stick")
    nt(ax, -12, Y_CB + 1, "clamps\n#1 clear")
    nt(ax, 5, Y_S + 1, "alerts to\nRB vert")
    nt(ax, -5, Y_S + 1, "controls\n#2 sit")
    nt(ax, 7.5, Y_APEX + 1, "reads\n#2 flat")
    nt(ax, 1.5, Y_LB + 1, "sits on\nstick")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Zone drops cover all 3 levels. B reads flat, M sits on stick.")
    R["stick_ninja"] = save_fig(fig, "stick_ninja")

    # OREGON
    fig, ax = new_fig("Stick vs OREGON (Cover 1, Post = FS)")
    concept_base(ax); routes(ax); std_def(ax)
    or_drops(ax, (-7, 5), (10, 2.5))
    nt(ax, 12, Y_CB + 1, "man #1\nstick")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "post/MOF\nhelps vert")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 11, 1.7, "B: man\n#2F flat")
    nt(ax, 1.5, Y_LB + 1, "RB vert\nfunnel")
    bn(ax, "Man on each level. Post safety reads QB and helps on vert push.")
    R["stick_oregon"] = save_fig(fig, "stick_oregon")

    # OKLAHOMA
    fig, ax = new_fig("Stick vs OKLAHOMA (Cover 1, Post = D)")
    concept_base(ax); routes(ax); ok_def(ax)
    ok_drops(ax, (-7, 5), (10, 2.5))
    nt(ax, 12, Y_CB + 1, "man #1\nstick")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 0.5, 10.8, "D: post/MOF\nhelps vert")
    nt(ax, -8, 4.2, "FS: man\n#2B sit")
    nt(ax, 11, 1.7, "B: man\n#2F flat")
    bn(ax, "D (field) to post. Man covers each level cleanly.")
    R["stick_oklahoma"] = save_fig(fig, "stick_oklahoma")

    # ZEUS
    fig, ax = new_fig("Stick vs ZEUS (Cover 0)")
    concept_base(ax); routes(ax); std_def(ax)
    z_drops(ax, (10, 2.5), (-7, 5))
    nt(ax, 12, Y_CB + 1, "man #1\nstick")
    nt(ax, -12, Y_CB + 1, "man #1\nclear")
    nt(ax, 11, 1.7, "FS: man\n#2F flat")
    nt(ax, -8, 4.2, "D: man\n#2B sit")
    nt(ax, 0.5, Y_LB - 0.5, "M: rush")
    nt(ax, 7.5, Y_APEX + 1, "B: RB\nfunnel")
    bn(ax, "Man on each. Need pressure before QB reads 3-level prog.")
    R["stick_zeus"] = save_fig(fig, "stick_zeus")

    # VIKING
    fig, ax = new_fig("Stick vs VIKING (Cover 3)")
    concept_base(ax); routes(ax); std_def(ax)
    vik_drops(ax)
    nt(ax, 11.5, 11.3, "deep 1/3")
    nt(ax, -11.5, 11.3, "deep 1/3")
    nt(ax, 0, 11.3, "deep mid\ncaps vert")
    nt(ax, -7, 4.7, "D: seam\ncurl-flat")
    nt(ax, 7.5, Y_APEX + 1, "B: flat or\nstick?")
    nt(ax, 1.5, Y_LB + 1, "sits on\nstick")
    nt(ax, -1.5, Y_LB + 1, "hook")
    bn(ax, "Similar to curl/flat conflict. B must choose flat or stick.")
    R["stick_viking"] = save_fig(fig, "stick_viking")

    return R


# =============================================================================
# GENERATE
# =============================================================================

def generate_diagrams():
    paths = {}
    paths["pre_snap"] = save_fig(draw_pre_snap_shell(), "pre_snap")
    paths["ninja_2x2"] = save_fig(draw_ninja_post_snap(), "ninja_2x2_ps")
    paths["ninja_3x1"] = save_fig(draw_ninja_3x1_post_snap(), "ninja_3x1_ps")
    paths["oregon"] = save_fig(draw_oregon_post_snap(), "oregon_ps")
    paths["oklahoma"] = save_fig(draw_oklahoma_post_snap(), "oklahoma_ps")
    paths["ohio"] = save_fig(draw_ohio_exception(), "ohio_ps")
    paths["zeus"] = save_fig(draw_zeus_post_snap(), "zeus_ps")
    paths["zorro"] = save_fig(draw_zorro_exception(), "zorro_ps")
    paths["zunnel"] = save_fig(draw_zunnel_post_snap(), "zunnel_ps")
    paths["viking"] = save_fig(draw_viking_post_snap(), "viking_ps")

    # Concept x Coverage diagrams
    for gen_fn in [gen_4verts, gen_mesh, gen_smash, gen_flood,
                   gen_levels, gen_curl_flat, gen_slant_flat,
                   gen_rpo, gen_double_post, gen_stick]:
        paths.update(gen_fn())

    return paths


# =============================================================================
# BUILD DOC
# =============================================================================

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def make_header_row(table, headers):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, NAVY_HEX)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True


def build_doc(paths):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.font.color.rgb = DARK

    def h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs: r.font.color.rgb = NAVY

    def h2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs: r.font.color.rgb = NAVY

    def h3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs: r.font.color.rgb = NAVY

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)

    def img(key, width=5.5):
        if key in paths:
            doc.add_picture(paths[key], width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === TITLE ===
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("PRE-SNAP DISGUISE &\nCOVERAGE RULES")
    r.bold = True; r.font.size = Pt(24); r.font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = y.add_run("2026 Season")
    r.font.size = Pt(18); r.font.color.rgb = NAVY
    doc.add_page_break()

    # === PRINCIPLE ===
    h1("PRE-SNAP DISGUISE PRINCIPLE")
    body("Every coverage we run \u2014 NINJA, Oregon, Oklahoma, Zeus, Zunnel, Zill, Zike, Viking \u2014 "
         "looks the same before the snap. The offense sees the same two-high shell every time. "
         "DBs execute their coverage responsibilities post-snap.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("The offense should never know what coverage we are in until after the ball is snapped.")
    r.bold = True
    r.font.size = Pt(13)

    h2("Exceptions (2)")
    bullet("B moves from apex to ~8 yards deep in the middle of the formation to become the post/MOF safety. His movement tips the coverage.", bold_prefix="OHIO: ")
    bullet("B moves from apex to ~8 yards deep in the middle of the formation to key the RB. His alignment change tips the coverage.", bold_prefix="ZORRO: ")

    doc.add_page_break()

    # === PRE-SNAP SHELL ===
    h1("UNIVERSAL PRE-SNAP ALIGNMENT")
    body("This is the look the offense sees on every snap (except Ohio and Zorro).")

    img("pre_snap")

    h2("Alignment Rules")
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(table, ["Position", "Pre-Snap Alignment", "Key"])
    align_data = [
        ("FC / BC", "Inside leverage, 6 yards on #1", "Do not tip coverage with depth or leverage changes"),
        ("FS", "10-12 yds deep, inside foot on #2 (~1 yd inside #2)", "If no #2: hash-to-middle, eyes on #3/QB"),
        ("D", "10-12 yds deep, inside foot on #2 (~1 yd inside #2)", "If no #2: hash-to-middle, eyes on #3/QB"),
        ("B", "Apex to field, split between #2 and box", "Same spot every coverage (except Ohio/Zorro)"),
        ("M / W", "Standard LB depth over assigned gap", "Run-fit alignment does not change by coverage"),
    ]
    for i, (pos, align, key) in enumerate(align_data):
        table.rows[i + 1].cells[0].text = pos
        table.rows[i + 1].cells[1].text = align
        table.rows[i + 1].cells[2].text = key

    doc.add_paragraph()
    bullet("Compressed/stack splits: safeties tighten 1-2 steps, ready for BANJO. This is formation-driven, not coverage-driven \u2014 same adjustment in every coverage.")
    bullet("Motion: NINJA re-checks (POACH/MOD/CLAMP). Man coverages use BUMP-BUMP. Either way, the pre-snap shell resets the same.")

    doc.add_page_break()

    # === COVERAGE BY COVERAGE ===
    h1("COVERAGE RULES \u2014 POST-SNAP DROPS")
    body("Below is every coverage with its post-snap responsibilities. "
         "Pre-snap alignment is identical for all (except Ohio and Zorro). "
         "Solid circles = pre-snap position. Dashed circles = post-snap drop.")

    doc.add_page_break()

    # -- NINJA 2x2 --
    h2("NINJA vs 2x2 \u2014 MOD (Field) / CLAMP (Boundary)")
    body("No rotation needed. The pre-snap look IS the coverage alignment. DBs execute pattern-match rules at the snap.")
    img("ninja_2x2")

    h3("Coverage Rules")
    bullet("Man-match #1. Stay on top, deny explosives. #1 shallow = pass off if rules allow, look for next threat.", bold_prefix="FC: ")
    bullet("Clamp technique on #1. Physical, deny release. #1 vertical = stay on top.", bold_prefix="BC: ")
    bullet("MOD rules: top-down on #2. #2 vertical = match and stay on top. #2 out/flat = drive with control, don't open seams.", bold_prefix="FS: ")
    bullet("CLAMP rules: control #2 and help corner. #2 vertical = match. #2 out fast = drive with leverage.", bold_prefix="D: ")
    bullet("Apex rules to #2/#3 threats. Eliminate quick RPO access. Carry seams long enough for safety help.", bold_prefix="B: ")
    bullet("Hook. Eyes to QB, relate to threats.", bold_prefix="M / W: ")
    doc.add_page_break()

    # -- NINJA 3x1 --
    h2("NINJA vs 3x1 \u2014 POACH")
    body("D is the poach safety when trips go to the field. FS is the poach safety when trips go to the boundary. "
         "Poach safety slides post-snap toward the trips side.")
    img("ninja_3x1")

    h3("Coverage Rules")
    bullet("#3 vertical: poach player takes it (no free seam).")
    bullet("#3 shallow/under: poach player communicates and overlaps crossers.")
    bullet("Backside stays sound \u2014 no free go balls or glance RPO.")
    bullet("Poach player (D in diagram) slides post-snap. Pre-snap he is still at 10-12 over #2 boundary.", bold_prefix="Key: ")
    doc.add_page_break()

    # -- OREGON --
    h2("OREGON (Cover 1 \u2014 Post = FS)")
    body("Pre-snap: same two-high. Post-snap: FS rotates to MOF as the post/fixer. D and B drop to man assignments.")
    img("oregon")

    h3("Coverage Rules")
    bullet("Man on #1 to your side.", bold_prefix="FC / BC: ")
    bullet("Post-snap: rotate to MOF. Fixer \u2014 make the corner right.", bold_prefix="FS: ")
    bullet("Post-snap: drop to man on #2 boundary, or #3 away if no #2 boundary.", bold_prefix="D: ")
    bullet("Post-snap: drop to man on #2 field, or #3 away if no #2 field.", bold_prefix="B: ")
    bullet("RB funnel. RB to your side = take. RB away = rush. TAMPA: away LB = RAT.", bold_prefix="M / W: ")
    doc.add_page_break()

    # -- OKLAHOMA --
    h2("OKLAHOMA (Cover 1 \u2014 Post = D)")
    body("Pre-snap: FS and D switch sides \u2014 D aligns to the field, FS aligns to the boundary. "
         "Post-snap: D rotates from field to MOF. FS drops from boundary to man #2 boundary.")
    img("oklahoma")

    h3("Coverage Rules")
    bullet("Man on #1 to your side.", bold_prefix="FC / BC: ")
    bullet("Post-snap: rotate to MOF. Fixer.", bold_prefix="D: ")
    bullet("Aligns to boundary pre-snap. Post-snap: man on #2 boundary, or #3 away.", bold_prefix="FS: ")
    bullet("Post-snap: drop to man on #2 field, or #3 away if no #2 field.", bold_prefix="B: ")
    bullet("RB funnel.", bold_prefix="M / W: ")
    doc.add_page_break()

    # -- OHIO (EXCEPTION) --
    h2("OHIO (Cover 1 \u2014 Post = B) \u2014 EXCEPTION")
    body("B aligns at ~8 yards deep in the middle of the formation to become the MOF post/fixer. "
         "This changes the pre-snap picture \u2014 the offense may read B's alignment out of the apex. "
         "Both safeties drop to man on #2.")
    img("ohio")

    h3("Coverage Rules")
    bullet("Man on #1 to your side.", bold_prefix="FC / BC: ")
    bullet("Post-snap: rotate from apex to MOF. Fixer \u2014 make the corner right.", bold_prefix="B: ")
    bullet("Post-snap: drop to man on #2 field, or #3 away if no #2 field.", bold_prefix="FS: ")
    bullet("Post-snap: drop to man on #2 boundary, or #3 away if no #2 boundary.", bold_prefix="D: ")
    bullet("RB funnel.", bold_prefix="M / W: ")
    body("Note: Because B vacates the apex, this is one of two coverages that changes the pre-snap look.")
    doc.add_page_break()

    # -- ZEUS --
    h2("ZEUS (Cover 0 \u2014 Delayed Pressure)")
    body("Pre-snap: same two-high. Post-snap: both safeties drop to man. No deep help. "
         "Run-first read; on pass, DL and M rush, B/W funnel RB.")
    img("zeus")

    h3("Coverage Rules")
    bullet("Man on #1 to your side.", bold_prefix="FC / BC: ")
    bullet("Post-snap: drop to man on #2 field; if no #2 field, #3 away.", bold_prefix="FS: ")
    bullet("Post-snap: drop to man on #2 boundary; if no #2 boundary, #3 away.", bold_prefix="D: ")
    bullet("RB funnel. RB to your side = take (man). RB away = rush. RB middle = W takes.", bold_prefix="B / W: ")
    bullet("Primary add-on rusher on pass read.", bold_prefix="M: ")
    bullet("Cage/contain \u2014 do not run past QB depth. Force step-up.", bold_prefix="A / E: ")
    bullet("Vertical push / collapse pocket.", bold_prefix="T / N: ")
    doc.add_page_break()

    # -- ZORRO (EXCEPTION) --
    h2("ZORRO (Cover 0) \u2014 EXCEPTION")
    body("B has the RB, so B aligns at ~8 yards deep in the middle of the formation instead of at his normal apex. "
         "This changes the pre-snap picture. M/W have no pass responsibility unless tagged.")
    img("zorro")

    h3("Coverage Rules")
    bullet("Man on #1 to your side.", bold_prefix="FC / BC: ")
    bullet("Post-snap: drop to man on #2 field; if no #2, #3 away.", bold_prefix="FS: ")
    bullet("Post-snap: drop to man on #2 boundary; if no #2, #3 away.", bold_prefix="D: ")
    bullet("Has RB (man). Aligns at ~8 yds in the middle of the formation to key RB.", bold_prefix="B: ")
    bullet("No pass responsibility unless tagged. TAMPA = become droppers. SPY = spy QB.", bold_prefix="M / W: ")
    body("Note: Because B aligns at ~8 yds in the middle instead of the apex, this is one of two coverages that changes the pre-snap look.")
    doc.add_page_break()

    # -- ZUNNEL / ZILL / ZIKE --
    h2("ZUNNEL / ZILL / ZIKE (Cover 0)")
    body("Pre-snap: same two-high. Post-snap: both safeties drop to man. RB handled per call.")
    img("zunnel")

    h3("Coverage Rules (Shared)")
    bullet("Man on #1.", bold_prefix="FC / BC: ")
    bullet("Man on #2 field (or #3 away).", bold_prefix="FS: ")
    bullet("Man on #2 boundary (or #3 away).", bold_prefix="D: ")
    bullet("Normal apex pre-snap. Post-snap: per call assignment.", bold_prefix="B: ")

    h3("RB Responsibility by Call")
    ztable = doc.add_table(rows=4, cols=2)
    ztable.style = "Light Grid Accent 1"
    ztable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ztable, ["Call", "RB Player"])
    for i, (c, p) in enumerate([("Zunnel", "M/W funnel (to you = take, away = rush)"),
                                 ("Zill", "W has RB (man)"), ("Zike", "M has RB (man)")]):
        ztable.rows[i + 1].cells[0].text = c
        ztable.rows[i + 1].cells[1].text = p
    doc.add_page_break()

    # -- VIKING --
    h2("VIKING (Cover 3 \u2014 Spot Drop)")
    body("Pre-snap: same two-high. Post-snap: corners bail to deep thirds, FS rotates to deep middle third. "
         "D drops to seam-curl-flat. B plays curl/flat. M/W play hook.")
    img("viking")

    h3("Coverage Rules")
    bullet("Post-snap: bail to deep 1/3 (field).", bold_prefix="FC: ")
    bullet("Post-snap: bail to deep 1/3 (boundary).", bold_prefix="BC: ")
    bullet("Post-snap: rotate to deep middle 1/3.", bold_prefix="FS: ")
    bullet("Post-snap: drop to seam-curl-flat (boundary).", bold_prefix="D: ")
    bullet("Post-snap: curl/flat (field).", bold_prefix="B: ")
    bullet("Hook. Eyes to QB, relate to #3.", bold_prefix="M: ")
    bullet("Hook.", bold_prefix="W: ")

    h3("Tags")
    bullet("Match #2 vertical rules.", bold_prefix="VIKING SEAM: ")
    bullet("Distribution adjustment vs 3x1.", bold_prefix="VIKING PUSH: ")
    bullet("Carry crosser to depth before passing.", bold_prefix="VIKING CROSS: ")
    bullet("Trigger/replace rules vs now/bubble.", bold_prefix="VIKING SCREEN: ")
    doc.add_page_break()

    # === SUMMARY TABLE ===
    h1("COVERAGE SUMMARY \u2014 PRE-SNAP vs POST-SNAP")

    summ = doc.add_table(rows=11, cols=5)
    summ.style = "Light Grid Accent 1"
    summ.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(summ, ["Coverage", "Pre-Snap", "FS Post-Snap", "D Post-Snap", "B Post-Snap"])
    sdata = [
        ("NINJA 2x2", "Two-high", "MOD: top-down #2", "CLAMP: control #2", "Apex: #2/#3 threats"),
        ("NINJA 3x1\n(trips field)", "Two-high", "Hold field half", "POACH toward trips", "Apex: trips threats"),
        ("NINJA 3x1\n(trips bnd)", "Two-high", "POACH toward trips", "Hold boundary half", "Apex: trips threats"),
        ("OREGON", "Two-high", "Rotate to MOF (post)", "Man #2 bnd", "Man #2 field"),
        ("OKLAHOMA", "Two-high", "Man #2 bnd", "Rotate to MOF (post)", "Man #2 field"),
        ("OHIO *", "Changed", "Man #2 field", "Man #2 bnd", "Rotate to MOF (post)"),
        ("ZEUS", "Two-high", "Man #2 field", "Man #2 bnd", "RB funnel"),
        ("ZORRO *", "Changed", "Man #2 field", "Man #2 bnd", "Has RB (~8 yds middle)"),
        ("ZUNNEL", "Two-high", "Man #2 field", "Man #2 bnd", "Per call"),
        ("VIKING", "Two-high", "Rotate to deep mid 1/3", "Seam-curl-flat", "Curl/flat"),
    ]
    for i, row_data in enumerate(sdata):
        for j, val in enumerate(row_data):
            summ.rows[i + 1].cells[j].text = val

    doc.add_paragraph()
    body("* = Exception. Pre-snap look changes from the standard two-high shell.")
    doc.add_page_break()

    # === PASS CONCEPTS SECTION (DIAGRAM-BASED) ===
    h1("PASS CONCEPT BREAKDOWN BY COVERAGE")
    body("How each coverage defends common pass concepts. Each concept is shown "
         "with a diagram for every coverage: NINJA, Oregon, Oklahoma, Zeus, and Viking. "
         "Red dashed lines = offensive routes. Navy circles = defensive players and drops.")

    # --- FOUR VERTICALS ---
    doc.add_page_break()
    h2("1. FOUR VERTICALS (4 Verts)")
    body("All four eligible receivers push vertical. Stretches the defense deep.")
    img("4v_ninja", 5.0); img("4v_oregon", 5.0)
    doc.add_page_break()
    img("4v_oklahoma", 5.0); img("4v_zeus", 5.0)
    doc.add_page_break()
    img("4v_viking", 5.0)
    doc.add_page_break()

    # --- MESH / CROSSERS ---
    h2("2. MESH / CROSSERS")
    body("Two receivers cross underneath. Creates pick/rub action.")
    img("mesh_ninja", 5.0); img("mesh_oregon", 5.0)
    doc.add_page_break()
    img("mesh_oklahoma", 5.0); img("mesh_zeus", 5.0)
    doc.add_page_break()
    img("mesh_viking", 5.0)
    doc.add_page_break()

    # --- SMASH ---
    h2("3. SMASH (Corner / Hitch)")
    body("#1 hitches at 5-6 yds. #2 runs corner behind it. Attacks the flat defender.")
    img("smash_ninja", 5.0); img("smash_oregon", 5.0)
    doc.add_page_break()
    img("smash_oklahoma", 5.0); img("smash_zeus", 5.0)
    doc.add_page_break()
    img("smash_viking", 5.0)
    doc.add_page_break()

    # --- FLOOD / SAIL ---
    h2("4. FLOOD / SAIL (3-Level Stretch)")
    body("Three receivers to one side at three levels: deep, intermediate, short.")
    img("flood_ninja", 5.0); img("flood_oregon", 5.0)
    doc.add_page_break()
    img("flood_oklahoma", 5.0); img("flood_zeus", 5.0)
    doc.add_page_break()
    img("flood_viking", 5.0)
    doc.add_page_break()

    # --- LEVELS ---
    h2("5. LEVELS (Hi-Lo Crossers)")
    body("Two crossers at different depths. Reads the hook/curl defender.")
    img("levels_ninja", 5.0); img("levels_oregon", 5.0)
    doc.add_page_break()
    img("levels_oklahoma", 5.0); img("levels_zeus", 5.0)
    doc.add_page_break()
    img("levels_viking", 5.0)
    doc.add_page_break()

    # --- CURL/FLAT ---
    h2("6. CURL / FLAT")
    body("#1 curls at 12 yds. #2 runs flat. Classic 2-man read of the flat defender.")
    img("cf_ninja", 5.0); img("cf_oregon", 5.0)
    doc.add_page_break()
    img("cf_oklahoma", 5.0); img("cf_zeus", 5.0)
    doc.add_page_break()
    img("cf_viking", 5.0)
    doc.add_page_break()

    # --- SLANTS ---
    h2("7. SLANT / FLAT (Quick Game)")
    body("#1 slants inside. #2 runs flat. Fast-developing quick game concept.")
    img("sf_ninja", 5.0); img("sf_oregon", 5.0)
    doc.add_page_break()
    img("sf_oklahoma", 5.0); img("sf_zeus", 5.0)
    doc.add_page_break()
    img("sf_viking", 5.0)
    doc.add_page_break()

    # --- RPO CONCEPTS ---
    h2("8. RPO CONCEPTS (Run-Pass Options)")
    body("QB reads a defender. If he bites on the run, QB throws the quick pass behind him.")
    img("rpo_ninja", 5.0); img("rpo_oregon", 5.0)
    doc.add_page_break()
    img("rpo_oklahoma", 5.0); img("rpo_zeus", 5.0)
    doc.add_page_break()
    img("rpo_viking", 5.0)
    doc.add_page_break()

    # --- DOUBLE POST ---
    h2("9. DOUBLE POST / POST-DIG")
    body("Two receivers run post routes. Attacks MOF, stresses single-high safeties.")
    img("dp_ninja", 5.0); img("dp_oregon", 5.0)
    doc.add_page_break()
    img("dp_oklahoma", 5.0); img("dp_zeus", 5.0)
    doc.add_page_break()
    img("dp_viking", 5.0)
    doc.add_page_break()

    # --- STICK ---
    h2("10. STICK (Quick Game)")
    body("#1 sticks at 6 yds. #2 flat. RB vertical push. A 3-level quick concept.")
    img("stick_ninja", 5.0); img("stick_oregon", 5.0)
    doc.add_page_break()
    img("stick_oklahoma", 5.0); img("stick_zeus", 5.0)
    doc.add_page_break()
    img("stick_viking", 5.0)
    doc.add_page_break()

    # === CONCEPT COVERAGE MATRIX ===
    h1("CONCEPT vs COVERAGE QUICK REFERENCE")
    body("Quick reference: which coverages are strongest and weakest against each common concept. "
         "\u2705 = Strong matchup. \u26A0\uFE0F = Caution/conflict. \u274C = Vulnerable.")

    matrix = doc.add_table(rows=11, cols=5)
    matrix.style = "Light Grid Accent 1"
    matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(matrix, ["Concept", "NINJA (Cov 7)", "Cover 1 (OR/OK)", "Cover 0 (Zeus)", "VIKING (Cov 3)"])
    mdata = [
        ("Four Verts",        "Strong \u2014 two-high caps seams",     "Caution \u2014 single post vs 4 seams",   "Vulnerable \u2014 no help deep",          "Strong \u2014 3 deep zones cap it"),
        ("Mesh / Crossers",   "Strong \u2014 zone handles picks",     "Caution \u2014 picks trouble man",        "Vulnerable \u2014 picks + no help",       "Strong \u2014 zone handles picks"),
        ("Smash",             "Strong \u2014 FS on corner, CB on hitch","Strong \u2014 man matches both",         "Caution \u2014 no help on corner",        "Strong \u2014 deep 1/3 caps corner"),
        ("Flood / Sail",      "Caution \u2014 3rd level stress",      "Caution \u2014 post must commit",         "Strong \u2014 man runs with each",        "Caution \u2014 flat vulnerable"),
        ("Levels",            "Strong \u2014 zone reads hi-lo",       "Caution \u2014 long chases across field", "Vulnerable \u2014 trail + no help",       "Strong \u2014 zone layers"),
        ("Curl / Flat",       "Strong \u2014 two-high covers both",   "Strong \u2014 man on each",               "Strong \u2014 man on each",               "Caution \u2014 flat conflict"),
        ("Slant / Flat",      "Strong \u2014 inside leverage + FS",   "Strong \u2014 CB physical on slant",      "Strong \u2014 pressure + man",            "Caution \u2014 corners bail"),
        ("RPO",               "Caution \u2014 B must read, not bite", "Strong \u2014 man accounts all",          "Strong \u2014 man accounts all",          "Caution \u2014 flat discipline"),
        ("Double Post",       "Strong \u2014 two-high splits posts",  "Vulnerable \u2014 single post vs two",    "Vulnerable \u2014 no help deep",          "Strong \u2014 FS splits posts"),
        ("Stick",             "Strong \u2014 zone covers all levels", "Strong \u2014 man on each level",         "Strong \u2014 man on each",               "Caution \u2014 flat conflict"),
    ]
    for i, row_data in enumerate(mdata):
        for j, val in enumerate(row_data):
            matrix.rows[i + 1].cells[j].text = val

    doc.add_paragraph()
    body("Use this matrix to game-plan coverage calls based on opponent tendencies. "
         "If they love four verts and double posts, lean NINJA and VIKING. "
         "If they run heavy mesh/crossers, NINJA and VIKING handle picks best. "
         "If they live on quick game (slant/flat, curl/flat, stick), any coverage works but pressure helps most.")

    # Footer
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run("\u2014 END OF DOCUMENT \u2014")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    doc.add_paragraph()
    g = doc.add_paragraph()
    g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = g.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    r.font.size = Pt(9); r.font.color.rgb = GRAY_RGB

    out = "/home/ksc4130/src/defensive_playbook/River_Valley_Vikings_PreSnap_Disguise_Coverage_Rules.docx"
    doc.save(out)
    print(f"Saved to: {out}")


if __name__ == "__main__":
    print("Generating diagrams...")
    paths = generate_diagrams()
    print(f"Generated {len(paths)} diagrams.")
    print("Building document...")
    build_doc(paths)
