#!/usr/bin/env python3
"""
V2 Offseason Meal Log FORM (tracking pages only — no educational content).
Optimized for 6 AM lift → school schedule.
Print multiples of this. Pair with the separate guide document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- Colors --
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_BLUE_HEX = "D6EAF8"
LIGHT_GOLD_HEX = "FFF3CC"
LIGHT_GRAY_HEX = "F2F2F2"
LIGHT_GREEN_HEX = "D5F5E3"
NAVY_HEX = "002D62"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_row_height(row, inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_header_row(table, headers, font_size=10):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        set_cell_shading(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = WHITE
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.name = "Calibri"


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # -- Helpers --
    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def small_note(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True
        return p

    # ================================================================
    # COVER / PLAYER INFO (page 1)
    # ================================================================
    for _ in range(2):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("OFFSEASON MEAL LOG")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = GOLD

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("4-Week Recovery Tracking Form")
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY

    body("")

    # Player info
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    info_labels = [
        ("Name:", "", "Position:", ""),
        ("Grade:", "", "Height:", ""),
        ("Current Weight:", "", "Start Date:", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_labels):
        row = info_table.rows[i]
        set_row_height(row, 0.3)
        c0, c1, c2, c3 = row.cells
        set_cell_shading(c0, LIGHT_BLUE_HEX)
        set_cell_shading(c2, LIGHT_BLUE_HEX)
        c0.text = l1
        c1.text = v1
        c2.text = l2
        c3.text = v2
        for cell in [c0, c2]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)

    body("")

    # Starting lifts
    lift_table = doc.add_table(rows=2, cols=5)
    lift_table.style = "Table Grid"
    make_header_row(lift_table, ["Squat", "Bench", "Deadlift", "Power Clean", "Bodyweight"], font_size=9)
    set_row_height(lift_table.rows[1], 0.35)

    body("")

    # Quick reference reminder (compact)
    ref_table = doc.add_table(rows=2, cols=4)
    ref_table.style = "Table Grid"
    make_header_row(ref_table, [
        "Post-Lift (7:00 AM)",
        "Protein Target",
        "Water Target",
        "Bedtime",
    ], font_size=8)
    row = ref_table.rows[1]
    row.cells[0].text = "Protein + carbs\nwithin 30 min"
    row.cells[1].text = "1g per lb BW\n(every meal)"
    row.cells[2].text = "Half BW in oz\n(start at wake-up)"
    row.cells[3].text = "Bed by 9:00 PM\nLights out 9:15"
    set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)
    set_cell_shading(row.cells[3], LIGHT_GOLD_HEX)
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.bold = True

    doc.add_page_break()

    # ================================================================
    # WEEKLY PROGRESS TRACKER
    # ================================================================
    heading2("12-Week Progress Tracker")
    small_note("Weigh in at the same time each week (morning, before eating). Record lift numbers when tested.")

    prog_table = doc.add_table(rows=13, cols=7)
    prog_table.style = "Table Grid"
    make_header_row(
        prog_table,
        ["Wk", "Date", "Weight", "Squat", "Bench", "Clean", "Notes"],
        font_size=9,
    )

    for i in range(1, 13):
        row = prog_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_row_height(row, 0.3)
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY_HEX)

    body("")

    # Coach check-in slots (compact — 4 weeks on one page)
    heading2("Weekly Coach Check-In")
    for week in range(1, 5):
        p = doc.add_paragraph()
        r = p.add_run(f"Week {week}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY

        notes_table = doc.add_table(rows=3, cols=2)
        notes_table.style = "Table Grid"
        labels = [
            "Meals/day (avg):",
            "Post-lift meals logged:",
            "Adjust for next week:",
        ]
        for i, label in enumerate(labels):
            row = notes_table.rows[i]
            set_row_height(row, 0.3)
            row.cells[0].text = label
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
            set_cell_width(row.cells[0], 1.8)
            set_cell_width(row.cells[1], 4.7)
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

    doc.add_page_break()

    # ================================================================
    # DAILY MEAL LOG PAGES (4 weeks × 7 days)
    # ================================================================

    # Meal labels for the 6 AM schedule
    LIFT_DAY_MEALS = [
        ("Wake-Up Bite", LIGHT_GOLD_HEX),
        ("Post-Lift Fuel", LIGHT_GREEN_HEX),
        ("School Breakfast", LIGHT_BLUE_HEX),
        ("Mid-Morning", LIGHT_GRAY_HEX),
        ("Lunch", LIGHT_BLUE_HEX),
        ("After School", LIGHT_GRAY_HEX),
        ("Dinner", LIGHT_BLUE_HEX),
        ("Bedtime Fuel", LIGHT_GOLD_HEX),
    ]

    REST_DAY_MEALS = [
        ("Breakfast", LIGHT_BLUE_HEX),
        ("Mid-Morning", LIGHT_GRAY_HEX),
        ("Lunch", LIGHT_BLUE_HEX),
        ("Afternoon", LIGHT_GRAY_HEX),
        ("Dinner", LIGHT_BLUE_HEX),
        ("Bedtime Fuel", LIGHT_GOLD_HEX),
    ]

    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

    for week in range(1, 5):
        heading2(f"Week {week}")

        for day_idx in range(7):
            day_name = day_names[day_idx]
            day_num = (week - 1) * 7 + day_idx + 1
            is_weekend = day_name in ["Saturday", "Sunday"]

            # Day header
            day_head = doc.add_paragraph()
            r = day_head.add_run(f"{day_name}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            day_head.add_run("  ")
            r2 = day_head.add_run("Date: ____________")
            r2.font.size = Pt(9)
            r2.font.color.rgb = GRAY

            if not is_weekend:
                day_head.add_run("   ")
                r3 = day_head.add_run("Lift Day?  \u2610 Yes  \u2610 No")
                r3.font.size = Pt(9)
                r3.font.color.rgb = DARK

            # Choose meal template
            meals = REST_DAY_MEALS if is_weekend else LIFT_DAY_MEALS

            # Meal table
            meal_table = doc.add_table(rows=len(meals) + 1, cols=4)
            meal_table.style = "Table Grid"
            make_header_row(
                meal_table,
                ["Meal", "Food / Drink", "\u2714 Protein?", "Time"],
                font_size=8,
            )

            for i, (label, bg_color) in enumerate(meals):
                row = meal_table.rows[i + 1]
                row.cells[0].text = label
                set_cell_shading(row.cells[0], bg_color)
                set_row_height(row, 0.32)
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        r.bold = True
                # Protein checkbox
                row.cells[2].text = "\u2610"
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Column widths
                set_cell_width(row.cells[0], 1.05)
                set_cell_width(row.cells[1], 3.9)
                set_cell_width(row.cells[2], 0.65)
                set_cell_width(row.cells[3], 0.6)

            # Bottom tracking row
            track_p = doc.add_paragraph()
            track_items = [
                "Water (oz): ______",
                "Sleep (hrs): ______",
                "Soreness (1-5): ______",
                "Energy (1-5): ______",
            ]
            for idx, item in enumerate(track_items):
                r = track_p.add_run(item)
                r.font.size = Pt(8)
                r.bold = True
                if idx < len(track_items) - 1:
                    track_p.add_run("  ")

            # Page breaks: 2 days per page on lift days (8 meals = tall),
            # 3 days per page on weekends (6 meals = shorter)
            # Simple rule: break after every 2 days
            if day_idx < 6:
                add_horizontal_line(doc)
            if day_idx in [1, 3, 5]:
                doc.add_page_break()

        # Page break between weeks
        if week < 4:
            doc.add_page_break()

    # ================================================================
    # SAVE
    # ================================================================
    out = "River_Valley_Vikings_Offseason_Meal_Log_Form.docx"
    doc.save(out)
    print(f"Saved  {out}")


if __name__ == "__main__":
    main()
