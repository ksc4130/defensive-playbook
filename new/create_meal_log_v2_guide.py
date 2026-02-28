#!/usr/bin/env python3
"""
V2 Offseason Recovery Nutrition GUIDE (information only, no log pages).
Optimized for 6 AM lift → school schedule.
Print this once per player. Pair with the separate meal-log form.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- Colors --
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_BLUE_HEX = "D6EAF8"
LIGHT_GOLD_HEX = "FFF3CC"
LIGHT_GRAY_HEX = "F2F2F2"
LIGHT_GREEN_HEX = "D5F5E3"
LIGHT_RED_HEX = "FADBD8"
NAVY_HEX = "002D62"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_row_height(row, inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_header_row(table, headers, font_size=10):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        set_cell_shading(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = WHITE
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.name = "Calibri"


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # -- Helpers --
    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def bold_body(bold_text, normal_text=""):
        p = doc.add_paragraph()
        r = p.add_run(bold_text)
        r.bold = True
        if normal_text:
            p.add_run(normal_text)
        return p

    def small_note(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        r.italic = True
        return p

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("OFFSEASON FUEL")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GOLD

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Recovery Nutrition Guide")
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run(
        '"You don\'t get stronger in the weight room. '
        'You get stronger when you recover from the weight room."'
    )
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    doc.add_page_break()

    # ================================================================
    # SECTION 1: YOUR SCHEDULE IS YOUR ADVANTAGE
    # ================================================================
    heading1("Your Schedule Is Your Advantage")

    body(
        "You lift at 6 AM. That means you are doing something most high school "
        "athletes never will — training before the sun is up, before school, "
        "before anyone else is awake. That commitment is rare. But the work in "
        "the weight room is only half the equation."
    )

    body(
        "Lifting tears your muscles down. Recovery builds them back stronger. "
        "And recovery runs on food, water, and sleep. If you lift hard at 6 AM "
        "and then skip meals, eat junk, or sleep five hours — you are wasting "
        "that early alarm. The guys who make real gains between now and August "
        "are the ones who fuel the recovery."
    )

    heading3("Your Day at a Glance")

    glance_table = doc.add_table(rows=9, cols=2)
    glance_table.style = "Table Grid"
    make_header_row(glance_table, ["Time", "What Happens"])

    glance_data = [
        ("~5:10 AM", "Wake up. Water (16 oz) + quick bite if you can."),
        ("5:30 AM", "Travel to school / weight room."),
        ("6:00-7:00 AM", "LIFT. Sip water throughout."),
        ("7:00-7:30 AM", "POST-LIFT FUEL. This is the most important meal of your day."),
        ("~8:00 AM - 3:00 PM", "School. Eat a real breakfast, a mid-morning snack, and a real lunch."),
        ("3:00-4:00 PM", "After-school snack. Protein + carbs. Rebuild continues."),
        ("6:00-7:00 PM", "Dinner. Full plate. Your body is still repairing."),
        ("8:30-9:00 PM", "Bedtime fuel. Slow protein (milk, yogurt, cottage cheese). Lights out by 9:15."),
    ]
    for i, (time, what) in enumerate(glance_data):
        row = glance_table.rows[i + 1]
        row.cells[0].text = time
        row.cells[1].text = what
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.bold = True
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        # Color code key moments
        if i == 0:  # wake
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        elif i == 2:  # lift
            set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        elif i == 3:  # post-lift
            set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)
            set_cell_shading(row.cells[1], LIGHT_GREEN_HEX)
        elif i == 7:  # bedtime
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 2: THE SCIENCE (SIMPLE VERSION)
    # ================================================================
    heading1("Why Food = Strength Gains")

    body(
        "Lifting creates micro-tears in your muscle fibers. That sounds bad, but "
        "it's the whole point — when those fibers heal, they come back thicker "
        "and stronger. But healing requires raw materials. No materials, no "
        "rebuilding. You just broke yourself down for nothing."
    )

    sci_table = doc.add_table(rows=5, cols=2)
    sci_table.style = "Table Grid"
    make_header_row(sci_table, ["What Happens During/After Lifting", "What Your Body Needs"])

    sci_data = [
        (
            "Muscle fibers tear (micro-damage)",
            "PROTEIN rebuilds fibers thicker and stronger. "
            "Without it, you tore muscle for nothing.",
        ),
        (
            "Glycogen (stored energy) gets depleted",
            "CARBS refill glycogen so you have energy for "
            "the next session and don't feel flat.",
        ),
        (
            "Cortisol (stress hormone) spikes",
            "FOOD within 45 min post-lift brings cortisol "
            "down and flips your body into rebuild mode.",
        ),
        (
            "Growth hormone releases during sleep",
            "SLEEP (8+ hrs) is when your body does the "
            "actual construction. Cut sleep = cut gains.",
        ),
    ]
    for i, (what, need) in enumerate(sci_data):
        row = sci_table.rows[i + 1]
        row.cells[0].text = what
        row.cells[1].text = need
        set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    body("")
    bold_body(
        "Bottom line: ",
        "Lifting is the stimulus. Food is the builder. Sleep is the "
        "construction crew. You need all three — and your 6 AM schedule "
        "means you have to be intentional about all of them.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 3: THE RECOVERY PLATE
    # ================================================================
    heading1("The Recovery Plate")

    body(
        "You don't need to count calories or weigh food. Just make sure every "
        "meal has protein and carbs — those are the two non-negotiables. Use "
        "your hands to estimate portions."
    )

    heading3("Hand Portions")
    bullet("Palm = 1 serving of protein (chicken, beef, fish, eggs)")
    bullet("Fist = 1 serving of carbs (rice, pasta, potatoes, oatmeal)")
    bullet("Cupped hand = 1 serving of fruit or snack")
    bullet("Thumb = 1 serving of fats (butter, oil, peanut butter)")
    bullet("Two fists = vegetables (unlimited)")

    body("")

    plate_table = doc.add_table(rows=5, cols=3)
    plate_table.style = "Table Grid"
    make_header_row(plate_table, ["Pillar", "Why It Matters for Recovery", "Best Sources"])

    plate_data = [
        [
            "PROTEIN\n(Rebuild)",
            "Repairs torn muscle fibers. Without\nprotein your body can't rebuild\nwhat you broke down.",
            "Chicken, beef, turkey, eggs,\nGreek yogurt, milk, tuna,\nprotein shake, cottage cheese",
        ],
        [
            "CARBS\n(Reload)",
            "Refills glycogen so you're not\nrunning on empty. Prevents your\nbody from burning muscle for fuel.",
            "Rice, potatoes, oatmeal, pasta,\nbread, bananas, tortillas,\nfruit, cereal, granola",
        ],
        [
            "VEGETABLES\n(Recover)",
            "Vitamins, minerals, and anti-\ninflammatory compounds that reduce\nsoreness and speed healing.",
            "Broccoli, spinach, peppers,\ncarrots, green beans, salad,\nsweet potatoes",
        ],
        [
            "FATS\n(Regulate)",
            "Supports testosterone and growth\nhormone production. Keeps joints\nhealthy under heavy loads.",
            "Peanut butter, avocado, nuts,\ncheese, olive oil, whole eggs,\nwhole milk",
        ],
    ]
    for i, row_data in enumerate(plate_data):
        row = plate_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, LIGHT_GOLD_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    body("")
    bold_body(
        "Minimum per meal: ",
        "1 palm of protein + 1 fist of carbs. That is the floor.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 4: DAILY RECOVERY TARGETS
    # ================================================================
    heading1("Daily Recovery Targets")

    body(
        "These are simple guidelines. If you are consistently sore, tired, or "
        "stalling on your lifts, you are probably not eating enough."
    )

    target_table = doc.add_table(rows=4, cols=3)
    target_table.style = "Table Grid"
    make_header_row(target_table, ["Nutrient", "Daily Target", "Why"])

    target_data = [
        [
            "Protein",
            "1g per pound of bodyweight\n(180 lb = ~180g protein/day)",
            "Rebuilds muscle. #1 priority.\nSpread across every meal.",
        ],
        [
            "Carbs",
            "2-3g per pound of bodyweight\n(180 lb = ~360-540g carbs/day)",
            "Refuels glycogen. More on lift\ndays. Carbs are not the enemy.",
        ],
        [
            "Water",
            "Half your bodyweight in ounces\n(180 lb = 90 oz minimum/day)",
            "Muscle is 75% water. Dehydration =\nweaker lifts, slower recovery.",
        ],
    ]
    for i, row_data in enumerate(target_data):
        row = target_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    body("")
    heading3("Protein Cheat Sheet")

    protein_table = doc.add_table(rows=11, cols=2)
    protein_table.style = "Table Grid"
    make_header_row(protein_table, ["Food", "Protein"])

    protein_data = [
        ("Chicken breast (1 palm / ~4 oz)", "~30g"),
        ("Ground beef or turkey (1 palm / ~4 oz)", "~28g"),
        ("3 whole eggs", "~18g"),
        ("Scoop of protein powder (in water or milk)", "~25g"),
        ("Greek yogurt (1 cup)", "~15-20g"),
        ("Can of tuna", "~20g"),
        ("Glass of whole milk (8 oz)", "~8g"),
        ("Chocolate milk (8 oz)", "~8g"),
        ("PB&J sandwich", "~12g"),
        ("String cheese (1 stick)", "~7g"),
    ]
    for i, (food, protein) in enumerate(protein_data):
        row = protein_table.rows[i + 1]
        row.cells[0].text = food
        row.cells[1].text = protein
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)
            set_cell_shading(row.cells[1], LIGHT_GRAY_HEX)

    small_note(
        "A 180 lb player needs ~180g protein/day. That's roughly 6 palms spread "
        "across your meals. It adds up fast if you eat protein at every meal."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 5: LIFT-DAY FUELING (6 AM schedule)
    # ================================================================
    heading1("Lift-Day Fueling Schedule")

    body(
        "Your 6 AM lift changes everything about meal timing. You are training "
        "in a fasted or near-fasted state — that makes what you eat the night "
        "before and immediately after the lift the two most critical meals of "
        "your day. Every other meal supports the recovery that started at 7 AM."
    )

    heading2("The 7-Meal Lift Day")

    timing_table = doc.add_table(rows=8, cols=3)
    timing_table.style = "Table Grid"
    make_header_row(timing_table, ["Meal", "Time", "What & Why"])

    timing_data = [
        [
            "WAKE-UP\nBITE",
            "~5:10 AM",
            "Small, fast-digesting. Just enough to blunt cortisol\n"
            "and give your brain glucose. NOT a full meal.\n"
            "Ex: banana, granola bar, handful of cereal, glass\n"
            "of juice or milk, spoonful of peanut butter + bread",
        ],
        [
            "LIFT",
            "6:00-7:00 AM",
            "Sip water throughout. 16-24 oz minimum during\n"
            "the session. You woke up dehydrated — fix that.",
        ],
        [
            "POST-LIFT\nFUEL",
            "7:00-7:30 AM\n(within 30 min)",
            "MOST IMPORTANT MEAL OF THE DAY.\n"
            "Protein + fast carbs. Portable — eat in the car,\n"
            "hallway, or first period. This is non-negotiable.\n"
            "Ex: protein shake + banana, chocolate milk + PB&J,\n"
            "Greek yogurt + granola bar",
        ],
        [
            "SCHOOL\nBREAKFAST",
            "~8:00-8:30 AM",
            "Full meal. Protein + carbs + fruit. This is your\n"
            "second recovery meal — your muscles are still\n"
            "absorbing nutrients from the session.\n"
            "Ex: eggs + toast + fruit, oatmeal + milk + banana",
        ],
        [
            "MID-MORNING\nSNACK",
            "~10:00-10:30 AM\n(between classes)",
            "Protein + carb. Keep it in your bag.\n"
            "Ex: PB&J, string cheese + crackers, trail mix,\n"
            "protein bar, beef jerky + granola bar",
        ],
        [
            "LUNCH",
            "~11:30-12:30 PM",
            "Full plate: protein + carbs + veggies. Biggest\n"
            "school meal. Not chips and a cookie.\n"
            "Ex: chicken + rice + green beans, sub sandwich\n"
            "+ fruit, burger + side + milk",
        ],
        [
            "AFTER-\nSCHOOL",
            "~3:00-4:00 PM",
            "Protein + carbs. Recovery is still happening.\n"
            "Bridge the gap to dinner.\n"
            "Ex: PB&J + milk, leftovers, yogurt + fruit,\n"
            "protein shake + crackers",
        ],
    ]
    for i, row_data in enumerate(timing_data):
        row = timing_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if j == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
        # Color code
        if i == 0:  # wake-up bite
            set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        elif i == 1:  # lift
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        elif i == 2:  # post-lift
            set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)
            set_cell_shading(row.cells[1], LIGHT_GREEN_HEX)
        elif i == 3:  # school breakfast
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        else:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)

    doc.add_page_break()

    # Dinner + bedtime on next page (continuation)
    heading2("Evening Recovery (The Hours That Build You)")

    body(
        "Dinner and bedtime fuel are where most of the overnight repair happens. "
        "Growth hormone peaks during deep sleep — and it needs protein and carbs "
        "to do its job."
    )

    eve_table = doc.add_table(rows=3, cols=3)
    eve_table.style = "Table Grid"
    make_header_row(eve_table, ["Meal", "Time", "What & Why"])

    eve_data = [
        [
            "DINNER",
            "~6:00-7:00 PM",
            "Full plate: protein + carbs + veggies + fats.\n"
            "This is your largest meal. Your body is actively\n"
            "repairing from this morning's session.\n"
            "Ex: beef + pasta + salad + bread + glass of milk,\n"
            "chicken + rice + broccoli + butter",
        ],
        [
            "BEDTIME\nFUEL",
            "~8:30-9:00 PM\n(30 min before\nsleep)",
            "Slow-digesting protein. Feeds muscle repair\n"
            "overnight while you sleep. This is NOT junk food.\n"
            "Ex: cottage cheese + fruit, Greek yogurt,\n"
            "glass of milk + peanut butter on toast,\n"
            "casein shake, string cheese + crackers",
        ],
    ]
    for i, row_data in enumerate(eve_data):
        row = eve_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if j == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
        if i == 0:
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        elif i == 1:
            set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
            set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)

    body("")
    bold_body(
        "Why bedtime fuel matters: ",
        "You're about to fast for 8 hours, then lift at 6 AM on whatever is "
        "left in your system. Slow-digesting protein (casein — found in milk, "
        "cottage cheese, Greek yogurt) feeds your muscles while you sleep. "
        "This is free gains."
    )

    body("")
    heading3("Non-Lift Days (Weekends & Rest Days)")
    body(
        "Your body is still rebuilding from the previous session. Recovery does "
        "not stop because you are not in the weight room. Eat the same meals — "
        "skip the wake-up bite and post-lift fuel, but keep breakfast, lunch, "
        "snacks, dinner, and bedtime fuel. Do not use rest days as an excuse "
        "to eat less."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 6: POST-LIFT FUEL (dedicated section)
    # ================================================================
    heading1("Post-Lift Fuel: The One You Can't Skip")

    body(
        "In the 30-45 minutes after lifting, your muscles are primed to absorb "
        "protein and carbs faster than any other time of day. Cortisol (the "
        "stress hormone that breaks down muscle) is spiking — food brings it "
        "down and flips your body into rebuild mode. Skip this and your body "
        "stays in breakdown mode through first period."
    )

    heading3("The Rules")
    bullet("Eat within 30 minutes of your last rep. Don't wait until you get to class.")
    bullet("20-40g protein + 30-60g fast carbs. That's the formula.")
    bullet("It has to be PORTABLE. You're eating this in the car or hallway.")
    bullet("This is not the time for a salad or a big sit-down meal. Fast fuel.")

    body("")
    heading3("Grab-and-Go Post-Lift Combos")

    combo_table = doc.add_table(rows=9, cols=3)
    combo_table.style = "Table Grid"
    make_header_row(combo_table, ["Option", "Protein", "Carbs"])

    combos = [
        ("Protein shake (milk) + banana", "~30g", "~40g"),
        ("Chocolate milk (16 oz) + PB&J", "~16g", "~60g"),
        ("Greek yogurt + granola bar", "~22g", "~35g"),
        ("Protein bar + banana", "~20g", "~45g"),
        ("2 string cheese + apple + granola bar", "~15g", "~45g"),
        ("Deli turkey wrap + chocolate milk", "~28g", "~45g"),
        ("Overnight oats w/ protein powder (premade)", "~30g", "~50g"),
        ("Cottage cheese cup + fruit + crackers", "~20g", "~35g"),
    ]
    for i, (option, prot, carb) in enumerate(combos):
        row = combo_table.rows[i + 1]
        row.cells[0].text = option
        row.cells[1].text = prot
        row.cells[2].text = carb
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GREEN_HEX)

    body("")
    small_note(
        "Pack your post-lift fuel the night before. Put it in your bag next to "
        "your shoes. Make it impossible to forget."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 7: SAMPLE FULL DAY
    # ================================================================
    heading1("Sample Lift-Day (Full Day)")

    body(
        "Here's what a complete day looks like for a 6 AM lifter whose only "
        "goal is getting stronger. This is not a diet. It's fueling."
    )

    sample_table = doc.add_table(rows=10, cols=4)
    sample_table.style = "Table Grid"
    make_header_row(sample_table, ["Time", "Meal", "Example", "Protein"])

    sample_data = [
        ("5:15 AM", "Wake-Up Bite", "Banana + glass of milk", "~10g"),
        ("6:00-7:00", "Lift", "Water (16-24 oz during session)", "—"),
        ("7:10 AM", "Post-Lift Fuel", "Protein shake (milk) + banana", "~33g"),
        ("8:15 AM", "School Breakfast", "3 eggs + 2 toast + juice", "~22g"),
        ("10:15 AM", "Snack", "PB&J + string cheese", "~19g"),
        ("12:00 PM", "Lunch", "Chicken + rice + green beans + milk", "~38g"),
        ("3:15 PM", "After School", "Greek yogurt + granola + banana", "~18g"),
        ("6:30 PM", "Dinner", "Ground beef + pasta + salad + bread + milk", "~42g"),
        ("8:45 PM", "Bedtime Fuel", "Cottage cheese + berries", "~15g"),
    ]
    for i, (time, meal, example, protein) in enumerate(sample_data):
        row = sample_table.rows[i + 1]
        row.cells[0].text = time
        row.cells[1].text = meal
        row.cells[2].text = example
        row.cells[3].text = protein
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.bold = True
        # Color code
        if i == 0:
            set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)
        elif i == 1:
            set_cell_shading(row.cells[1], LIGHT_BLUE_HEX)
        elif i == 2:
            set_cell_shading(row.cells[1], LIGHT_GREEN_HEX)
            set_cell_shading(row.cells[2], LIGHT_GREEN_HEX)
        elif i == 8:
            set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)

    body("")
    p = doc.add_paragraph()
    r = p.add_run("Daily total: ~197g protein across 8 eating opportunities.")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        " That's enough for a 190+ lb player. The key is frequency — protein "
        "at every meal, never skipping the post-lift window."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 8: RECOVERY KILLERS
    # ================================================================
    heading1("Recovery Killers")

    body(
        "These habits undo your work in the weight room. You put in the reps "
        "at 6 AM — don't waste them."
    )

    killer_table = doc.add_table(rows=9, cols=2)
    killer_table.style = "Table Grid"
    make_header_row(killer_table, ["Recovery Killer", "What It Costs You"])

    killers = [
        (
            "Skipping post-lift fuel",
            "Your #1 recovery window — gone. Cortisol stays elevated, "
            "muscle breakdown continues. You lifted for nothing.",
        ),
        (
            "Not eating until lunch",
            "You lifted at 6 AM. By noon, you've gone 5+ hours with no "
            "protein. Your body is eating its own muscle for fuel.",
        ),
        (
            "Sleeping less than 8 hours\n(bed after 9:15 PM)",
            "Growth hormone peaks in deep sleep. A 5 AM alarm with a "
            "midnight bedtime = 5 hours = half the recovery.",
        ),
        (
            "Eating 1-2 meals/day",
            "Not enough total fuel. You'll be sore longer, weaker "
            "next session, and plateau within weeks.",
        ),
        (
            "No protein at meals",
            "A plate of pasta with no meat is just energy — no building "
            "blocks. Every meal needs a palm of protein.",
        ),
        (
            "Soda / energy drinks as\nmain beverages",
            "Empty calories, dehydration, sugar crashes. Water and "
            "milk are the only drinks that help recovery.",
        ),
        (
            "Dehydration",
            "You wake up dehydrated, then lift. If you don't drink "
            "water immediately and all day, recovery tanks.",
        ),
        (
            "No bedtime fuel",
            "You're about to fast 8 hours then lift fasted. Going to "
            "bed on empty = going to the rack on empty.",
        ),
    ]
    for i, (killer, effect) in enumerate(killers):
        row = killer_table.rows[i + 1]
        row.cells[0].text = killer
        row.cells[1].text = effect
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_RED_HEX)
            set_cell_shading(row.cells[1], LIGHT_RED_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 9: BUDGET-FRIENDLY RECOVERY FOODS
    # ================================================================
    heading1("Budget-Friendly Recovery Foods")

    body(
        "You don't need supplements or expensive groceries. These staples "
        "are cheap, effective, and available at any store."
    )

    budget_table = doc.add_table(rows=11, cols=3)
    budget_table.style = "Table Grid"
    make_header_row(budget_table, ["Food", "Cost", "Why It Works"])

    budget_data = [
        ("Eggs (1 dozen)", "~$3", "Best protein per dollar. 6g per egg. Easy to cook."),
        ("Rice (5 lb bag)", "~$4", "Carbs for an entire week. Pairs with everything."),
        ("Peanut butter (jar)", "~$3", "Protein + healthy fats + calories. Lasts forever."),
        ("Oatmeal (canister)", "~$4", "Breakfast for weeks. Add milk and PB for protein."),
        ("Whole milk (gallon)", "~$4", "8g protein per glass. Cheap easy calories."),
        ("Frozen chicken (3 lb bag)", "~$8", "Bulk protein. Bake a batch on Sunday."),
        ("Bananas (each)", "~$0.25", "Perfect portable carb. Pre-lift and post-lift."),
        ("Chocolate milk (carton)", "~$1", "Best cheap post-workout drink. Right ratio."),
        ("Canned tuna (can)", "~$1", "20g protein, no cooking, fits in your bag."),
        ("Beans / lentils (can)", "~$1", "Protein + carbs + fiber. Stretch any meal."),
    ]
    for i, (food, cost, why) in enumerate(budget_data):
        row = budget_table.rows[i + 1]
        row.cells[0].text = food
        row.cells[1].text = cost
        row.cells[2].text = why
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 10: SLEEP & HYDRATION
    # ================================================================
    heading1("Sleep & Hydration")

    heading2("Sleep: The Free Performance Enhancer")

    body(
        "With a 5:10 AM alarm, you need to be asleep by 9:00-9:15 PM to get "
        "8 hours. That is not a suggestion — it is the math. Growth hormone "
        "peaks during stages 3 and 4 of deep sleep. Cut your sleep to 6 hours "
        "and you are cutting your recovery nearly in half."
    )

    sleep_table = doc.add_table(rows=5, cols=2)
    sleep_table.style = "Table Grid"
    make_header_row(sleep_table, ["Sleep Rule", "Why"])

    sleep_data = [
        ("In bed by 9:00 PM, lights out by 9:15", "8 hours before a 5:10 AM alarm. Non-negotiable."),
        ("Phone down at 8:30 PM", "Blue light from screens delays melatonin by 30-60 min."),
        ("No caffeine after 2:00 PM", "Caffeine half-life is 5-6 hours. 3 PM coffee is still in you at 9."),
        ("Dark, cool room", "Your brain sleeps deeper when it's dark and below 68 degrees."),
    ]
    for i, (rule, why) in enumerate(sleep_data):
        row = sleep_table.rows[i + 1]
        row.cells[0].text = rule
        row.cells[1].text = why
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_BLUE_HEX)

    body("")
    heading2("Hydration")

    body(
        "You wake up dehydrated from 8 hours of no water, then immediately go "
        "lift. If you don't start drinking water the moment you wake up, you are "
        "training in a deficit."
    )

    bullet("16 oz water as soon as you wake up (before anything else).", bold_prefix="5:10 AM: ")
    bullet("16-24 oz during the lift session.", bold_prefix="During lift: ")
    bullet("Water with every meal and between classes.", bold_prefix="All day: ")
    bullet("Half your bodyweight in ounces by end of day (180 lb = 90 oz).", bold_prefix="Daily target: ")
    bullet("Soda, energy drinks, and juice are not hydration. Water and milk.", bold_prefix="Cut the junk: ")
    bullet("If your pee is dark yellow, you're behind.", bold_prefix="Self-check: ")

    doc.add_page_break()

    # ================================================================
    # SECTION 11: DAILY RECOVERY CHECKLIST
    # ================================================================
    heading1("Daily Recovery Checklist")

    body(
        "Nutrition is the biggest piece, but recovery is a system. Hit these "
        "every day and your lifts will go up."
    )

    check_table = doc.add_table(rows=10, cols=2)
    check_table.style = "Table Grid"
    make_header_row(check_table, ["Checklist Item", "Target"])

    checks = [
        ("\u2610  Drank 16 oz water at wake-up", "Before you do anything else"),
        ("\u2610  Ate wake-up bite before lift", "Small carb + protein at 5:15 AM"),
        ("\u2610  Ate post-lift fuel within 30 min", "Protein + carbs. Non-negotiable."),
        ("\u2610  Ate school breakfast", "Second recovery meal. Full meal with protein."),
        ("\u2610  Had protein at every meal", "1 palm minimum per meal, every meal"),
        ("\u2610  Ate 5+ times today", "Wake-up, post-lift, breakfast, snack, lunch, after-school, dinner, bedtime"),
        ("\u2610  Drank enough water", "Half bodyweight in ounces"),
        ("\u2610  Ate bedtime fuel with protein", "Slow protein before sleep (milk, yogurt, cottage cheese)"),
        ("\u2610  In bed by 9:00 PM, lights out 9:15", "8 hours of sleep = full recovery"),
    ]
    for i, (check, target) in enumerate(checks):
        row = check_table.rows[i + 1]
        row.cells[0].text = check
        row.cells[1].text = target
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i in [2, 7, 8]:  # highlight critical items
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GREEN_HEX)
        elif i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 12: COMMITMENT
    # ================================================================
    heading1("My Commitment")

    body("")
    body(
        "I understand that my body gets stronger through recovery, not just "
        "lifting. I commit to fueling my recovery by eating enough protein, "
        "carbs, and whole foods every day — especially after I train. I will "
        "pack my post-lift fuel the night before, eat throughout the school "
        "day, fuel before bed, and be asleep by 9:15 PM. I will not waste "
        "my 6 AM work by neglecting my nutrition."
    )

    body("")

    heading3("My Offseason Strength Goals")
    goal_table = doc.add_table(rows=2, cols=4)
    goal_table.style = "Table Grid"
    make_header_row(goal_table, ["Squat Goal", "Bench Goal", "Clean Goal", "Weight Goal"])
    set_row_height(goal_table.rows[1], 0.45)

    body("")

    # Player info
    heading3("Player Information")
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    info_labels = [
        ("Name:", "", "Position:", ""),
        ("Grade:", "", "Height:", ""),
        ("Current Weight:", "", "Start Date:", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_labels):
        row = info_table.rows[i]
        set_row_height(row, 0.35)
        c0, c1, c2, c3 = row.cells
        set_cell_shading(c0, LIGHT_BLUE_HEX)
        set_cell_shading(c2, LIGHT_BLUE_HEX)
        c0.text = l1
        c1.text = v1
        c2.text = l2
        c3.text = v2
        for cell in [c0, c2]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

    body("")

    # Starting lifts
    heading3("Starting Lift Numbers")
    lift_table = doc.add_table(rows=2, cols=5)
    lift_table.style = "Table Grid"
    make_header_row(lift_table, ["Squat", "Bench", "Deadlift", "Power Clean", "Bodyweight"])
    set_row_height(lift_table.rows[1], 0.4)

    body("")
    body("")

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"
    for i, (l1, l2) in enumerate([("Player Signature:", "Date:"), ("Coach Signature:", "Date:")]):
        row = sig_table.rows[i]
        set_row_height(row, 0.6)
        row.cells[0].text = l1
        row.cells[1].text = l2
        set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

    body("")
    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = close.add_run(
        '"The alarm goes off at 5. The gains happen in the 23 hours after."'
    )
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    # ================================================================
    # SAVE
    # ================================================================
    out = "River_Valley_Vikings_Offseason_Fuel_Guide.docx"
    doc.save(out)
    print(f"Saved  {out}")


if __name__ == "__main__":
    main()
