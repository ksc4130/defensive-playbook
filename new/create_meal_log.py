#!/usr/bin/env python3
"""Generate a meal-log handout for River Valley Vikings football players."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- Colors (Columbia Blue / Gold / White / Navy) --
NAVY = RGBColor(0x00, 0x2D, 0x62)
COLUMBIA_BLUE = RGBColor(0x6C, 0xAC, 0xE4)
GOLD = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_BLUE_HEX = "D6EAF8"
LIGHT_GOLD_HEX = "FFF3CC"
LIGHT_GRAY_HEX = "F2F2F2"
NAVY_HEX = "002D62"
GOLD_HEX = "CFA700"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_row_height(row, inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_header_row(table, headers, font_size=10):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        set_cell_shading(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = WHITE
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.name = "Calibri"


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # -- Helpers --
    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def bold_body(bold_text, normal_text=""):
        p = doc.add_paragraph()
        r = p.add_run(bold_text)
        r.bold = True
        if normal_text:
            p.add_run(normal_text)
        return p

    def small_note(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        r.italic = True
        return p

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("FOOTBALL NUTRITION")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GOLD

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Meal Log & Fueling Guide")
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run('"You can\'t out-train a bad diet. Fuel the machine."')
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    doc.add_page_break()

    # ================================================================
    # SECTION 1: PLAYER INFO
    # ================================================================
    heading1("Player Information")

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = "Table Grid"
    info_labels = [
        ("Name:", "", "Position:", ""),
        ("Grade:", "", "Height:", ""),
        ("Starting Weight:", "", "Goal Weight:", ""),
        ("Start Date:", "", "Target Date:", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_labels):
        row = info_table.rows[i]
        set_row_height(row, 0.35)
        c0, c1, c2, c3 = row.cells
        set_cell_shading(c0, LIGHT_BLUE_HEX)
        set_cell_shading(c2, LIGHT_BLUE_HEX)
        for cell in [c0, c2]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        c0.text = l1
        c1.text = v1
        c2.text = l2
        c3.text = v2

    body("")

    # Goal checkboxes
    bold_body("My Goal (check one):")
    goal_table = doc.add_table(rows=1, cols=3)
    goal_table.style = "Table Grid"
    goals = [
        "\u2610  Gain Weight / Add Mass",
        "\u2610  Lose Body Fat / Lean Out",
        "\u2610  Maintain / Recomposition",
    ]
    for j, g in enumerate(goals):
        cell = goal_table.rows[0].cells[j]
        cell.text = g
        set_cell_shading(cell, LIGHT_GOLD_HEX)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_page_break()

    # ================================================================
    # SECTION 2: HOW TO USE THIS LOG
    # ================================================================
    heading1("How to Use This Log")

    body(
        "This log is your accountability partner. Tracking what you eat makes you "
        "aware of your habits, helps coaches guide you, and keeps you honest with "
        "yourself. The players who commit to fueling right are the ones who show up "
        "bigger, faster, and more prepared in August."
    )

    heading3("The Rules")
    bullet("Write down everything you eat and drink. No skipping meals, no hiding junk.")
    bullet("Log your meals the same day. Don't try to remember three days later.")
    bullet("Be honest. This is for YOU. Nobody is grading your food — we're building habits.")
    bullet("Bring your log to weekly weigh-ins. Your position coach will check it.")
    bullet("Water intake matters. Track it daily. Minimum 80 oz (half your bodyweight in ounces is the goal).")

    heading3("What Counts as a Serving?")
    body("You don't need a food scale. Use your hands:")
    bullet("Palm = 1 serving of protein (chicken, beef, fish, eggs)")
    bullet("Fist = 1 serving of carbs (rice, pasta, potatoes, oatmeal)")
    bullet("Cupped hand = 1 serving of fruit or snack")
    bullet("Thumb = 1 serving of fats (butter, oil, peanut butter)")
    bullet("Two fists = 1 serving of vegetables (eat as much as you want)")

    doc.add_page_break()

    # ================================================================
    # SECTION 3: NUTRITION GUIDELINES
    # ================================================================
    heading1("Vikings Nutrition Guidelines")

    heading2("Daily Targets by Goal")
    body(
        "These are general guidelines. Every player is different. The important thing "
        "is consistency — hit these targets most days and you will see results."
    )

    # Target table
    target_table = doc.add_table(rows=4, cols=5)
    target_table.style = "Table Grid"
    make_header_row(target_table, ["Goal", "Protein", "Carbs", "Fats", "Calories (approx.)"])

    target_data = [
        ["Gain Weight", "1g per lb BW\n(e.g., 180 lb = 180g)", "2-3g per lb BW", "0.5g per lb BW", "BW x 18-20"],
        ["Lose Fat", "1.2g per lb BW", "1-1.5g per lb BW", "0.4g per lb BW", "BW x 13-15"],
        ["Maintain", "1g per lb BW", "1.5-2g per lb BW", "0.45g per lb BW", "BW x 15-17"],
    ]
    for i, row_data in enumerate(target_data):
        row = target_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    small_note("BW = body weight in pounds. These are starting points — adjust based on results after 2-3 weeks.")

    body("")

    heading2("The Vikings Plate")
    body("At every meal, build your plate like this:")

    plate_table = doc.add_table(rows=5, cols=3)
    plate_table.style = "Table Grid"
    make_header_row(plate_table, ["Section", "What to Eat", "How Much"])

    plate_data = [
        ["Protein\n(Build muscle)", "Chicken, ground beef/turkey, eggs,\nfish, Greek yogurt, protein shake", "1-2 palms"],
        ["Carbs\n(Fuel performance)", "Rice, potatoes, oatmeal, pasta,\nbread, fruit, tortillas", "1-2 fists"],
        ["Vegetables\n(Recover & stay healthy)", "Broccoli, spinach, peppers, green\nbeans, salad, carrots", "2+ fists (unlimited)"],
        ["Fats\n(Hormones & energy)", "Peanut butter, avocado, cheese,\nnuts, olive oil, butter", "1-2 thumbs"],
    ]
    for i, row_data in enumerate(plate_data):
        row = plate_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, LIGHT_GOLD_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_page_break()

    # ================================================================
    # SECTION 4: MEAL TIMING
    # ================================================================
    heading1("Meal Timing for Football Players")

    body(
        "Eat 4-5 times per day. Do not skip breakfast. Do not skip your post-workout "
        "meal. These two meals are the most important for building and recovering."
    )

    timing_table = doc.add_table(rows=7, cols=3)
    timing_table.style = "Table Grid"
    make_header_row(timing_table, ["Meal", "When", "What to Prioritize"])

    timing_data = [
        ["Breakfast", "Within 30 min\nof waking up", "Protein + carbs. Sets the tone for the day.\nExamples: eggs + toast, oatmeal + protein shake"],
        ["Snack 1", "Mid-morning\nor before lunch", "Protein + carb. Keep it simple.\nExamples: PB&J, Greek yogurt + granola, trail mix"],
        ["Lunch", "Noon", "Full plate: protein + carbs + veggies.\nThis is a full meal, not chips and a cookie."],
        ["Pre-Workout", "60-90 min\nbefore training", "Carbs + small protein. Easy to digest.\nExamples: banana + PB, granola bar, rice + chicken"],
        ["Post-Workout", "Within 45 min\nafter training", "Protein + fast carbs. Recovery window.\nExamples: protein shake + banana, chocolate milk"],
        ["Dinner", "Evening", "Full plate: protein + carbs + veggies + fats.\nBiggest meal if you're trying to gain weight."],
    ]
    for i, row_data in enumerate(timing_data):
        row = timing_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(10)

    body("")
    heading3("Hydration")
    bullet("Minimum 80 oz of water per day (about 5 standard water bottles).")
    bullet("Goal: half your body weight in ounces (200 lb player = 100 oz).")
    bullet("Drink water with every meal and throughout practice.")
    bullet("Limit soda, energy drinks, and juice. They are not hydration.", bold_prefix="")
    bullet("If your pee is dark yellow, you are dehydrated. Clear to light yellow = good.", bold_prefix="")

    doc.add_page_break()

    # ================================================================
    # SECTION 5: GOOD vs BAD CHOICES
    # ================================================================
    heading1("Smart Choices vs. Bad Habits")

    body(
        "You don't have to be perfect. You have to be consistent. But know what "
        "helps you and what hurts you."
    )

    choice_table = doc.add_table(rows=9, cols=2)
    choice_table.style = "Table Grid"
    make_header_row(choice_table, ["Eat More Of (Fuel)", "Eat Less Of (Junk)"])

    choices = [
        ("Grilled/baked chicken, turkey, beef", "Fried foods (fried chicken, fries every day)"),
        ("Eggs (whole eggs are fine)", "Fast food as a daily habit"),
        ("Rice, potatoes, oatmeal, whole grain bread", "Candy, cookies, chips as meal replacements"),
        ("Fruits: bananas, apples, berries", "Soda and energy drinks (Monster, etc.)"),
        ("Vegetables at every meal", "Skipping meals (especially breakfast)"),
        ("Water, milk, protein shakes", "Ice cream / milkshakes as daily snacks"),
        ("Peanut butter, nuts, avocado", "Gas station food as your main nutrition"),
        ("Greek yogurt, cottage cheese, string cheese", "Eating only one meal a day"),
    ]
    for i, (good, bad) in enumerate(choices):
        row = choice_table.rows[i + 1]
        row.cells[0].text = good
        row.cells[1].text = bad
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)
            set_cell_shading(row.cells[1], LIGHT_GRAY_HEX)

    body("")
    heading3("Budget-Friendly Staples")
    body(
        "We know not everyone has the same grocery budget. These are cheap, effective, "
        "and available everywhere:"
    )
    bullet("Eggs (dozen for ~$3) — best protein per dollar")
    bullet("Rice (5 lb bag for ~$4) — carbs for days")
    bullet("Beans/lentils (canned or dry) — protein + carbs + fiber")
    bullet("Peanut butter — healthy fats + protein, lasts forever")
    bullet("Oatmeal (big canister for ~$4) — breakfast every day")
    bullet("Bananas (~$0.25 each) — pre/post workout carbs")
    bullet("Whole milk — calories + protein for weight gain")
    bullet("Frozen chicken breasts or thighs — buy in bulk")
    bullet("Frozen vegetables — just as nutritious as fresh, cheaper")
    bullet("Tuna cans — cheap portable protein")

    doc.add_page_break()

    # ================================================================
    # SECTION 6: SAMPLE MEAL PLANS
    # ================================================================
    heading1("Sample Meal Plans")

    # --- GAIN WEIGHT ---
    heading2("Gaining Weight (~3,000-3,500 cal)")
    gain_table = doc.add_table(rows=7, cols=2)
    gain_table.style = "Table Grid"
    make_header_row(gain_table, ["Meal", "Example"])

    gain_meals = [
        ("Breakfast", "4 eggs scrambled + 2 pieces toast + banana + glass of milk"),
        ("Snack 1", "PB&J sandwich + protein shake"),
        ("Lunch", "Double portion chicken + rice + broccoli + water"),
        ("Pre-Workout", "Granola bar + banana + water"),
        ("Post-Workout", "Protein shake + chocolate milk + banana"),
        ("Dinner", "8oz ground beef + pasta + salad + bread + glass of milk"),
    ]
    for i, (meal, example) in enumerate(gain_meals):
        row = gain_table.rows[i + 1]
        row.cells[0].text = meal
        row.cells[1].text = example
        set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    small_note("Key: eat big at every meal. Add milk, peanut butter, and extra carbs to increase calories.")

    body("")

    # --- LOSE FAT ---
    heading2("Losing Fat / Leaning Out (~2,000-2,400 cal)")
    lose_table = doc.add_table(rows=7, cols=2)
    lose_table.style = "Table Grid"
    make_header_row(lose_table, ["Meal", "Example"])

    lose_meals = [
        ("Breakfast", "3 eggs + 1 slice toast + fruit"),
        ("Snack 1", "Greek yogurt + handful of almonds"),
        ("Lunch", "Grilled chicken + big salad + light dressing + water"),
        ("Pre-Workout", "Apple + small scoop peanut butter"),
        ("Post-Workout", "Protein shake (water, not milk)"),
        ("Dinner", "Grilled chicken or fish + veggies + small portion rice"),
    ]
    for i, (meal, example) in enumerate(lose_meals):
        row = lose_table.rows[i + 1]
        row.cells[0].text = meal
        row.cells[1].text = example
        set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    small_note("Key: protein stays high. Cut back on carbs and fats slightly. Eliminate liquid calories (soda, juice).")

    doc.add_page_break()

    # ================================================================
    # SECTION 7: WEEKLY WEIGH-IN TRACKER
    # ================================================================
    heading1("Weekly Weigh-In Tracker")
    body("Record your weight at the same time each week (morning, before eating).")

    weigh_table = doc.add_table(rows=13, cols=4)
    weigh_table.style = "Table Grid"
    make_header_row(weigh_table, ["Week", "Date", "Weight (lbs)", "Notes / How I Feel"])

    for i in range(1, 13):
        row = weigh_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_row_height(row, 0.35)
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 8: DAILY MEAL LOG PAGES (4 weeks = 28 days)
    # ================================================================
    heading1("Daily Meal Log")
    body(
        "Fill out one row per meal. Be specific — \"chicken and rice\" is better than "
        "\"food.\" Track your water intake at the bottom of each day."
    )

    add_horizontal_line(doc)

    for week in range(1, 5):
        heading2(f"Week {week}")

        for day in range(1, 8):
            day_names = ["", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
            day_name = day_names[day]

            # Day header
            day_head = doc.add_paragraph()
            r = day_head.add_run(f"Day {(week - 1) * 7 + day}: {day_name}")
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = NAVY

            r2 = day_head.add_run("     Date: _______________")
            r2.font.size = Pt(10)
            r2.font.color.rgb = GRAY

            # Meal table
            meal_table = doc.add_table(rows=7, cols=4)
            meal_table.style = "Table Grid"
            make_header_row(meal_table, ["Meal", "Food / Drink", "Protein?", "Time"], font_size=9)

            meal_labels = ["Breakfast", "Snack 1", "Lunch", "Pre-Workout", "Post-Workout", "Dinner"]
            for i, label in enumerate(meal_labels):
                row = meal_table.rows[i + 1]
                row.cells[0].text = label
                set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
                set_row_height(row, 0.4)
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.bold = True
                # Protein checkbox
                row.cells[2].text = "\u2610"
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Set column widths
                set_cell_width(row.cells[0], 1.1)
                set_cell_width(row.cells[1], 3.8)
                set_cell_width(row.cells[2], 0.7)
                set_cell_width(row.cells[3], 0.7)

            # Water + sleep row
            water_p = doc.add_paragraph()
            r = water_p.add_run("Water (oz): ________")
            r.font.size = Pt(10)
            r.bold = True
            water_p.add_run("     ")
            r2 = water_p.add_run("Sleep (hrs): ________")
            r2.font.size = Pt(10)
            r2.bold = True
            water_p.add_run("     ")
            r3 = water_p.add_run("Energy Level (1-5): ________")
            r3.font.size = Pt(10)
            r3.bold = True

            # Only add a thin separator, not a page break for every day
            if day < 7:
                add_horizontal_line(doc)
            # But if it's day 3 or 6 within the week, consider page break for print layout
            if day in [3, 6]:
                doc.add_page_break()

        # Page break between weeks
        if week < 4:
            doc.add_page_break()

    doc.add_page_break()

    # ================================================================
    # SECTION 9: ACCOUNTABILITY NOTES
    # ================================================================
    heading1("Weekly Check-In Notes")
    body(
        "At each weekly weigh-in, your position coach will review your log and "
        "write notes here. This is how we help you stay on track."
    )

    for week in range(1, 5):
        heading3(f"Week {week} Coach Notes")
        notes_table = doc.add_table(rows=3, cols=2)
        notes_table.style = "Table Grid"

        labels = [
            "What went well:",
            "What to improve:",
            "Adjustments for next week:",
        ]
        for i, label in enumerate(labels):
            row = notes_table.rows[i]
            set_row_height(row, 0.5)
            row.cells[0].text = label
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
            set_cell_width(row.cells[0], 2.0)
            set_cell_width(row.cells[1], 4.5)
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)

        body("")

    doc.add_page_break()

    # ================================================================
    # SECTION 10: COMMITMENT PAGE
    # ================================================================
    heading1("My Commitment")

    body("")
    body(
        "I commit to fueling my body the right way to become the best player and "
        "teammate I can be. I understand that what I eat directly affects how I "
        "perform, recover, and develop. I will hold myself accountable by tracking "
        "my meals honestly and consistently."
    )

    body("")
    body("")

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"

    sig_labels = [
        ("Player Signature:", "Date:"),
        ("Coach Signature:", "Date:"),
    ]
    for i, (l1, l2) in enumerate(sig_labels):
        row = sig_table.rows[i]
        set_row_height(row, 0.6)
        row.cells[0].text = l1
        row.cells[1].text = l2
        set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

    body("")
    body("")

    # Closing quote
    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = close.add_run('"Champions are built in the kitchen, not just the weight room."')
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    # ================================================================
    # SAVE
    # ================================================================
    out = "River_Valley_Vikings_Meal_Log.docx"
    doc.save(out)
    print(f"Saved  {out}")


if __name__ == "__main__":
    main()
