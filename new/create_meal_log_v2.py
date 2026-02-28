#!/usr/bin/env python3
"""Generate V2 meal-log: offseason lifting recovery & fueling guide."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- Colors (Columbia Blue / Gold / White / Navy) --
NAVY = RGBColor(0x00, 0x2D, 0x62)
COLUMBIA_BLUE = RGBColor(0x6C, 0xAC, 0xE4)
GOLD = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_BLUE_HEX = "D6EAF8"
LIGHT_GOLD_HEX = "FFF3CC"
LIGHT_GRAY_HEX = "F2F2F2"
LIGHT_GREEN_HEX = "D5F5E3"
NAVY_HEX = "002D62"
GOLD_HEX = "CFA700"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_row_height(row, inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_header_row(table, headers, font_size=10):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        set_cell_shading(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = WHITE
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.name = "Calibri"


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # -- Helpers --
    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def bold_body(bold_text, normal_text=""):
        p = doc.add_paragraph()
        r = p.add_run(bold_text)
        r.bold = True
        if normal_text:
            p.add_run(normal_text)
        return p

    def small_note(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        r.italic = True
        return p

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("OFFSEASON FUEL")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GOLD

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Recovery Nutrition & Meal Log")
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run(
        '"You don\'t get stronger in the weight room. '
        'You get stronger when you recover from the weight room."'
    )
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    doc.add_page_break()

    # ================================================================
    # SECTION 1: WHY RECOVERY NUTRITION MATTERS
    # ================================================================
    heading1("Why Recovery Nutrition Matters")

    body(
        "Lifting breaks your muscles down. That's the point. But the gains you're "
        "chasing — the strength, the size, the speed — those don't happen under the "
        "bar. They happen in the hours and days after, when your body rebuilds what "
        "you just tore up. The only way your body can rebuild is if you give it the "
        "raw materials: protein to repair muscle fibers, carbs to refill your energy "
        "stores, and enough total food to fuel the process."
    )

    body(
        "If you lift hard but eat like garbage — or worse, skip meals — you are "
        "wasting your time in the weight room. You will not get stronger. You will "
        "stay sore longer. You will plateau. The guys who make the biggest jumps "
        "between February and August are the guys who eat to recover."
    )

    heading3("The Science (Simple Version)")

    # Recovery science table
    sci_table = doc.add_table(rows=4, cols=2)
    sci_table.style = "Table Grid"
    make_header_row(sci_table, ["What Happens", "What Your Body Needs"])

    sci_data = [
        (
            "Muscle fibers tear during lifting\n(micro-damage = good)",
            "PROTEIN to rebuild fibers thicker and stronger.\n"
            "Goal: protein at every meal, especially post-workout.",
        ),
        (
            "Glycogen (stored energy) gets depleted\nduring sets and reps",
            "CARBS to refill glycogen stores so you have\n"
            "energy for your next session and don't feel flat.",
        ),
        (
            "Inflammation and stress hormones rise\nafter hard training",
            "SLEEP + WATER + WHOLE FOODS to bring\n"
            "inflammation down and let adaptation happen.",
        ),
    ]
    for i, (what, need) in enumerate(sci_data):
        row = sci_table.rows[i + 1]
        row.cells[0].text = what
        row.cells[1].text = need
        set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    body("")
    bold_body("Bottom line: ", "Lifting is the stimulus. Food is the builder. Sleep is the construction crew. You need all three.")

    doc.add_page_break()

    # ================================================================
    # SECTION 2: PLAYER INFO
    # ================================================================
    heading1("Player Information")

    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    info_labels = [
        ("Name:", "", "Position:", ""),
        ("Grade:", "", "Height:", ""),
        ("Current Weight:", "", "Current Maxes:", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_labels):
        row = info_table.rows[i]
        set_row_height(row, 0.35)
        c0, c1, c2, c3 = row.cells
        set_cell_shading(c0, LIGHT_BLUE_HEX)
        set_cell_shading(c2, LIGHT_BLUE_HEX)
        for cell in [c0, c2]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        c0.text = l1
        c1.text = v1
        c2.text = l2
        c3.text = v2

    body("")

    # Lift maxes tracker
    heading3("Starting Lift Numbers")
    small_note("Record your current maxes so you can track progress alongside your nutrition.")

    lift_table = doc.add_table(rows=2, cols=5)
    lift_table.style = "Table Grid"
    make_header_row(lift_table, ["Squat", "Bench", "Deadlift", "Power Clean", "Bodyweight"])
    set_row_height(lift_table.rows[1], 0.4)

    doc.add_page_break()

    # ================================================================
    # SECTION 3: THE RECOVERY PLATE
    # ================================================================
    heading1("The Recovery Plate")

    body(
        "Every time you eat, you are either helping or hurting your recovery. Build "
        "your plate around these four pillars. You don't need to count calories or "
        "weigh food. Just make sure every meal has protein and carbs — those are "
        "the two non-negotiables for getting stronger."
    )

    heading3("What Counts as a Serving? (Use Your Hands)")
    bullet("Palm = 1 serving of protein (chicken, beef, fish, eggs)")
    bullet("Fist = 1 serving of carbs (rice, pasta, potatoes, oatmeal)")
    bullet("Cupped hand = 1 serving of fruit or snack")
    bullet("Thumb = 1 serving of fats (butter, oil, peanut butter)")
    bullet("Two fists = 1 serving of vegetables (eat as much as you want)")

    body("")

    # Recovery plate table
    plate_table = doc.add_table(rows=5, cols=3)
    plate_table.style = "Table Grid"
    make_header_row(plate_table, ["Pillar", "Why It Matters for Recovery", "Best Sources"])

    plate_data = [
        [
            "PROTEIN\n(Rebuild)",
            "Repairs torn muscle fibers.\nWithout it, you broke down\nmuscle for nothing.",
            "Chicken, beef, turkey, eggs,\nGreek yogurt, milk, tuna,\nprotein shake",
        ],
        [
            "CARBS\n(Reload)",
            "Refills glycogen so you're not\nrunning on empty next session.\nPrevents muscle breakdown.",
            "Rice, potatoes, oatmeal,\npasta, bread, bananas,\ntortillas, fruit",
        ],
        [
            "VEGETABLES\n(Recover)",
            "Anti-inflammatory. Vitamins &\nminerals that speed healing\nand reduce soreness.",
            "Broccoli, spinach, peppers,\ncarrots, green beans, salad,\nsweet potatoes",
        ],
        [
            "FATS\n(Regulate)",
            "Supports hormone production\n(testosterone, growth hormone).\nKeeps joints healthy.",
            "Peanut butter, avocado, nuts,\ncheese, olive oil, eggs\n(yolks included)",
        ],
    ]
    for i, row_data in enumerate(plate_data):
        row = plate_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, LIGHT_GOLD_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    body("")
    bold_body(
        "Minimum per meal: ",
        "1 palm of protein + 1 fist of carbs. That is the floor. "
        "Add vegetables and fats to build a complete plate.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 4: DAILY RECOVERY TARGETS
    # ================================================================
    heading1("Daily Recovery Targets")

    body(
        "These numbers are simple guidelines — not a strict diet. The goal is to "
        "give your body enough fuel to recover and adapt from offseason lifting. "
        "If you are consistently sore, tired, or stalling on your lifts, you are "
        "probably not eating enough."
    )

    target_table = doc.add_table(rows=4, cols=3)
    target_table.style = "Table Grid"
    make_header_row(target_table, ["Nutrient", "Daily Target", "Why"])

    target_data = [
        [
            "Protein",
            "1g per pound of bodyweight\n(180 lb = 180g protein)",
            "Rebuilds muscle. This is the #1\npriority. Spread across all meals.",
        ],
        [
            "Carbs",
            "2-3g per pound of bodyweight\n(180 lb = 360-540g carbs)",
            "Refuels glycogen. More on heavy\nlift days. Don't fear carbs.",
        ],
        [
            "Water",
            "Half your bodyweight in oz\n(180 lb = 90 oz minimum)",
            "Dehydration = weaker lifts,\nslower recovery, more cramps.",
        ],
    ]
    for i, row_data in enumerate(target_data):
        row = target_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    body("")

    heading3("Protein Cheat Sheet")
    body("How much protein is in common foods:")

    protein_table = doc.add_table(rows=11, cols=2)
    protein_table.style = "Table Grid"
    make_header_row(protein_table, ["Food", "Protein (approx.)"])

    protein_data = [
        ("Chicken breast (1 palm / 4 oz)", "~30g"),
        ("Ground beef (1 palm / 4 oz)", "~28g"),
        ("3 whole eggs", "~18g"),
        ("Greek yogurt (1 cup)", "~15-20g"),
        ("Glass of whole milk (8 oz)", "~8g"),
        ("Scoop of protein powder", "~25g"),
        ("Can of tuna", "~20g"),
        ("PB&J sandwich", "~12g"),
        ("String cheese (1 stick)", "~7g"),
        ("Chocolate milk (8 oz)", "~8g"),
    ]
    for i, (food, protein) in enumerate(protein_data):
        row = protein_table.rows[i + 1]
        row.cells[0].text = food
        row.cells[1].text = protein
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)
            set_cell_shading(row.cells[1], LIGHT_GRAY_HEX)

    small_note(
        "A 180 lb player needs ~180g protein/day. That's roughly 6 palms of protein "
        "spread across your meals. It's not hard if you eat protein at every meal."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 5: MEAL TIMING AROUND LIFTING
    # ================================================================
    heading1("Meal Timing Around Lifting")

    body(
        "When you eat matters almost as much as what you eat — especially around "
        "your training window. The two most important meals of your day are the "
        "one before you lift and the one right after."
    )

    heading2("Lift-Day Fueling Schedule")

    timing_table = doc.add_table(rows=7, cols=3)
    timing_table.style = "Table Grid"
    make_header_row(timing_table, ["Meal", "When", "What & Why"])

    timing_data = [
        [
            "Breakfast",
            "Within 30 min\nof waking up",
            "Break the overnight fast. Protein + carbs.\n"
            "Your body has been fasting 8+ hrs — feed it.\n"
            "Ex: eggs + toast + banana, oatmeal + milk",
        ],
        [
            "Mid-Morning",
            "Between\nbreakfast\nand lunch",
            "Keep fuel coming. Protein + carb snack.\n"
            "Ex: PB&J, Greek yogurt + granola, trail mix,\n"
            "string cheese + crackers",
        ],
        [
            "Lunch",
            "Noon",
            "Full plate: protein + carbs + veggies. This is\n"
            "a real meal, not a bag of chips.\n"
            "Ex: chicken + rice + green beans",
        ],
        [
            "PRE-LIFT",
            "60-90 min\nbefore lifting",
            "Carbs + moderate protein. Easy to digest.\n"
            "This is your fuel for the session.\n"
            "Ex: banana + PB, granola bar + milk, rice + chicken",
        ],
        [
            "POST-LIFT",
            "Within 45 min\nafter lifting",
            "MOST IMPORTANT MEAL. Protein + fast carbs.\n"
            "This is the recovery window — don't miss it.\n"
            "Ex: protein shake + banana, chocolate milk + PB&J",
        ],
        [
            "Dinner",
            "Evening",
            "Full plate: protein + carbs + veggies + fats.\n"
            "Your body does most of its repair at night.\n"
            "Ex: beef + pasta + salad + bread + milk",
        ],
    ]
    for i, row_data in enumerate(timing_data):
        row = timing_table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE_HEX)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(10)
        # Highlight pre/post lift rows
        if i == 3:  # pre-lift
            set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        if i == 4:  # post-lift
            set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)

    body("")
    heading3("Non-Lift Days")
    body(
        "You still need to eat. Recovery doesn't stop because you're not in the "
        "weight room. Your body is still rebuilding from the previous session. "
        "Eat the same meals — just skip the pre/post-lift timing. Don't use rest "
        "days as an excuse to eat less."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 6: POST-WORKOUT WINDOW
    # ================================================================
    heading1("The Post-Workout Window")

    body(
        "This is the one meal you cannot afford to skip. In the 30-60 minutes "
        "after lifting, your muscles are primed to absorb protein and carbs at "
        "a higher rate than any other time of day. Miss this window and you are "
        "leaving gains on the table."
    )

    heading3("What to Eat Post-Lift")
    bullet("20-40g protein (shake, chocolate milk, chicken, eggs)")
    bullet("30-60g fast carbs (banana, bread, rice, fruit, Gatorade)")
    bullet("This is NOT the time for a salad — your body needs fuel, not fiber")

    body("")
    heading3("Quick Post-Lift Combos (No Cooking Required)")

    combo_table = doc.add_table(rows=7, cols=2)
    combo_table.style = "Table Grid"
    make_header_row(combo_table, ["Option", "What You Get"])

    combos = [
        ("Protein shake + banana", "~30g protein, ~30g carbs"),
        ("Chocolate milk (16 oz) + PB&J", "~16g protein, ~60g carbs"),
        ("Greek yogurt + granola + berries", "~20g protein, ~40g carbs"),
        ("2 string cheese + apple + granola bar", "~15g protein, ~45g carbs"),
        ("Protein bar + Gatorade", "~20g protein, ~50g carbs"),
        ("Deli turkey sandwich + milk", "~25g protein, ~40g carbs"),
    ]
    for i, (option, gets) in enumerate(combos):
        row = combo_table.rows[i + 1]
        row.cells[0].text = option
        row.cells[1].text = gets
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)
            set_cell_shading(row.cells[1], LIGHT_GRAY_HEX)

    small_note(
        "Chocolate milk is one of the best post-workout drinks — cheap, portable, "
        "and has the right protein-to-carb ratio. Keep one in your bag."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 7: RECOVERY KILLERS
    # ================================================================
    heading1("Recovery Killers")

    body(
        "These are the habits that undo your work in the weight room. You put in "
        "the reps — don't waste them."
    )

    killer_table = doc.add_table(rows=9, cols=2)
    killer_table.style = "Table Grid"
    make_header_row(killer_table, ["Recovery Killer", "What It Does to You"])

    killers = [
        (
            "Skipping breakfast",
            "You've fasted 8-10 hours. Not eating = your body burns muscle for energy instead of building it.",
        ),
        (
            "Skipping post-workout meal",
            "Biggest missed opportunity. Your muscles are starving for protein and carbs. "
            "Feed them or lose the session.",
        ),
        (
            "Eating only 1-2 meals/day",
            "Not enough fuel to recover. You'll be sore longer, weaker in the next "
            "session, and plateau fast.",
        ),
        (
            "Soda and energy drinks\nas main beverages",
            "Sugar crashes, dehydration, zero recovery benefit. Water and milk are "
            "the only drinks that help.",
        ),
        (
            "Sleeping less than 7 hours",
            "Growth hormone peaks during deep sleep. Less sleep = less recovery = "
            "less strength gains. Period.",
        ),
        (
            "No protein at meals",
            "A plate of pasta with no chicken is just energy — no building blocks. "
            "Every meal needs protein.",
        ),
        (
            "Fast food as your only\nfood source",
            "Low protein, high sodium, inflammatory fats. Occasional is fine — "
            "daily habit is a problem.",
        ),
        (
            "Dehydration",
            "Muscle is 75% water. Dehydrated = weaker contractions, more cramps, "
            "slower nutrient delivery.",
        ),
    ]
    for i, (killer, effect) in enumerate(killers):
        row = killer_table.rows[i + 1]
        row.cells[0].text = killer
        row.cells[1].text = effect
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GRAY_HEX)
            set_cell_shading(row.cells[1], LIGHT_GRAY_HEX)

    doc.add_page_break()

    # ================================================================
    # SECTION 8: SAMPLE RECOVERY DAY
    # ================================================================
    heading1("Sample Lift-Day Meals")

    body(
        "Here's what a full day of eating looks like when your only goal is "
        "recovering from offseason lifting and getting stronger. This is not a "
        "diet — it's fueling."
    )

    sample_table = doc.add_table(rows=7, cols=3)
    sample_table.style = "Table Grid"
    make_header_row(sample_table, ["Meal", "Example", "Protein"])

    sample_data = [
        ("Breakfast", "4 eggs scrambled + 2 toast + banana\n+ glass of milk", "~35g"),
        ("Mid-Morning", "PB&J sandwich + string cheese", "~19g"),
        ("Lunch", "Chicken breast + rice + green beans\n+ water", "~35g"),
        ("Pre-Lift", "Granola bar + banana + water", "~5g"),
        ("Post-Lift", "Protein shake (milk) + banana\n+ handful of crackers", "~35g"),
        ("Dinner", "Ground beef + pasta + salad\n+ bread + glass of milk", "~40g"),
    ]
    for i, (meal, example, protein) in enumerate(sample_data):
        row = sample_table.rows[i + 1]
        row.cells[0].text = meal
        row.cells[1].text = example
        row.cells[2].text = protein
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        for p in row.cells[2].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.bold = True
        # Highlight post-lift
        if i == 4:
            set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)

    body("")
    p = doc.add_paragraph()
    r = p.add_run("Daily total: ~169g protein")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        " — close to target for a 170 lb player. Add a glass of milk or extra "
        "eggs at any meal to bump it higher."
    )

    body("")
    heading3("Budget-Friendly Recovery Foods")
    body(
        "You don't need expensive supplements. These staples are cheap and do the job:"
    )
    bullet("Eggs", bold_prefix="~$3/dozen: ")
    bullet("best protein per dollar", bold_prefix="")
    bullet("Rice", bold_prefix="~$4 for 5 lbs: ")
    bullet("carbs that last all week", bold_prefix="")
    bullet("Peanut butter", bold_prefix="~$3/jar: ")
    bullet("protein + healthy fats + calories", bold_prefix="")
    bullet("Oatmeal", bold_prefix="~$4/canister: ")
    bullet("breakfast every day for weeks", bold_prefix="")
    bullet("Whole milk", bold_prefix="~$4/gallon: ")
    bullet("protein + carbs + calories + cheap", bold_prefix="")
    bullet("Frozen chicken", bold_prefix="~$8 for 3 lbs: ")
    bullet("bulk protein, easy to cook", bold_prefix="")
    bullet("Bananas", bold_prefix="~$0.25 each: ")
    bullet("perfect pre/post workout carb", bold_prefix="")
    bullet("Chocolate milk", bold_prefix="~$1/carton: ")
    bullet("the best cheap post-workout drink", bold_prefix="")
    bullet("Canned tuna", bold_prefix="~$1/can: ")
    bullet("20g protein, no cooking", bold_prefix="")
    bullet("Beans/lentils", bold_prefix="~$1/can: ")
    bullet("protein + carbs + fiber", bold_prefix="")

    doc.add_page_break()

    # ================================================================
    # SECTION 9: SLEEP & RECOVERY CHECKLIST
    # ================================================================
    heading1("The Recovery Checklist")

    body(
        "Nutrition is the biggest piece, but recovery is a system. Check these "
        "boxes every day and you will get stronger."
    )

    check_table = doc.add_table(rows=8, cols=2)
    check_table.style = "Table Grid"
    make_header_row(check_table, ["Daily Checklist", "Target"])

    checks = [
        ("\u2610  Ate breakfast with protein", "Within 30 min of waking"),
        ("\u2610  Ate protein at every meal", "1 palm minimum per meal"),
        ("\u2610  Ate pre-lift meal", "60-90 min before training"),
        ("\u2610  Ate post-lift meal", "Within 45 min after training"),
        ("\u2610  Drank enough water", "Half bodyweight in ounces"),
        ("\u2610  Ate 4+ meals today", "Don't skip — feed the recovery"),
        ("\u2610  Got 7+ hours of sleep", "Growth hormone peaks in deep sleep"),
    ]
    for i, (check, target) in enumerate(checks):
        row = check_table.rows[i + 1]
        row.cells[0].text = check
        row.cells[1].text = target
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(row.cells[0], LIGHT_GREEN_HEX)
            set_cell_shading(row.cells[1], LIGHT_GREEN_HEX)

    body("")
    heading3("Sleep Is Not Optional")
    body(
        "Your body releases the most growth hormone during deep sleep (stages 3-4). "
        "If you are sleeping 5-6 hours, you are cutting your recovery in half. "
        "Getting 8 hours of sleep is the single easiest thing you can do to "
        "get stronger — and it's free."
    )
    bullet("Set a consistent bedtime. Your body adapts to routine.")
    bullet("Put the phone down 30 minutes before bed. Blue light delays sleep.")
    bullet("No caffeine after 3 PM (pre-workout drinks count).")
    bullet("Dark, cool room. That's the formula.")

    doc.add_page_break()

    # ================================================================
    # SECTION 10: WEEKLY WEIGH-IN & LIFT TRACKER
    # ================================================================
    heading1("Weekly Progress Tracker")
    body(
        "Weigh yourself at the same time each week (morning, before eating). "
        "Track your lifts alongside your weight — the goal is to see both go up."
    )

    prog_table = doc.add_table(rows=13, cols=7)
    prog_table.style = "Table Grid"
    make_header_row(
        prog_table,
        ["Wk", "Date", "Weight", "Squat", "Bench", "Clean", "Notes"],
        font_size=9,
    )

    for i in range(1, 13):
        row = prog_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_row_height(row, 0.35)
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY_HEX)

    small_note(
        "If weight is going up and lifts are going up, you're recovering properly. "
        "If lifts stall while you're eating, check sleep and water first."
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 11: DAILY MEAL LOG (4 weeks)
    # ================================================================
    heading1("Daily Meal Log")
    body(
        "Log every meal. Be specific. \"Chicken and rice\" is useful. \"Food\" is not. "
        "Check the protein box if the meal had at least 1 palm of protein. Track "
        "your water, sleep, and soreness at the bottom."
    )

    add_horizontal_line(doc)

    for week in range(1, 5):
        heading2(f"Week {week}")

        for day in range(1, 8):
            day_names = [
                "",
                "Monday",
                "Tuesday",
                "Wednesday",
                "Thursday",
                "Friday",
                "Saturday",
                "Sunday",
            ]
            day_name = day_names[day]

            # Day header
            day_head = doc.add_paragraph()
            r = day_head.add_run(f"Day {(week - 1) * 7 + day}: {day_name}")
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = NAVY

            r2 = day_head.add_run("     Date: _______________")
            r2.font.size = Pt(10)
            r2.font.color.rgb = GRAY

            # Lift day indicator
            if day_name in ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]:
                r3 = day_head.add_run("     Lift Day?  \u2610 Yes  \u2610 No")
                r3.font.size = Pt(10)
                r3.font.color.rgb = DARK

            # Meal table
            meal_table = doc.add_table(rows=7, cols=4)
            meal_table.style = "Table Grid"
            make_header_row(
                meal_table,
                ["Meal", "Food / Drink", "Protein?", "Time"],
                font_size=9,
            )

            meal_labels = [
                "Breakfast",
                "Mid-Morning",
                "Lunch",
                "Pre-Lift",
                "Post-Lift",
                "Dinner",
            ]
            for i, label in enumerate(meal_labels):
                row = meal_table.rows[i + 1]
                row.cells[0].text = label
                bg = LIGHT_BLUE_HEX
                if label == "Pre-Lift":
                    bg = LIGHT_GOLD_HEX
                elif label == "Post-Lift":
                    bg = LIGHT_GREEN_HEX
                set_cell_shading(row.cells[0], bg)
                set_row_height(row, 0.4)
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.bold = True
                # Protein checkbox
                row.cells[2].text = "\u2610"
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Column widths
                set_cell_width(row.cells[0], 1.0)
                set_cell_width(row.cells[1], 3.8)
                set_cell_width(row.cells[2], 0.7)
                set_cell_width(row.cells[3], 0.7)

            # Bottom tracking row
            track_p = doc.add_paragraph()
            r = track_p.add_run("Water (oz): ________")
            r.font.size = Pt(10)
            r.bold = True
            track_p.add_run("   ")
            r2 = track_p.add_run("Sleep (hrs): ________")
            r2.font.size = Pt(10)
            r2.bold = True
            track_p.add_run("   ")
            r3 = track_p.add_run("Soreness (1-5): ________")
            r3.font.size = Pt(10)
            r3.bold = True
            track_p.add_run("   ")
            r4 = track_p.add_run("Energy (1-5): ________")
            r4.font.size = Pt(10)
            r4.bold = True

            if day < 7:
                add_horizontal_line(doc)
            if day in [3, 6]:
                doc.add_page_break()

        if week < 4:
            doc.add_page_break()

    doc.add_page_break()

    # ================================================================
    # SECTION 12: COACH CHECK-IN
    # ================================================================
    heading1("Weekly Coach Check-In")
    body(
        "Bring your log to your weekly check-in. Your coach will review your "
        "meals, track your progress, and help you adjust."
    )

    for week in range(1, 5):
        heading3(f"Week {week}")
        notes_table = doc.add_table(rows=4, cols=2)
        notes_table.style = "Table Grid"

        labels = [
            "Meals per day (avg):",
            "Post-lift meals logged:",
            "What to keep doing:",
            "What to improve:",
        ]
        for i, label in enumerate(labels):
            row = notes_table.rows[i]
            set_row_height(row, 0.45)
            row.cells[0].text = label
            set_cell_shading(row.cells[0], LIGHT_BLUE_HEX)
            set_cell_width(row.cells[0], 2.2)
            set_cell_width(row.cells[1], 4.3)
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)

        body("")

    doc.add_page_break()

    # ================================================================
    # SECTION 13: COMMITMENT
    # ================================================================
    heading1("My Commitment")

    body("")
    body(
        "I understand that my body gets stronger through recovery, not just "
        "lifting. I commit to fueling my recovery by eating enough protein, "
        "carbs, and whole foods every day — especially after I train. I will "
        "track my meals honestly, drink enough water, and prioritize sleep. "
        "I will not waste my work in the weight room by neglecting my nutrition."
    )

    body("")

    # Lift goals mini-table
    heading3("My Offseason Strength Goals")
    goal_table = doc.add_table(rows=2, cols=4)
    goal_table.style = "Table Grid"
    make_header_row(goal_table, ["Squat Goal", "Bench Goal", "Clean Goal", "Weight Goal"])
    set_row_height(goal_table.rows[1], 0.45)

    body("")
    body("")

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"

    sig_labels = [
        ("Player Signature:", "Date:"),
        ("Coach Signature:", "Date:"),
    ]
    for i, (l1, l2) in enumerate(sig_labels):
        row = sig_table.rows[i]
        set_row_height(row, 0.6)
        row.cells[0].text = l1
        row.cells[1].text = l2
        set_cell_shading(row.cells[0], LIGHT_GOLD_HEX)
        set_cell_shading(row.cells[1], LIGHT_GOLD_HEX)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

    body("")
    body("")

    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = close.add_run(
        '"The work is in the weight room. The results are in the kitchen and the bed."'
    )
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    # ================================================================
    # SAVE
    # ================================================================
    out = "River_Valley_Vikings_Meal_Log_V2_Recovery.docx"
    doc.save(out)
    print(f"Saved  {out}")


if __name__ == "__main__":
    main()
