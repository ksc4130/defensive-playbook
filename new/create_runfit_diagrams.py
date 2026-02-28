#!/usr/bin/env python3
"""
Generate run-fit diagrams for all fronts × formations × stunts.
Landscape .docx, one diagram per page, gap responsibility lines for all 11 defenders.
"""

import os, copy, datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# COLORS
# =============================================================================
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD_RGB = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_RGB = RGBColor(0x88, 0x88, 0x88)
NAVY_HEX = "002D62"

C_NAVY = "#002D62"
C_COLUMBIA = "#6CACE4"
C_GOLD = "#CFA700"
C_GRAY = "#888888"
C_OL = "#AAAAAA"
C_WR = "#666666"
C_WHITE = "#FFFFFF"
C_RED = "#CC3333"

# =============================================================================
# Y-DEPTHS
# =============================================================================
Y_LOS = 0
Y_DL = 1.2
Y_LB = 3.5
Y_OLB = 2.5
Y_APEX = 3.0
Y_CB = 5.5
Y_S = 8.0

# =============================================================================
# GAP TARGET X COORDINATES
# =============================================================================
GAP_A_FIELD = 1.0
GAP_A_BND = -1.0
GAP_B_FIELD = 3.0
GAP_B_BND = -3.0
GAP_C_FIELD = 5.5
GAP_C_BND = -5.5
CONTAIN_FIELD = 6.5
CONTAIN_BND = -6.5
FORCE_FIELD = 8.0
FORCE_BND = -8.0
ALLEY_FIELD = 9.5
ALLEY_BND = -9.5

# =============================================================================
# OUTPUT DIRECTORIES
# =============================================================================
DIAGRAM_DIR = "/tmp/rv_runfit_diagrams"
os.makedirs(DIAGRAM_DIR, exist_ok=True)
OUTPUT_DIR = "/home/ksc4130/src/defensive_playbook/diagram_docs/runfit"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# =============================================================================
# DRAWING PRIMITIVES
# =============================================================================

def new_fig(title="", figsize=(11, 7.5)):
    """Landscape-proportioned figure."""
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-16, 16)
    ax.set_ylim(-5, 11)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.axhline(y=0.5, color=C_GRAY, linewidth=1, linestyle="--", alpha=0.4)
    ax.text(14, 10, "FIELD \u2192", fontsize=8, color=C_GRAY,
            ha="right", va="top", style="italic")
    ax.text(-14, 10, "\u2190 BOUNDARY", fontsize=8, color=C_GRAY,
            ha="left", va="top", style="italic")
    if title:
        ax.set_title(title, fontsize=13, fontweight="bold",
                     color=C_NAVY, pad=10)
    return fig, ax


def draw_player(ax, x, y, label, color=C_NAVY, fontsize=8):
    ax.plot(x, y, "o", color=color, markersize=14,
            markeredgecolor=C_NAVY, markeredgewidth=1.2)
    ax.text(x, y, label, fontsize=fontsize, ha="center", va="center",
            color=C_WHITE, fontweight="bold")


def draw_gap_fit_line(ax, x1, y1, x2, y2, color=C_NAVY):
    """Dotted line, no arrow — gap responsibility."""
    ax.plot([x1, x2], [y1, y2], linestyle=":", color=color,
            linewidth=1.8, alpha=0.7)


def draw_stunt_arrow(ax, x1, y1, x2, y2, color=C_RED, lw=2.5):
    """Solid line with arrow — stunt/blitz movement."""
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=lw))


def note(ax, x, y, text, fontsize=7, color=C_NAVY):
    ax.text(x, y, text, fontsize=fontsize, ha="center",
            color=color, style="italic")


def bottom_note(ax, text):
    ax.text(0, -4.3, text, fontsize=8, ha="center",
            color=C_NAVY, fontweight="bold")


def save_fig(fig, name, subdir=DIAGRAM_DIR):
    path = os.path.join(subdir, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# =============================================================================
# OFFENSIVE FORMATION DATA & DRAWING
# =============================================================================
OL_BASE = [(0, 0), (2, 0), (-2, 0), (4, 0), (-4, 0)]


def draw_ol(ax, te_field=False, te_bnd=False):
    for x, y in OL_BASE:
        ax.add_patch(plt.Rectangle((x - 0.4, y - 0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
    if te_field:
        ax.add_patch(plt.Rectangle((6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
        ax.text(6, -1.0, "TE", fontsize=7, ha="center", color="#555")
    if te_bnd:
        ax.add_patch(plt.Rectangle((-6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
        ax.text(-6, -1.0, "TE", fontsize=7, ha="center", color="#555")


def draw_gap_labels(ax, te_field=False, te_bnd=False):
    """Draw gap letters between OL at LOS level."""
    if te_field and te_bnd:
        gaps = [
            (7.5, "D"), (5, "C"), (3, "B"), (1, "A"),
            (-1, "A"), (-3, "B"), (-5, "C"), (-7.5, "D")
        ]
    elif te_field:
        gaps = [
            (7.5, "D"), (5, "C"), (3, "B"), (1, "A"),
            (-1, "A"), (-3, "B"), (-5.5, "C")
        ]
    elif te_bnd:
        gaps = [
            (5.5, "C"), (3, "B"), (1, "A"),
            (-1, "A"), (-3, "B"), (-5, "C"), (-7.5, "D")
        ]
    else:
        gaps = [
            (5.5, "C"), (3, "B"), (1, "A"), (-1, "A"), (-3, "B"), (-5.5, "C")
        ]
    for gx, glabel in gaps:
        ax.text(gx, 0.0, glabel, fontsize=9, ha="center", va="center",
                color=C_NAVY, fontweight="bold", alpha=0.5)


def draw_qb(ax, x=0, y=-1.5):
    ax.plot(x, y, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(x, y - 0.7, "QB", fontsize=7, ha="center", color="#555")


def draw_rb(ax, x=0, y=-3, label="RB"):
    ax.plot(x, y, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(x, y - 0.7, label, fontsize=7, ha="center", color="#555")


def draw_wr(ax, x, y, label):
    ax.plot(x, y, "s", color=C_WR, markersize=7, markeredgecolor="#333")
    ax.text(x, y - 0.7, label, fontsize=6, ha="center", color="#555")


FORMATIONS = {
    "2x2": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"), (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [(0, -3, "RB")], "show_qb": True,
        "label": "2\u00d72 Spread (10 pers.)",
    },
    "3x1_field": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (7, 0, "#3"), (-13, 0, "#1")],
        "rbs": [(0, -3, "RB")], "show_qb": True,
        "label": "3\u00d71 Trips Field (10 pers.)",
    },
    "3x1_bnd": {
        "te_field": False, "te_bnd": False,
        "wrs": [(-13, 0, "#1"), (-9, 0, "#2"), (-7, 0, "#3"), (13, 0, "#1")],
        "rbs": [(0, -3, "RB")], "show_qb": True,
        "label": "3\u00d71 Trips Boundary (10 pers.)",
    },
    "2x2_te_field": {
        "te_field": True, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [(0, -3, "RB")], "show_qb": True,
        "label": "2\u00d72 + TE Field (11 pers.)",
    },
    "2x2_te_bnd": {
        "te_field": False, "te_bnd": True,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"), (-13, 0, "#1"), (-9, 0, "#2")],
        "rbs": [(0, -3, "RB")], "show_qb": True,
        "label": "2\u00d72 + TE Boundary (11 pers.)",
    },
    "empty_3x2": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (7, 0, "#3"),
                (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [], "show_qb": True,
        "label": "Empty 3\u00d72 (10 pers.)",
    },
    "empty_2x3": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"),
                (-13, 0, "#1"), (-9, 0, "#2"), (-7, 0, "#3")],
        "rbs": [], "show_qb": True,
        "label": "Empty 2\u00d73 (10 pers.)",
    },
    "21_pers": {
        "te_field": True, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (-13, 0, "#1")],
        "rbs": [(0, -3, "RB"), (-1.5, -2.5, "FB")], "show_qb": True,
        "label": "21 Personnel (I-Form, TE Field)",
    },
    "22_pers": {
        "te_field": True, "te_bnd": True,
        "wrs": [(13, 0, "#1"), (-13, 0, "#1")],
        "rbs": [(0, -3, "RB"), (-1.5, -2.5, "FB")], "show_qb": True,
        "label": "22 Personnel (Heavy, 2 TE)",
    },
}


def draw_formation(ax, form_key):
    f = FORMATIONS[form_key]
    draw_ol(ax, te_field=f["te_field"], te_bnd=f["te_bnd"])
    draw_gap_labels(ax, te_field=f["te_field"], te_bnd=f["te_bnd"])
    if f["show_qb"]:
        draw_qb(ax)
    for x, y, lbl in f.get("wrs", []):
        draw_wr(ax, x, y, lbl)
    for x, y, lbl in f.get("rbs", []):
        draw_rb(ax, x, y, lbl)


# =============================================================================
# DEFENSIVE FRONT DATA WITH GAP FITS
# =============================================================================
# Each player: (x, y, label, color, gap_target_x)

def _std_secondary():
    return [
        (12, Y_CB, "FC", C_COLUMBIA, ALLEY_FIELD),
        (-12, Y_CB, "BC", C_COLUMBIA, ALLEY_BND),
        (5, Y_S, "FS", C_GOLD, FORCE_FIELD),
        (-5, Y_S, "D", C_GOLD, FORCE_BND),
    ]


FRONTS = {
    "SHADE": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (2.5, Y_DL, "T", C_NAVY, GAP_B_FIELD),
            (-1.5, Y_DL, "N", C_NAVY, GAP_A_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_B_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "UNDER": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (1.5, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-2.5, Y_DL, "N", C_NAVY, GAP_B_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_B_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "EYES": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (1.5, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-1.5, Y_DL, "N", C_NAVY, GAP_A_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_B_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_B_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "WIDE": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (2.5, Y_DL, "T", C_NAVY, GAP_B_FIELD),
            (-2.5, Y_DL, "N", C_NAVY, GAP_B_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "DEUCES": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (2.0, Y_DL, "T", C_NAVY, GAP_B_FIELD),
            (-2.0, Y_DL, "N", C_NAVY, GAP_B_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "GRIZZLY": {
        "dl": [
            (3.5, Y_DL, "A", C_NAVY, GAP_B_FIELD),
            (1.5, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-1.5, Y_DL, "N", C_NAVY, GAP_A_BND),
            (-3.5, Y_DL, "E", C_NAVY, GAP_B_BND),
        ],
        "lbs": [
            (0, Y_LB, "M", C_COLUMBIA, GAP_C_FIELD),
            (-6.5, Y_OLB, "W", C_COLUMBIA, CONTAIN_BND),
            (6.5, Y_OLB, "B", C_COLUMBIA, CONTAIN_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "BOSS": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (2.5, Y_DL, "T", C_NAVY, GAP_B_FIELD),
            (0.5, Y_DL, "N", C_NAVY, GAP_A_FIELD),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (-0.5, Y_LB, "M", C_COLUMBIA, GAP_A_BND),
            (-2.5, Y_LB, "W", C_COLUMBIA, GAP_B_BND),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "BOSS UNDER": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (-0.5, Y_DL, "T", C_NAVY, GAP_A_BND),
            (-2.5, Y_DL, "N", C_NAVY, GAP_B_BND),
            (-4.7, Y_DL, "E", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (0.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (2.5, Y_LB, "W", C_COLUMBIA, GAP_B_FIELD),
            (7.5, Y_APEX, "B", C_COLUMBIA, FORCE_FIELD),
        ],
        "secondary": _std_secondary(),
    },
    "MINT": {
        "dl": [
            (3.5, Y_DL, "A", C_NAVY, GAP_B_FIELD),
            (0, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-3.5, Y_DL, "N", C_NAVY, GAP_B_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_OLB, "B", C_COLUMBIA, CONTAIN_FIELD),
            (-7.5, Y_OLB, "E", C_COLUMBIA, CONTAIN_BND),
        ],
        "secondary": _std_secondary(),
    },
    "ACE": {
        "dl": [
            (4.0, Y_DL, "A", C_NAVY, GAP_B_FIELD),
            (0, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-4.0, Y_DL, "N", C_NAVY, GAP_B_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_A_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_OLB, "B", C_COLUMBIA, CONTAIN_FIELD),
            (-7.5, Y_OLB, "E", C_COLUMBIA, CONTAIN_BND),
        ],
        "secondary": _std_secondary(),
    },
    "JET": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (0, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-4.7, Y_DL, "N", C_NAVY, CONTAIN_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_B_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_B_BND),
            (7.5, Y_OLB, "B", C_COLUMBIA, FORCE_FIELD),
            (-7.5, Y_OLB, "E", C_COLUMBIA, FORCE_BND),
        ],
        "secondary": _std_secondary(),
    },
    "SLIP": {
        "dl": [
            (4.7, Y_DL, "A", C_NAVY, CONTAIN_FIELD),
            (0, Y_DL, "T", C_NAVY, GAP_A_FIELD),
            (-3.5, Y_DL, "N", C_NAVY, GAP_B_BND),
        ],
        "lbs": [
            (1.5, Y_LB, "M", C_COLUMBIA, GAP_B_FIELD),
            (-1.5, Y_LB, "W", C_COLUMBIA, GAP_A_BND),
            (7.5, Y_OLB, "B", C_COLUMBIA, FORCE_FIELD),
            (-7.5, Y_OLB, "E", C_COLUMBIA, CONTAIN_BND),
        ],
        "secondary": _std_secondary(),
    },
}


def get_front(front_key, has_te):
    """Return front dict, adjusting B to LB depth field side if TE present."""
    front = copy.deepcopy(FRONTS[front_key])
    if has_te:
        new_lbs = []
        for item in front["lbs"]:
            if item[2] == "B":
                new_lbs.append((5.0, Y_LB, "B", C_COLUMBIA, item[4]))
            else:
                new_lbs.append(item)
        front["lbs"] = new_lbs
    return front


# =============================================================================
# STUNT DATA
# =============================================================================

STUNTS_4DOWN = {
    "SLANT": {
        "desc": "T+N slant to field",
        "moves": {"T": (1.5, -0.8), "N": (1.5, -0.8)},
        "gap_overrides": {},
        "legality": {
            "UNDER": "BEST", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK",
            "BOSS UNDER": "BEST",
        },
    },
    "ANGLE": {
        "desc": "T+N slant to boundary",
        "moves": {"T": (-1.5, -0.8), "N": (-1.5, -0.8)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "BEST", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "BEST",
        },
    },
    "PINCH": {
        "desc": "T+N pinch A gaps",
        "moves": {"T": (-1.0, -0.8), "N": (1.0, -0.8)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "BEST", "UNDER": "BEST", "EYES": "OK",
            "WIDE": "BEST", "DEUCES": "BEST", "GRIZZLY": "OK",
            "BOSS": "OK", "BOSS UNDER": "OK",
        },
    },
    "JACKS": {
        "desc": "T+N shoot B gaps",
        "moves": {"T": (1.5, -0.8), "N": (-1.5, -0.8)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "BEST", "UNDER": "BEST", "EYES": "BEST",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK",
            "BOSS UNDER": "OK",
        },
    },
    "SPLIT": {
        "desc": "N crosses C face to opposite A gap",
        "moves": {"N": (-1.5, -0.8)},
        "gap_overrides": {},
        "legality": {
            "GRIZZLY": "BEST", "BOSS": "BEST", "BOSS UNDER": "BEST",
        },
    },
    "CRASH": {
        "desc": "A+E\u2192B gaps. T+N\u2192A gaps.\nB=field contain. W=bnd contain. M=free LB.",
        "moves": {
            "A": (-1.4, -0.8), "E": (1.4, -0.8),
            "T": (-1.0, -0.8), "N": (1.0, -0.8),
        },
        "gap_overrides": {
            "B": CONTAIN_FIELD,
            "W": CONTAIN_BND,
            "M": GAP_A_FIELD,
        },
        "legality": {
            "SHADE": "BEST", "UNDER": "BEST", "EYES": "OK",
            "WIDE": "BEST", "DEUCES": "OK", "BOSS": "OK",
            "BOSS UNDER": "OK",
        },
    },
    "ANCHOR ATTACK": {
        "desc": "A washes OL one gap inside",
        "moves": {"A": (-1.5, -0.3)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "OK", "UNDER": "OK", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK", "BOSS UNDER": "OK",
        },
    },
    "EDGE ATTACK": {
        "desc": "E attacks OL one gap inside",
        "moves": {"E": (1.5, -0.3)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "OK", "UNDER": "OK", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK", "BOSS UNDER": "OK",
        },
    },
    "ANCHOR RAVEN": {
        "desc": "A shoots B gap from 5-tech",
        "moves": {"A": (-1.2, -0.8)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "OK", "UNDER": "OK", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK", "BOSS UNDER": "OK",
        },
    },
    "EDGE RAVEN": {
        "desc": "E shoots B gap from 5-tech",
        "moves": {"E": (1.2, -0.8)},
        "gap_overrides": {},
        "legality": {
            "SHADE": "OK", "UNDER": "OK", "EYES": "OK", "WIDE": "OK",
            "DEUCES": "OK", "GRIZZLY": "OK", "BOSS": "OK", "BOSS UNDER": "OK",
        },
    },
}

STUNTS_3DOWN = {
    "SLANT": {
        "desc": "A+T+N slant to field",
        "moves": {"A": (1.5, -0.8), "T": (1.5, -0.8), "N": (1.5, -0.8)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK"},
    },
    "ANGLE": {
        "desc": "A+T+N slant to boundary",
        "moves": {"A": (-1.5, -0.8), "T": (-1.5, -0.8), "N": (-1.5, -0.8)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK", "SLIP": "OK"},
    },
    "PINCH": {
        "desc": "A+N pinch B gaps",
        "moves": {"A": (-1.0, -0.8), "N": (1.0, -0.8)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK"},
    },
    "JACKS": {
        "desc": "A+N expand to C gaps",
        "moves": {"A": (1.5, -0.8), "N": (-1.5, -0.8)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK"},
    },
    "ANCHOR ATTACK": {
        "desc": "A washes OL inside",
        "moves": {"A": (-1.5, -0.3)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK"},
    },
    "ANCHOR RAVEN": {
        "desc": "A shoots B gap",
        "moves": {"A": (-1.2, -0.8)},
        "gap_overrides": {},
        "legality": {"MINT": "OK", "ACE": "OK", "JET": "OK"},
    },
}


# =============================================================================
# CORE RENDERING
# =============================================================================

def draw_defense_runfit(ax, front_dict, stunt=None):
    """Draw all defensive players with gap-fit dotted lines.
    If stunt provided, draw solid arrows for movers, override gap fits."""
    movers = stunt["moves"] if stunt else {}
    gap_overrides = stunt.get("gap_overrides", {}) if stunt else {}

    all_players = []
    for x, y, lbl, color, gap_x in front_dict["dl"]:
        all_players.append((x, y, lbl, color, gap_x, "dl"))
    for item in front_dict["lbs"]:
        x, y, lbl, color, gap_x = item
        all_players.append((x, y, lbl, color, gap_x, "lb"))
    for item in front_dict["secondary"]:
        x, y, lbl, color, gap_x = item
        all_players.append((x, y, lbl, color, gap_x, "sec"))

    for x, y, lbl, color, gap_x, group in all_players:
        draw_player(ax, x, y, lbl, color=color)

        if lbl in movers:
            dx, dy = movers[lbl]
            draw_stunt_arrow(ax, x, y, x + dx, y + dy, color=C_RED)
        else:
            final_gap = gap_overrides.get(lbl, gap_x)
            draw_gap_fit_line(ax, x, y, final_gap, Y_LOS, color=color)


def render_front_diagram(front_key, form_key, stunt_key=None, stunt_dict=None):
    """Render one complete diagram: front vs formation, optional stunt."""
    f = FORMATIONS[form_key]
    has_te = f["te_field"] or f["te_bnd"]
    front = get_front(front_key, has_te)

    if stunt_key:
        title = f"{stunt_key} on {front_key} vs {f['label']}"
    else:
        title = f"{front_key} vs {f['label']}"

    fig, ax = new_fig(title)
    draw_formation(ax, form_key)
    draw_defense_runfit(ax, front, stunt=stunt_dict)

    if stunt_dict:
        bottom_note(ax, stunt_dict["desc"])

    return fig


# =============================================================================
# LANDSCAPE DOCX BUILDER
# =============================================================================

def new_doc():
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    return doc


def doc_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_body(doc, text):
    return doc.add_paragraph(text)

def doc_img(doc, path, width=9.0):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def doc_title_page(doc, subtitle):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = y.add_run("RUN-FIT DIAGRAMS \u2014 2026 Season")
    r.font.size = Pt(16); r.font.color.rgb = NAVY
    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("CONFIDENTIAL \u2014 COACHING STAFF ONLY")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    doc.add_page_break()

def doc_footer(doc, diagram_count):
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  "
                   f"\u2014  {diagram_count} diagrams")
    r.font.size = Pt(9); r.font.color.rgb = GRAY_RGB


# =============================================================================
# PER-FRONT DOCUMENT GENERATOR
# =============================================================================

FRONT_DESCRIPTIONS = {
    "SHADE": "A=5, T=3, N=2i, E=5. M=open A (field), W=open B (boundary). Base front \u2014 set field.",
    "UNDER": "A=5, T=2i, N=3, E=5. M=open B (field), W=open A (boundary). Set boundary.",
    "EYES": "A=5, T=2i, N=2i, E=5. M=open B, W=open B. Balanced vs zone/duo.",
    "WIDE": "A=5, T=3, N=3, E=5. M=open A, W=open A. Force bounce; vs B-gap heavy.",
    "DEUCES": "A=5, T=2, N=2, E=5. M/W react to DT. Vs Wing-T / pullers.",
    "GRIZZLY": "A=4i, T=2i, N=2i, E=4i. B/W=OLBs. M=10-tech. Vs power/counter, red zone.",
    "BOSS": "A=5, T=3(F), N=1(F), E=5. Bigs field. M=A bnd, W=B bnd.",
    "BOSS UNDER": "A=5, T=1(B), N=3(B), E=5. Bigs boundary. M=A fld, W=B fld.",
    "MINT": "3-down: A=4i, T=0, N=4i. B/E=OLBs. M/W=A gaps.",
    "ACE": "3-down: A=4, T=0, N=4. All DL 2-gapping. B/E=OLBs, contain.",
    "JET": "3-down: A=5, T=0, N=5. A/N=contain. M=B fld, W=B bnd.",
    "SLIP": "3-down: A=5, T=0, N=4i. 1-gap penetrating. A/E=contain. B=coverage clean.",
}

ALL_FORMATIONS = [
    "2x2", "3x1_field", "3x1_bnd",
    "2x2_te_field", "2x2_te_bnd",
    "empty_3x2", "empty_2x3",
    "21_pers", "22_pers",
]

FOUR_DOWN = ["SHADE", "UNDER", "EYES", "WIDE", "DEUCES",
             "GRIZZLY", "BOSS", "BOSS UNDER"]
THREE_DOWN = ["MINT", "ACE", "JET", "SLIP"]


def _safe(name):
    return name.replace(" ", "_").replace("(", "").replace(")", "")


def generate_front_docx(front_key):
    """Generate one landscape docx for a front."""
    count = 0
    fdir = os.path.join(DIAGRAM_DIR, _safe(front_key))
    os.makedirs(fdir, exist_ok=True)

    is_3down = front_key in THREE_DOWN
    stunt_catalog = STUNTS_3DOWN if is_3down else STUNTS_4DOWN

    legal_stunts = []
    for sk, sdata in stunt_catalog.items():
        rating = sdata["legality"].get(front_key)
        if rating:
            legal_stunts.append((sk, sdata, rating))
    legal_stunts.sort(key=lambda x: (0 if x[2] == "BEST" else 1, x[0]))

    doc = new_doc()
    doc_title_page(doc, f"{front_key} FRONT")

    doc_h1(doc, f"{front_key}")
    doc_body(doc, FRONT_DESCRIPTIONS.get(front_key, ""))
    doc_h2(doc, "Base Alignment \u2014 Run Fits vs Formations")

    for fm in ALL_FORMATIONS:
        fig = render_front_diagram(front_key, fm)
        img_path = save_fig(fig, f"{_safe(front_key)}_{fm}", subdir=fdir)
        count += 1
        doc_img(doc, img_path)
        doc.add_page_break()

    if legal_stunts:
        doc_h2(doc, f"Stunts on {front_key} \u2014 Run Fits vs 2\u00d72")

        best = [(sk, sd, r) for sk, sd, r in legal_stunts if r == "BEST"]
        ok = [(sk, sd, r) for sk, sd, r in legal_stunts if r == "OK"]

        if best:
            doc_h3(doc, "Recommended Pairings")
            for sk, sdata, rating in best:
                fig = render_front_diagram(front_key, "2x2",
                                           stunt_key=sk, stunt_dict=sdata)
                img_path = save_fig(fig, f"{_safe(front_key)}_{_safe(sk)}", subdir=fdir)
                count += 1
                doc_img(doc, img_path)
                doc.add_page_break()

        if ok:
            doc_h3(doc, "Legal Pairings")
            for sk, sdata, rating in ok:
                fig = render_front_diagram(front_key, "2x2",
                                           stunt_key=sk, stunt_dict=sdata)
                img_path = save_fig(fig, f"{_safe(front_key)}_{_safe(sk)}", subdir=fdir)
                count += 1
                doc_img(doc, img_path)
                doc.add_page_break()

    doc_footer(doc, count)

    out = os.path.join(OUTPUT_DIR, f"{_safe(front_key)}.docx")
    doc.save(out)
    print(f"  {front_key}: {count} diagrams \u2192 {out}")
    return out, count


def generate_all():
    print("=" * 60)
    print("RIVER VALLEY VIKINGS \u2014 RUN-FIT DIAGRAMS")
    print("=" * 60)
    print(f"Output: {OUTPUT_DIR}\n")

    grand_total = 0
    all_fronts = FOUR_DOWN + THREE_DOWN

    for fk in all_fronts:
        _, c = generate_front_docx(fk)
        grand_total += c

    print(f"\n{'=' * 60}")
    print(f"DONE. {grand_total} total diagrams across {len(all_fronts)} files.")
    print(f"Output: {OUTPUT_DIR}/")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    generate_all()
