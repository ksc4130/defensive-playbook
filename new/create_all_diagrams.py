#!/usr/bin/env python3
"""
Generate comprehensive defensive playbook diagrams.

Creates diagrams for ALL fronts × offensive formations, ALL stunts (on their
best-fit front), ALL pressures/blitzes, and ALL coverages — assembled into
a single .docx organized by section.

Data-driven architecture: defensive alignments, stunt movements, pressure
assignments, and coverage rules are all defined as data.  A generic rendering
engine draws any combination.
"""

import os
import datetime
import itertools
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# COLORS
# =============================================================================
NAVY = RGBColor(0x00, 0x2D, 0x62)
GOLD_RGB = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_RGB = RGBColor(0x88, 0x88, 0x88)
NAVY_HEX = "002D62"

C_NAVY = "#002D62"
C_COLUMBIA = "#6CACE4"
C_GOLD = "#CFA700"
C_GRAY = "#888888"
C_OL = "#AAAAAA"
C_WR = "#666666"
C_WHITE = "#FFFFFF"
C_RED = "#CC3333"

DIAGRAM_DIR = "/tmp/rv_all_diagrams"
FRONT_DIR = os.path.join(DIAGRAM_DIR, "fronts")
STUNT_DIR = os.path.join(DIAGRAM_DIR, "stunts")
PRESS_DIR = os.path.join(DIAGRAM_DIR, "pressures")
COV_DIR = os.path.join(DIAGRAM_DIR, "coverages")
for _d in (DIAGRAM_DIR, FRONT_DIR, STUNT_DIR, PRESS_DIR, COV_DIR):
    os.makedirs(_d, exist_ok=True)

# =============================================================================
# Y-DEPTHS
# =============================================================================
Y_LOS = 0
Y_DL = 1.2
Y_LB = 3.5
Y_OLB = 2.5
Y_APEX = 3.0
Y_CB = 5.5
Y_S = 8.0

# =============================================================================
# DRAWING PRIMITIVES
# =============================================================================

def new_fig(title="", figsize=(10, 6)):
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-16, 16)
    ax.set_ylim(-5, 11)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.axhline(y=0.5, color=C_GRAY, linewidth=1, linestyle="--", alpha=0.4)
    ax.text(14, 10, "FIELD \u2192", fontsize=8, color=C_GRAY,
            ha="right", va="top", style="italic")
    ax.text(-14, 10, "\u2190 BOUNDARY", fontsize=8, color=C_GRAY,
            ha="left", va="top", style="italic")
    if title:
        ax.set_title(title, fontsize=12, fontweight="bold",
                     color=C_NAVY, pad=10)
    return fig, ax


def draw_player(ax, x, y, label, color=C_NAVY, fontsize=8):
    """Draw a defensive player (filled circle with label)."""
    ax.plot(x, y, "o", color=color, markersize=14,
            markeredgecolor=C_NAVY, markeredgewidth=1.2)
    ax.text(x, y, label, fontsize=fontsize, ha="center", va="center",
            color=C_WHITE, fontweight="bold")


def draw_ghost(ax, x, y, label, color=C_NAVY):
    """Dashed circle showing post-snap destination."""
    circle = plt.Circle((x, y), 0.5, fill=False, edgecolor=color,
                         linewidth=1.5, linestyle="--")
    ax.add_patch(circle)
    ax.text(x, y, label, fontsize=7, ha="center", va="center",
            color=color, fontweight="bold")


def draw_arrow(ax, x1, y1, x2, y2, color=C_NAVY, lw=2):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=lw))


def draw_curved_arrow(ax, x1, y1, x2, y2, color=C_NAVY):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=1.8,
                                connectionstyle="arc3,rad=0.3"))


def note(ax, x, y, text, fontsize=7, color=C_NAVY):
    ax.text(x, y, text, fontsize=fontsize, ha="center",
            color=color, style="italic")


def bottom_note(ax, text):
    ax.text(0, -4.3, text, fontsize=8, ha="center",
            color=C_NAVY, fontweight="bold")


def save_fig(fig, name, subdir=DIAGRAM_DIR):
    path = os.path.join(subdir, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# =============================================================================
# OFFENSIVE FORMATION DATA & DRAWING
# =============================================================================

# OL positions: (x, y)
OL_BASE = [(0, 0), (2, 0), (-2, 0), (4, 0), (-4, 0)]

def draw_ol(ax, te_field=False, te_bnd=False, te_both=False):
    """Draw OL boxes.  Optionally add TE(s)."""
    for x, y in OL_BASE:
        ax.add_patch(plt.Rectangle((x - 0.4, y - 0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
    if te_field or te_both:
        ax.add_patch(plt.Rectangle((6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
        ax.text(6, -1.0, "TE", fontsize=7, ha="center", color="#555")
    if te_bnd or te_both:
        ax.add_patch(plt.Rectangle((-6 - 0.4, -0.4), 0.8, 0.8,
                     fill=True, facecolor=C_OL, edgecolor="#555", lw=1.5))
        ax.text(-6, -1.0, "TE", fontsize=7, ha="center", color="#555")


def draw_qb(ax, x=0, y=-1.5):
    ax.plot(x, y, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(x, y - 0.7, "QB", fontsize=7, ha="center", color="#555")


def draw_rb(ax, x=0, y=-3, label="RB"):
    ax.plot(x, y, "s", color=C_WR, markersize=8, markeredgecolor="#333")
    ax.text(x, y - 0.7, label, fontsize=7, ha="center", color="#555")


def draw_wr(ax, x, y, label):
    ax.plot(x, y, "s", color=C_WR, markersize=7, markeredgecolor="#333")
    ax.text(x, y - 0.7, label, fontsize=6, ha="center", color="#555")


# ---------------------------------------------------------------------------
# Formation descriptors
# ---------------------------------------------------------------------------
# Each formation is a dict:
#   te_field, te_bnd: bool
#   wrs: list of (x, y, label)
#   rbs: list of (x, y, label) or empty
#   show_qb: bool
#   label: short description

FORMATIONS = {
    "2x2": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"), (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [(0, -3, "RB")],
        "show_qb": True,
        "label": "2\u00d72 Spread (10 pers.)",
    },
    "3x1_field": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (7, 0, "#3"),
                (-13, 0, "#1")],
        "rbs": [(0, -3, "RB")],
        "show_qb": True,
        "label": "3\u00d71 Trips Field (10 pers.)",
    },
    "3x1_bnd": {
        "te_field": False, "te_bnd": False,
        "wrs": [(-13, 0, "#1"), (-9, 0, "#2"), (-7, 0, "#3"),
                (13, 0, "#1")],
        "rbs": [(0, -3, "RB")],
        "show_qb": True,
        "label": "3\u00d71 Trips Boundary (10 pers.)",
    },
    "2x2_te_field": {
        "te_field": True, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [(0, -3, "RB")],
        "show_qb": True,
        "label": "2\u00d72 + TE Field (11 pers.)",
    },
    "2x2_te_bnd": {
        "te_field": False, "te_bnd": True,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"), (-13, 0, "#1"), (-9, 0, "#2")],
        "rbs": [(0, -3, "RB")],
        "show_qb": True,
        "label": "2\u00d72 + TE Boundary (11 pers.)",
    },
    "empty_3x2": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (9, 0, "#2"), (7, 0, "#3"),
                (-13, 0, "#1"), (-7, 0, "#2")],
        "rbs": [],
        "show_qb": True,
        "label": "Empty 3\u00d72 (10 pers.)",
    },
    "empty_2x3": {
        "te_field": False, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (7, 0, "#2"),
                (-13, 0, "#1"), (-9, 0, "#2"), (-7, 0, "#3")],
        "rbs": [],
        "show_qb": True,
        "label": "Empty 2\u00d73 (10 pers.)",
    },
    "21_pers": {
        "te_field": True, "te_bnd": False,
        "wrs": [(13, 0, "#1"), (-13, 0, "#1")],
        "rbs": [(0, -3, "RB"), (-1.5, -2.5, "FB")],
        "show_qb": True,
        "label": "21 Personnel (I-Form, TE Field)",
    },
    "22_pers": {
        "te_field": True, "te_bnd": True,
        "wrs": [(13, 0, "#1"), (-13, 0, "#1")],
        "rbs": [(0, -3, "RB"), (-1.5, -2.5, "FB")],
        "show_qb": True,
        "label": "22 Personnel (Heavy, 2 TE)",
    },
}

def draw_formation(ax, form_key):
    """Draw an offensive formation from FORMATIONS dict."""
    f = FORMATIONS[form_key]
    draw_ol(ax, te_field=f["te_field"], te_bnd=f["te_bnd"])
    if f["show_qb"]:
        draw_qb(ax)
    for x, y, lbl in f.get("wrs", []):
        draw_wr(ax, x, y, lbl)
    for x, y, lbl in f.get("rbs", []):
        draw_rb(ax, x, y, lbl)


# =============================================================================
# DEFENSIVE FRONT DATA
# =============================================================================
# Each front: dict with DL positions, LB positions, secondary positions.
# Positions are (x, y, label, color).
# For TE SET variants we compute adjustments dynamically.

def _std_secondary():
    """Standard secondary: FC, BC, FS, D, B at apex."""
    return [
        (12, Y_CB, "FC", C_COLUMBIA),
        (-12, Y_CB, "BC", C_COLUMBIA),
        (5, Y_S, "FS", C_GOLD),
        (-5, Y_S, "D", C_GOLD),
    ]

def _base_front(a_x, t_x, n_x, e_x, m_x=1.5, w_x=-1.5,
                b_x=7.5, b_y=Y_APEX, m_y=Y_LB, w_y=Y_LB,
                m_label="M", w_label="W"):
    """Build a standard 4-2-5 front dict."""
    return {
        "dl": [(a_x, Y_DL, "A"), (t_x, Y_DL, "T"),
               (n_x, Y_DL, "N"), (e_x, Y_DL, "E")],
        "lbs": [(m_x, m_y, m_label, C_COLUMBIA),
                (w_x, w_y, w_label, C_COLUMBIA),
                (b_x, b_y, "B", C_COLUMBIA)],
        "secondary": _std_secondary(),
    }

def _3down_front(a_x, t_x, n_x, b_x=7.5, e_x=-7.5,
                 m_x=1.5, w_x=-1.5):
    """Build a 3-down package front dict."""
    return {
        "dl": [(a_x, Y_DL, "A"), (t_x, Y_DL, "T"),
               (n_x, Y_DL, "N")],
        "lbs": [(m_x, Y_LB, "M", C_COLUMBIA),
                (w_x, Y_LB, "W", C_COLUMBIA),
                (b_x, Y_OLB, "B", C_COLUMBIA),
                (e_x, Y_OLB, "E", C_COLUMBIA)],
        "secondary": _std_secondary(),
    }


# 4-down fronts
FRONTS = {
    "SHADE":      _base_front(4.7, 2.5, -1.5, -4.7),
    "UNDER":      _base_front(4.7, 1.5, -2.5, -4.7),
    "EYES":       _base_front(4.7, 1.5, -1.5, -4.7),
    "WIDE":       _base_front(4.7, 2.5, -2.5, -4.7),
    "DEUCES":     _base_front(4.7, 2.0, -2.0, -4.7),
    "GRIZZLY": {
        "dl": [(3.5, Y_DL, "A"), (1.5, Y_DL, "T"),
               (-1.5, Y_DL, "N"), (-3.5, Y_DL, "E")],
        "lbs": [(5.5, Y_LB, "M", C_COLUMBIA),    # 10-tech at LB depth
                (-6.5, Y_OLB, "W", C_COLUMBIA),  # OLB bnd
                (6.5, Y_OLB, "B", C_COLUMBIA)],   # OLB field
        "secondary": _std_secondary(),
    },
    "BOSS":       _base_front(4.7, 2.5, 0.5, -4.7),
    "BOSS UNDER": _base_front(4.7, -0.5, -2.5, -4.7),
    # 3-down packages
    "MINT": _3down_front(3.5, 0, -3.5),
    "ACE":  _3down_front(4.0, 0, -4.0),
    "JET":  _3down_front(4.7, 0, -4.7),
    "SLIP": _3down_front(4.7, 0, -3.5),
}

# TE SET adjustments for 4-down fronts (not GRIZZLY)
# When TE is to field: A→7-tech(6.7), T→3-tech to TE, N→2i away, E stays.
#   B inserts at LB depth C gap field (~5.0)
# When TE is to boundary: E→7-tech(-6.7), N→3-tech to TE, T→2i away, A stays.
#   B inserts at LB depth C gap boundary (~-5.0)

def apply_te_set_field(front_key):
    """Return a modified front dict for TE to field."""
    base = FRONTS[front_key]
    if front_key == "GRIZZLY":
        return base  # GRIZZLY does not use TE SET
    dl = list(base["dl"])
    # A → 7-tech field (6.7)
    dl[0] = (6.7, Y_DL, "A")
    # T → 3-tech to TE side (field) = 2.5
    dl[1] = (2.5, Y_DL, "T")
    # N → 2i away from TE = -1.5
    dl[2] = (-1.5, Y_DL, "N")
    # E stays
    lbs = list(base["lbs"])
    # B inserts at LB depth over C gap field
    new_lbs = []
    for lb in lbs:
        if lb[2] == "B":
            new_lbs.append((5.0, Y_LB, "B", C_COLUMBIA))
        else:
            new_lbs.append(lb)
    return {"dl": dl, "lbs": new_lbs, "secondary": base["secondary"]}


def apply_te_set_bnd(front_key):
    """Return a modified front dict for TE to boundary."""
    base = FRONTS[front_key]
    if front_key == "GRIZZLY":
        return base
    dl = list(base["dl"])
    # E → 7-tech boundary (-6.7)
    dl[3] = (-6.7, Y_DL, "E")
    # N → 3-tech to TE side (boundary) = -2.5
    dl[2] = (-2.5, Y_DL, "N")
    # T → 2i away from TE = 1.5
    dl[1] = (1.5, Y_DL, "T")
    # A stays
    lbs = list(base["lbs"])
    new_lbs = []
    for lb in lbs:
        if lb[2] == "B":
            new_lbs.append((-5.0, Y_LB, "B", C_COLUMBIA))
        else:
            new_lbs.append(lb)
    return {"dl": dl, "lbs": new_lbs, "secondary": base["secondary"]}


def draw_defense(ax, front_dict):
    """Draw all defensive players from a front dict."""
    for x, y, lbl in front_dict["dl"]:
        draw_player(ax, x, y, lbl)
    for item in front_dict["lbs"]:
        x, y, lbl, color = item
        draw_player(ax, x, y, lbl, color=color)
    for item in front_dict["secondary"]:
        x, y, lbl, color = item
        draw_player(ax, x, y, lbl, color=color)


# =============================================================================
# STUNT DATA — arrows showing DL movement
# =============================================================================
# Each stunt: list of (player_label, dx, dy) offset from current position,
# plus any notes.  We also store "best_front" to draw it on.

STUNTS = {
    "SLANT": {
        "best_front": "UNDER",
        "desc": "T+N slant to field. Edges play normal.",
        "moves": {"T": (1.5, -0.8), "N": (1.5, -0.8)},
        "notes": "NOT with Shade.",
    },
    "ANGLE": {
        "best_front": "SHADE",
        "desc": "T+N slant to boundary. Edges play normal.",
        "moves": {"T": (-1.5, -0.8), "N": (-1.5, -0.8)},
        "notes": "NOT with Under/Boss Under.",
    },
    "PINCH": {
        "best_front": "SHADE",
        "desc": "T+N pinch A gaps. A and E normal.",
        "moves": {"T": (-1.8, -0.8), "N": (0.8, -0.8)},
        "notes": "Best: Shade, Under, Wide, Deuces.",
    },
    "JACKS": {
        "best_front": "EYES",
        "desc": "T+N shoot B gaps. A and E contain.",
        "moves": {"T": (2.0, -0.8), "N": (-1.5, -0.8)},
        "notes": "NOT with Wide.",
    },
    "SPLIT (BOSS)": {
        "best_front": "BOSS",
        "desc": "N (1-tech) crosses C face to opp A gap.",
        "moves": {"N": (-1.5, -0.8)},
        "notes": "Only legal with BOSS, BOSS UNDER, GRIZZLY.",
    },
    "CRASH": {
        "best_front": "WIDE",
        "desc": "A+E shoot B gaps. T+N shoot A gaps.\nB=field contain. W=bnd contain. M=free LB.",
        "moves": {"A": (-1.4, -0.8), "E": (1.4, -0.8),
                  "T": (-1.8, -0.8), "N": (1.8, -0.8)},
        "notes": "NOT with Grizzly. Best: Wide, Shade, Under.",
    },
    "ANCHOR ATTACK": {
        "best_front": "UNDER",
        "desc": "A washes OL one gap inside.",
        "moves": {"A": (-1.5, -0.3)},
        "notes": "NOT paired with Cobra.",
    },
    "EDGE ATTACK": {
        "best_front": "SHADE",
        "desc": "E attacks OL one gap inside.",
        "moves": {"E": (1.5, -0.3)},
        "notes": "CAN pair with Cobra.",
    },
    "ANCHOR RAVEN": {
        "best_front": "SHADE",
        "desc": "A shoots B gap from 5-tech.",
        "moves": {"A": (-1.2, -0.8)},
        "notes": "Not with Hammer or staB.",
    },
    "EDGE RAVEN": {
        "best_front": "SHADE",
        "desc": "E shoots B gap from 5-tech.",
        "moves": {"E": (1.2, -0.8)},
        "notes": "Signal: point boundary, flap flap.",
    },
}


# =============================================================================
# PRESSURE / BLITZ DATA
# =============================================================================
# Each pressure: who blitzes (arrows from their position toward LOS),
# plus notes.

PRESSURES = {
    "Mike": {
        "best_front": "SHADE",
        "desc": "M blitzes his gap.",
        "blitzers": ["M"],
        "notes": "",
    },
    "Will": {
        "best_front": "SHADE",
        "desc": "W blitzes his gap.",
        "blitzers": ["W"],
        "notes": "",
    },
    "Bandit": {
        "best_front": "SHADE",
        "desc": "B blitzes his gap (field edge).",
        "blitzers": ["B"],
        "notes": "",
    },
    "Dawg": {
        "best_front": "SHADE",
        "desc": "D blitzes his gap (boundary).",
        "blitzers": ["D"],
        "notes": "Game-plan special.",
    },
    "sWarM": {
        "best_front": "SHADE",
        "desc": "M + W blitz their gaps.",
        "blitzers": ["M", "W"],
        "notes": "NOT legal with Cover 1.",
    },
    "BooM": {
        "best_front": "SHADE",
        "desc": "B + M blitz.",
        "blitzers": ["B", "M"],
        "notes": "Cover 1: W takes RB solo.",
    },
    "BoW": {
        "best_front": "SHADE",
        "desc": "B + W blitz.",
        "blitzers": ["B", "W"],
        "notes": "Cover 1: M takes RB solo.",
    },
    "MaD": {
        "best_front": "SHADE",
        "desc": "M + D blitz.",
        "blitzers": ["M", "D"],
        "notes": "Game-plan special.",
    },
    "Eat": {
        "best_front": "SHADE",
        "desc": "M + W + B blitz (all underneath).",
        "blitzers": ["M", "W", "B"],
        "notes": "Cover 0 ONLY. RB unaccounted.",
    },
    "Hammer": {
        "best_front": "SHADE",
        "desc": "B edge blitz + Anchor Attack (A washes inside).",
        "blitzers": ["B"],
        "stunt_moves": {"A": (-1.5, -0.3)},
        "notes": "Packaged pressure (married to Anchor Attack).",
    },
    "Shave": {
        "best_front": "SHADE",
        "desc": "W edge blitz + Edge Attack (E attacks inside).",
        "blitzers": ["W"],
        "stunt_moves": {"E": (1.5, -0.3)},
        "notes": "Packaged pressure (married to Edge Attack).",
    },
    "staB": {
        "best_front": "SHADE",
        "desc": "B one gap inside EOL.",
        "blitzers": ["B"],
        "notes": "B to B/C gap (inside normal). Game-plan special.",
    },
    "Freebird (Oregon)": {
        "best_front": "SHADE",
        "desc": "FS blitzes (post safety rushes). Cover 1 only.",
        "blitzers": ["FS"],
        "notes": "No post safety \u2014 accept the risk.",
    },
    "Cobra (Zunnel)": {
        "best_front": "SHADE",
        "desc": "BC blitzes. D \u2192 #1 bnd. B \u2192 #2 bnd.\nFS \u2192 #2 field. M/W funnel RB.",
        "blitzers": ["BC"],
        "notes": "Cover 0 + Zunnel only. Hash-only call.",
    },
}

# =============================================================================
# COVERAGE DATA
# =============================================================================
# Each coverage defines post-snap positions/adjustments for secondary + LBs.

COVERAGES = {
    "NINJA_2x2": {
        "label": "NINJA vs 2\u00d72 (MOD/CLAMP)",
        "family": "Cover 7",
        "secondary_notes": {
            "FC": "man-match #1", "FS": "top-down #2",
            "B": "#2/#3 apex", "BC": "clamp #1", "D": "control #2",
            "M": "hook", "W": "hook",
        },
        "zones": [("MOD", 7, 10.2, C_COLUMBIA), ("CLAMP", -7, 10.2, C_GOLD)],
    },
    "NINJA_3x1": {
        "label": "NINJA vs 3\u00d71 (POACH)",
        "family": "Cover 7",
        "secondary_notes": {
            "FC": "man-match #1", "FS": "holds field",
            "B": "#2/#3 trips", "BC": "clamp #1",
            "D": "POACH\n(#3 vert?)", "M": "hook", "W": "hook",
        },
        "ghost_moves": [("D", -5, Y_S, -3, 7.5)],
        "zones": [("POACH (D)", -3, 10.2, C_GOLD)],
    },
    "NINJA_empty": {
        "label": "NINJA vs Empty 3\u00d72",
        "family": "Cover 7",
        "secondary_notes": {
            "FC": "#1 field", "FS": "#2 field",
            "B": "#3 field", "BC": "#1 bnd", "D": "#2 bnd",
            "M": "hook zone", "W": "hook zone",
        },
    },
    "OREGON": {
        "label": "OREGON (Cover 1, Post=FS)",
        "family": "Cover 1",
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "POST/MOF", "D": "man #2 bnd",
            "B": "man #2 field", "M": "RB funnel", "W": "RB funnel",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 0, 9.5),
            ("D", -5, Y_S, -7, 6),
            ("B", 7.5, Y_APEX, 7, 5),
        ],
    },
    "OKLAHOMA": {
        "label": "OKLAHOMA (Cover 1, Post=D)",
        "family": "Cover 1",
        "pre_swap": {"FS": (-5, Y_S), "D": (5, Y_S)},  # FS/D switch
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "D": "POST/MOF", "FS": "man #2 bnd",
            "B": "man #2 field", "M": "RB funnel", "W": "RB funnel",
        },
        "ghost_moves": [
            ("D", 5, Y_S, 0, 9.5),
            ("FS", -5, Y_S, -7, 6),
            ("B", 7.5, Y_APEX, 7, 5),
        ],
    },
    "OHIO": {
        "label": "OHIO (Cover 1, Post=B) \u2014 EXCEPTION",
        "family": "Cover 1",
        "pre_swap": {"B": (0, 8)},  # B to middle ~8 yds
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "B": "POST/MOF", "FS": "man #2 field",
            "D": "man #2 bnd", "M": "RB funnel", "W": "RB funnel",
        },
        "ghost_moves": [
            ("B", 0, 8, 0, 10),
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
    },
    "ZEUS": {
        "label": "ZEUS (Cover 0 \u2014 Run-First Pressure)",
        "family": "Cover 0",
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "man #2 field", "D": "man #2 bnd",
            "M": "run-first\nthen rush", "B": "run-first\nRB funnel",
            "W": "run-first\nRB funnel",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
        "rush_arrows": True,
    },
    "ZORRO": {
        "label": "ZORRO (Cover 0) \u2014 EXCEPTION",
        "family": "Cover 0",
        "pre_swap": {"B": (0, 8)},
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "man #2 field", "D": "man #2 bnd",
            "B": "has RB\n(~8 yds mid)", "M": "rush", "W": "rush",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
    },
    "ZUNNEL": {
        "label": "ZUNNEL (Cover 0)",
        "family": "Cover 0",
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "man #2 field", "D": "man #2 bnd",
            "M": "RB funnel", "W": "RB funnel",
            "B": "green-light\nrusher",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
    },
    "ZILL": {
        "label": "ZILL (Cover 0)",
        "family": "Cover 0",
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "man #2 field", "D": "man #2 bnd",
            "W": "has RB (man)", "M": "rush", "B": "rush",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
    },
    "ZIKE": {
        "label": "ZIKE (Cover 0)",
        "family": "Cover 0",
        "secondary_notes": {
            "FC": "man #1", "BC": "man #1",
            "FS": "man #2 field", "D": "man #2 bnd",
            "M": "has RB (man)", "W": "rush", "B": "rush",
        },
        "ghost_moves": [
            ("FS", 5, Y_S, 7, 6),
            ("D", -5, Y_S, -7, 6),
        ],
    },
    "VIKING": {
        "label": "VIKING (Cover 3 \u2014 Spot Drop)",
        "family": "Cover 3",
        "secondary_notes": {
            "FC": "deep 1/3 field", "BC": "deep 1/3 bnd",
            "FS": "deep 1/3 mid", "D": "seam-curl-flat",
            "B": "curl/flat", "M": "hook", "W": "hook",
        },
        "ghost_moves": [
            ("FC", 12, Y_CB, 11, 9),
            ("FS", 5, Y_S, 0, 9.5),
            ("BC", -12, Y_CB, -11, 9),
            ("D", -5, Y_S, -6, 5.5),
        ],
    },
}


# =============================================================================
# COMPOSITE RENDERING FUNCTIONS
# =============================================================================

def render_front_vs_formation(front_key, form_key, title=None):
    """Render a defensive front against an offensive formation."""
    f = FORMATIONS[form_key]
    if title is None:
        title = f"{front_key} vs {f['label']}"

    fig, ax = new_fig(title)
    draw_formation(ax, form_key)

    # Pick front dict — apply TE SET if needed
    is_te_field = f["te_field"]
    is_te_bnd = f["te_bnd"]
    is_4down = front_key not in ("MINT", "ACE", "JET", "SLIP")

    if is_4down and front_key != "GRIZZLY":
        if is_te_field and not is_te_bnd:
            front = apply_te_set_field(front_key)
            note(ax, 0, -4.3, f"TE SET applied (TE field). Base front: {front_key}.")
        elif is_te_bnd and not is_te_field:
            front = apply_te_set_bnd(front_key)
            note(ax, 0, -4.3, f"TE SET applied (TE boundary). Base front: {front_key}.")
        elif is_te_field and is_te_bnd:
            # Both TEs — use base front, TE SET sets to the first TE
            front = apply_te_set_field(front_key)
            note(ax, 0, -4.3, f"TE SET applied (primary TE field). Base front: {front_key}.")
        else:
            front = FRONTS[front_key]
    else:
        front = FRONTS[front_key]

    draw_defense(ax, front)
    return fig


# =============================================================================
# STUNT LEGALITY MATRIX  (Section 21 of playbook)
# =============================================================================
# "BEST" = recommended, "OK" = legal, None = NOT legal
# 3-down: MINT/ACE/JET use 3-down stunt variants; SLIP = no stunts.
STUNT_LEGALITY = {
    "SLANT":          {"SHADE": None,   "UNDER": "BEST", "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "BEST"},
    "ANGLE":          {"SHADE": "BEST", "UNDER": None,   "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "BEST", "BOSS UNDER": None},
    "PINCH":          {"SHADE": "BEST", "UNDER": "BEST", "EYES": "OK",   "WIDE": "BEST", "DEUCES": "BEST", "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "JACKS":          {"SHADE": "BEST", "UNDER": "BEST", "EYES": "BEST", "WIDE": None,   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "SPLIT (BOSS)":   {"SHADE": None,   "UNDER": None,   "EYES": None,   "WIDE": None,   "DEUCES": None,   "GRIZZLY": "BEST", "BOSS": "BEST", "BOSS UNDER": "BEST"},
    "CRASH":          {"SHADE": "BEST", "UNDER": "BEST", "EYES": "OK",   "WIDE": "BEST", "DEUCES": "OK",   "GRIZZLY": None,   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "ANCHOR ATTACK":  {"SHADE": "OK",   "UNDER": "OK",   "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "EDGE ATTACK":    {"SHADE": "OK",   "UNDER": "OK",   "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "ANCHOR RAVEN":   {"SHADE": "OK",   "UNDER": "OK",   "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
    "EDGE RAVEN":     {"SHADE": "OK",   "UNDER": "OK",   "EYES": "OK",   "WIDE": "OK",   "DEUCES": "OK",   "GRIZZLY": "OK",   "BOSS": "OK",   "BOSS UNDER": "OK"},
}


def get_legal_stunts(front_key):
    """Return list of (stunt_key, rating) legal on this front."""
    results = []
    for sk, legality in STUNT_LEGALITY.items():
        rating = legality.get(front_key)
        if rating is not None:
            results.append((sk, rating))
    # Sort: BEST first, then OK
    results.sort(key=lambda x: (0 if x[1] == "BEST" else 1, x[0]))
    return results


def render_stunt_on_front(stunt_key, front_key, rating="OK"):
    """Render a stunt on a SPECIFIC front (not just its default best_front)."""
    s = STUNTS[stunt_key]
    front = FRONTS[front_key]

    badge = "\u2605 BEST" if rating == "BEST" else "OK"
    title = f"{stunt_key} on {front_key}  [{badge}]"
    fig, ax = new_fig(title)

    draw_ol(ax)
    draw_qb(ax)

    # Draw DL + arrows for movers
    for x, y, lbl in front["dl"]:
        draw_player(ax, x, y, lbl)
        if lbl in s["moves"]:
            dx, dy = s["moves"][lbl]
            draw_arrow(ax, x, y, x + dx, y + dy, color=C_RED, lw=2.5)

    # Draw LBs
    for item in front["lbs"]:
        x, y, lbl, color = item
        draw_player(ax, x, y, lbl, color=color)

    bottom_note(ax, s["desc"])
    if s.get("notes"):
        note(ax, 0, -3.3, s["notes"], fontsize=7)

    return fig


def render_pressure(press_key):
    """Render a pressure/blitz diagram on its best front."""
    p = PRESSURES[press_key]
    front_key = p["best_front"]
    front = FRONTS[front_key]

    title = f"PRESSURE: {press_key}"
    fig, ax = new_fig(title)

    draw_formation(ax, "2x2")

    # Draw DL
    for x, y, lbl in front["dl"]:
        draw_player(ax, x, y, lbl)
        # If there are stunt_moves for packaged pressures
        if lbl in p.get("stunt_moves", {}):
            dx, dy = p["stunt_moves"][lbl]
            draw_arrow(ax, x, y, x + dx, y + dy, color=C_RED, lw=2.5)

    # Build position→coords map for blitzer arrows
    pos_map = {}
    for x, y, lbl in front["dl"]:
        pos_map[lbl] = (x, y)
    for item in front["lbs"]:
        x, y, lbl, color = item
        pos_map[lbl] = (x, y)
    for item in front["secondary"]:
        x, y, lbl, color = item
        pos_map[lbl] = (x, y)

    # Draw LBs + secondary, with blitz arrows
    for item in front["lbs"]:
        x, y, lbl, color = item
        draw_player(ax, x, y, lbl, color=color)
        if lbl in p["blitzers"]:
            # Arrow toward LOS
            ty = Y_DL - 0.5
            tx = x * 0.7  # angle toward center
            draw_arrow(ax, x, y, tx, ty, color=C_COLUMBIA, lw=2.5)

    for item in front["secondary"]:
        x, y, lbl, color = item
        draw_player(ax, x, y, lbl, color=color)
        if lbl in p["blitzers"]:
            ty = Y_DL - 0.5
            tx = x * 0.6
            draw_arrow(ax, x, y, tx, ty, color=C_GOLD, lw=2.5)

    bottom_note(ax, p["desc"])
    if p.get("notes"):
        note(ax, 0, -3.3, p["notes"], fontsize=7)

    return fig


def render_coverage(cov_key, form_key="2x2"):
    """Render a coverage diagram against a formation."""
    c = COVERAGES[cov_key]
    f = FORMATIONS[form_key]

    title = f"{c['label']}"
    fig, ax = new_fig(title)
    draw_formation(ax, form_key)

    # Use SHADE as base front for all coverage diagrams
    front = FRONTS["SHADE"]

    # Draw DL
    for x, y, lbl in front["dl"]:
        draw_player(ax, x, y, lbl)

    # Build position map for annotation and ghost moves
    pos_map = {}
    for x, y, lbl in front["dl"]:
        pos_map[lbl] = (x, y)

    # Handle pre-snap swaps (Oklahoma FS/D swap, Ohio/Zorro B to middle)
    swaps = c.get("pre_swap", {})
    lb_overrides = {}
    sec_overrides = {}
    for player, (sx, sy) in swaps.items():
        if player in ("FS", "D", "FC", "BC"):
            sec_overrides[player] = (sx, sy)
        else:
            lb_overrides[player] = (sx, sy)

    # Draw LBs
    for item in front["lbs"]:
        x, y, lbl, color = item
        if lbl in lb_overrides:
            x, y = lb_overrides[lbl]
        draw_player(ax, x, y, lbl, color=color)
        pos_map[lbl] = (x, y)
        if lbl in c.get("secondary_notes", {}):
            note(ax, x, y - 1.2 if y < 5 else y + 1.0,
                 c["secondary_notes"][lbl])

    # Draw secondary
    for item in front["secondary"]:
        x, y, lbl, color = item
        if lbl in sec_overrides:
            x, y = sec_overrides[lbl]
        draw_player(ax, x, y, lbl, color=color)
        pos_map[lbl] = (x, y)
        if lbl in c.get("secondary_notes", {}):
            note(ax, x, y + 1.0, c["secondary_notes"][lbl])

    # Draw ghost moves (post-snap rotations)
    for gm in c.get("ghost_moves", []):
        lbl, sx, sy, ex, ey = gm
        color = C_GOLD if lbl in ("FS", "D") else C_COLUMBIA
        draw_ghost(ax, ex, ey, lbl, color=color)
        if abs(ex - sx) > 3:
            draw_curved_arrow(ax, sx, sy, ex, ey, color=color)
        else:
            draw_arrow(ax, sx, sy, ex, ey, color=color, lw=1.8)

    # Draw zone labels
    for zl in c.get("zones", []):
        label, zx, zy, zcolor = zl
        ax.text(zx, zy, label, fontsize=11, ha="center", color=C_NAVY,
                fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.3", facecolor=zcolor, alpha=0.3))

    # Rush arrows for Zeus
    if c.get("rush_arrows"):
        for x, y, lbl in front["dl"]:
            draw_arrow(ax, x, y, x * 0.9, y - 1.5, color=C_NAVY, lw=1.5)

    return fig


# =============================================================================
# DOCX HELPERS
# =============================================================================

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def make_header_row(table, headers):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, NAVY_HEX)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True


# =============================================================================
# DOCX BUILDER HELPERS
# =============================================================================
OUTPUT_DIR = "/home/ksc4130/src/defensive_playbook/diagram_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

FRONT_DESCRIPTIONS = {
    "SHADE": "A=5, T=3, N=2i, E=5. M=open A (field), W=open B (boundary). Base front \u2014 set field.",
    "UNDER": "A=5, T=2i, N=3, E=5. M=open B (field), W=open A (boundary). Tendency breaker; boundary-run weeks.",
    "EYES": "A=5, T=2i, N=2i, E=5. M=open B, W=open B. Square interior vs zone/duo.",
    "WIDE": "A=5, T=3, N=3, E=5. M=open A, W=open A. Force bounce; vs B-gap heavy / gap schemes.",
    "DEUCES": "A=5, T=2, N=2, E=5. M/W react to DT. Vs Wing-T / pullers.",
    "BOSS": "A=5, T=3(F), N=1(F), E=5. Bigs to field. M=A bnd, W=B bnd. Pressure setup.",
    "BOSS UNDER": "A=5, T=1(B), N=3(B), E=5. Bigs to boundary. M=A fld, W=B fld.",
    "GRIZZLY": "A=4i, T=2i, N=2i, E=4i. B/W=OLBs. M at LB depth (10-tech). Vs power/counter, red zone, short yardage. TE SET does NOT override.",
    "MINT": "3-down: A=4i, T=0, N=4i. B/E=OLBs. Man contain=B; NINJA/VIKING=M is QB contain.",
    "ACE": "3-down: A=4, T=0, N=4. All DL 2-gapping (no stunt). Contains=B and E.",
    "JET": "3-down: A=5, T=0, N=5. T is 2-gapping. A/N contain (C-gap edges).",
    "SLIP": "3-down: A=5, T=0, N=4i. 1-gap penetrating. A/E contain. B plays coverage clean. No stunts with SLIP.",
}


def new_doc():
    """Create a fresh document with standard formatting."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    return doc


def doc_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.color.rgb = NAVY
    return h

def doc_body(doc, text):
    return doc.add_paragraph(text)

def doc_img(doc, path, width=5.0):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def doc_title_page(doc, subtitle):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = y.add_run("2026 Season")
    r.font.size = Pt(16); r.font.color.rgb = NAVY
    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("CONFIDENTIAL \u2014 COACHING STAFF ONLY")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    doc.add_page_break()

def doc_footer(doc, diagram_count):
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  "
                   f"\u2014  {diagram_count} diagrams")
    r.font.size = Pt(9); r.font.color.rgb = GRAY_RGB


# =============================================================================
# MAIN GENERATION — SEPARATE FILES
# =============================================================================

def _safe(name):
    """Sanitize a name for filenames."""
    return name.replace(" ", "_").replace("(", "").replace(")", "")


def generate_front_docx(front_key):
    """Generate a single docx for one front: alignments vs formations + legal stunts."""
    paths = {}
    count = 0
    fk = front_key
    fdir = os.path.join(FRONT_DIR, _safe(fk))
    os.makedirs(fdir, exist_ok=True)

    is_3down = fk in ("MINT", "ACE", "JET", "SLIP")
    is_grizzly = fk == "GRIZZLY"

    # Pick formations
    if is_3down:
        forms = ["2x2", "3x1_field", "empty_3x2"]
    elif is_grizzly:
        forms = ["2x2", "3x1_field", "21_pers", "22_pers", "2x2_te_field"]
    else:
        forms = ["2x2", "3x1_field", "2x2_te_field", "2x2_te_bnd",
                 "21_pers", "22_pers"]

    form_labels = {k: v["label"] for k, v in FORMATIONS.items()}

    # Generate front diagrams
    front_imgs = []
    for fm in forms:
        key = f"front_{_safe(fk)}_{fm}"
        fig = render_front_vs_formation(fk, fm)
        p = save_fig(fig, key, subdir=fdir)
        paths[key] = p
        front_imgs.append((fm, p))
        count += 1

    # Generate stunt-on-this-front diagrams (4-down only; SLIP = no stunts)
    stunt_imgs = []
    if not is_3down:
        legal = get_legal_stunts(fk)
        for sk, rating in legal:
            key = f"stunt_{_safe(sk)}_on_{_safe(fk)}"
            sdir = os.path.join(fdir, "stunts")
            os.makedirs(sdir, exist_ok=True)
            fig = render_stunt_on_front(sk, fk, rating)
            p = save_fig(fig, key, subdir=sdir)
            paths[key] = p
            stunt_imgs.append((sk, rating, p))
            count += 1

    # Build document
    doc = new_doc()
    doc_title_page(doc, f"{fk} FRONT")

    # Alignment description
    doc_h1(doc, f"{fk}")
    doc_body(doc, FRONT_DESCRIPTIONS.get(fk, ""))

    # Front vs formations
    doc_h2(doc, "Alignments vs Offensive Formations")
    for fm, img_path in front_imgs:
        doc_body(doc, f"vs {form_labels.get(fm, fm)}:")
        doc_img(doc, img_path)
        doc.add_page_break()

    # Legal stunts
    if stunt_imgs:
        doc_h2(doc, f"Legal Stunts on {fk}")
        best = [(sk, r, p) for sk, r, p in stunt_imgs if r == "BEST"]
        ok = [(sk, r, p) for sk, r, p in stunt_imgs if r == "OK"]

        if best:
            doc_h3(doc, "Recommended Pairings")
            for sk, rating, img_path in best:
                doc_body(doc, f"{sk}:")
                doc_body(doc, STUNTS[sk]["desc"])
                doc_img(doc, img_path)
                doc.add_page_break()

        if ok:
            doc_h3(doc, "Legal Pairings")
            for sk, rating, img_path in ok:
                doc_body(doc, f"{sk}:")
                doc_body(doc, STUNTS[sk]["desc"])
                doc_img(doc, img_path)
                doc.add_page_break()
    elif is_3down and fk == "SLIP":
        doc_h2(doc, "Stunts")
        doc_body(doc, "No stunts with SLIP. The package is designed for clean 1-gap penetration with B in coverage.")

    doc_footer(doc, count)

    out = os.path.join(OUTPUT_DIR, f"{_safe(fk)}.docx")
    doc.save(out)
    print(f"  {fk}: {count} diagrams \u2192 {out}")
    return out, count


def generate_pressures_docx():
    """Generate a single docx for all pressures/blitzes."""
    paths = {}
    count = 0

    doc = new_doc()
    doc_title_page(doc, "PRESSURES & BLITZES")

    doc_h1(doc, "PRESSURE / BLITZ DIAGRAMS")
    doc_body(doc, "Light blue arrows = LB/DB blitzers. "
                  "Red arrows = DL stunt movement (packaged pressures only).")

    press_keys = list(PRESSURES.keys())
    weekly = ["Mike", "Will", "Bandit", "sWarM", "BooM", "BoW", "Hammer", "Eat"]
    specials = [k for k in press_keys if k not in weekly]

    # Weekly
    doc_h2(doc, "Weekly Pressures (Locked Friday Menu)")
    for pk in weekly:
        key = f"press_{_safe(pk)}"
        fig = render_pressure(pk)
        p = save_fig(fig, key, subdir=PRESS_DIR)
        paths[key] = p
        count += 1

        doc_h3(doc, pk)
        doc_body(doc, PRESSURES[pk]["desc"])
        if PRESSURES[pk].get("notes"):
            doc_body(doc, f"Note: {PRESSURES[pk]['notes']}")
        doc_img(doc, p)
        doc.add_page_break()

    # Game-plan specials
    doc_h2(doc, "Game-Plan Specials")
    for pk in specials:
        key = f"press_{_safe(pk)}"
        fig = render_pressure(pk)
        p = save_fig(fig, key, subdir=PRESS_DIR)
        paths[key] = p
        count += 1

        doc_h3(doc, pk)
        doc_body(doc, PRESSURES[pk]["desc"])
        if PRESSURES[pk].get("notes"):
            doc_body(doc, f"Note: {PRESSURES[pk]['notes']}")
        doc_img(doc, p)
        doc.add_page_break()

    doc_footer(doc, count)

    out = os.path.join(OUTPUT_DIR, "Pressures.docx")
    doc.save(out)
    print(f"  Pressures: {count} diagrams \u2192 {out}")
    return out, count


def generate_coverages_docx():
    """Generate a single docx for all coverages."""
    paths = {}
    count = 0

    cov_formations = {
        "NINJA_2x2": "2x2", "NINJA_3x1": "3x1_field",
        "NINJA_empty": "empty_3x2",
        "OREGON": "2x2", "OKLAHOMA": "2x2", "OHIO": "2x2",
        "ZEUS": "2x2", "ZORRO": "2x2", "ZUNNEL": "2x2",
        "ZILL": "2x2", "ZIKE": "2x2", "VIKING": "2x2",
    }
    extra_cov = [
        ("OREGON", "3x1_field"), ("OKLAHOMA", "3x1_field"),
        ("OHIO", "3x1_field"), ("ZEUS", "3x1_field"),
        ("VIKING", "3x1_field"),
    ]

    doc = new_doc()
    doc_title_page(doc, "COVERAGES")

    doc_h1(doc, "COVERAGE DIAGRAMS")
    doc_body(doc, "Dashed circles = post-snap destinations. "
                  "Curved arrows = rotation. Straight arrows = drop.")

    cov_diagram_paths = {}  # cov_key -> list of (fm, img_path)

    cov_keys = list(COVERAGES.keys())
    for ck in cov_keys:
        fm = cov_formations.get(ck, "2x2")
        key = f"cov_{ck}_{fm}"
        family = COVERAGES[ck].get("family", "other").replace(" ", "_")
        cfdir = os.path.join(COV_DIR, family)
        os.makedirs(cfdir, exist_ok=True)
        fig = render_coverage(ck, fm)
        p = save_fig(fig, key, subdir=cfdir)
        paths[key] = p
        count += 1
        cov_diagram_paths.setdefault(ck, []).append((fm, p))

    for ck, fm in extra_cov:
        key = f"cov_{ck}_{fm}"
        if key not in paths:
            family = COVERAGES[ck].get("family", "other").replace(" ", "_")
            cfdir = os.path.join(COV_DIR, family)
            os.makedirs(cfdir, exist_ok=True)
            fig = render_coverage(ck, fm)
            p = save_fig(fig, key, subdir=cfdir)
            paths[key] = p
            count += 1
            cov_diagram_paths.setdefault(ck, []).append((fm, p))

    form_labels = {k: v["label"] for k, v in FORMATIONS.items()}

    # NINJA
    doc_h2(doc, "NINJA (Cover 7)")
    for ck in ["NINJA_2x2", "NINJA_3x1", "NINJA_empty"]:
        doc_h3(doc, COVERAGES[ck]["label"])
        for fm, img_path in cov_diagram_paths.get(ck, []):
            doc_img(doc, img_path)
        doc.add_page_break()

    # Cover 1
    doc_h2(doc, "Cover 1 Family")
    for ck in ["OREGON", "OKLAHOMA", "OHIO"]:
        doc_h3(doc, COVERAGES[ck]["label"])
        doc_body(doc, f"FC/BC: man #1. {COVERAGES[ck]['secondary_notes'].get('M', '')}")
        for fm, img_path in cov_diagram_paths.get(ck, []):
            doc_body(doc, f"vs {form_labels.get(fm, fm)}:")
            doc_img(doc, img_path)
        doc.add_page_break()

    # Cover 0 / Z-Family
    doc_h2(doc, "Cover 0 / Z-Family")
    for ck in ["ZEUS", "ZORRO", "ZUNNEL", "ZILL", "ZIKE"]:
        doc_h3(doc, COVERAGES[ck]["label"])
        doc_body(doc, "Cover 0 \u2014 No deep safety help.")
        for fm, img_path in cov_diagram_paths.get(ck, []):
            doc_body(doc, f"vs {form_labels.get(fm, fm)}:")
            doc_img(doc, img_path)
        doc.add_page_break()

    # VIKING
    doc_h2(doc, "VIKING (Cover 3)")
    doc_h3(doc, COVERAGES["VIKING"]["label"])
    doc_body(doc, "Deep 3: FC=1/3 field, FS=1/3 middle, BC=1/3 boundary. "
                  "Under: B=curl/flat, D=seam-curl-flat, M=hook, W=hook.")
    for fm, img_path in cov_diagram_paths.get("VIKING", []):
        doc_body(doc, f"vs {form_labels.get(fm, fm)}:")
        doc_img(doc, img_path)

    doc_footer(doc, count)

    out = os.path.join(OUTPUT_DIR, "Coverages.docx")
    doc.save(out)
    print(f"  Coverages: {count} diagrams \u2192 {out}")
    return out, count


def generate_all():
    """Generate all diagram documents — separate files."""
    print("=" * 60)
    print("RIVER VALLEY VIKINGS — DIAGRAM GENERATION")
    print("=" * 60)
    print(f"Output directory: {OUTPUT_DIR}\n")

    grand_total = 0

    # Per-front docx files (with legal stunts)
    all_fronts = ["SHADE", "UNDER", "EYES", "WIDE", "DEUCES",
                  "GRIZZLY", "BOSS", "BOSS UNDER",
                  "MINT", "ACE", "JET", "SLIP"]

    print("[1/3] Generating FRONT files (alignments + legal stunts)...")
    for fk in all_fronts:
        _, c = generate_front_docx(fk)
        grand_total += c

    print(f"\n[2/3] Generating PRESSURES file...")
    _, c = generate_pressures_docx()
    grand_total += c

    print(f"\n[3/3] Generating COVERAGES file...")
    _, c = generate_coverages_docx()
    grand_total += c

    print(f"\n{'=' * 60}")
    print(f"DONE. {grand_total} total diagrams across {len(all_fronts) + 2} files.")
    print(f"Output: {OUTPUT_DIR}/")
    print(f"{'=' * 60}")


# =============================================================================
if __name__ == "__main__":
    generate_all()
