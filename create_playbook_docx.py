#!/usr/bin/env python3
"""Generate a formatted .docx defensive playbook from the chat export."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# -- Colors (Columbia Blue / Gold / White / Navy) --
NAVY = RGBColor(0x00, 0x2D, 0x62)
COLUMBIA_BLUE = RGBColor(0x6C, 0xAC, 0xE4)
GOLD = RGBColor(0xCF, 0xA7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
NAVY_HEX = "002D62"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_header_row(table, headers):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, NAVY_HEX)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.font.color.rgb = DARK

    # Helpers
    def heading1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def heading3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def body(text):
        return doc.add_paragraph(text)

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def bold_body(bold_text, normal_text=""):
        p = doc.add_paragraph()
        r = p.add_run(bold_text)
        r.bold = True
        if normal_text:
            p.add_run(normal_text)
        return p

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RIVER VALLEY VIKINGS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("DEFENSIVE PLAYBOOK")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GOLD

    doc.add_paragraph()

    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = y.add_run("2026 Season")
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("CONFIDENTIAL \u2014 COACHING STAFF ONLY")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    heading1("TABLE OF CONTENTS")
    toc = [
        ("1.", "Defensive Philosophy & Identity"),
        ("2.", "Position Designations & Technique Chart"),
        ("3.", "4-Down Front Catalog"),
        ("4.", "3-Down Packages (Mint / Ace / Jet / Slip)"),
        ("5.", "Stunt Catalog"),
        ("6.", "Blitz / Pressure Tags"),
        ("7.", "Coverage: NINJA (Cover 7 Family)"),
        ("8.", "Coverage: Cover 1 Family (Oregon / Oklahoma / Ohio)"),
        ("9.", "Coverage: Cover 0 / Z-Family (Zeus / Zorro / Zunnel / Zill / Zike)"),
        ("10.", "Coverage: VIKING (Cover 3 Family)"),
        ("11.", "Autos & Alerts (BUMP / BANJO / EXCHANGE / UNDER)"),
        ("12.", "Run-Fit Rules"),
        ("13.", "Call Grammar, Signals & Wristband"),
        ("14.", "Deferred / Pending Items"),
    ]
    for num, item in toc:
        p = doc.add_paragraph()
        r = p.add_run(f"{num}  ")
        r.bold = True
        r.font.size = Pt(13)
        r2 = p.add_run(item)
        r2.font.size = Pt(13)

    doc.add_page_break()

    # =========================================================================
    # 1. PHILOSOPHY
    # =========================================================================
    heading1("1. DEFENSIVE PHILOSOPHY & IDENTITY")

    body("\"We play fast, physical, and disciplined football. Our job is to remove explosive plays, "
         "control the box, and force the offense to execute perfectly for an entire drive. We will "
         "pressure and change the picture, but never at the cost of leverage, gap integrity, or effort.\"")

    heading2("Base Structure")
    bullet("4-2-5 base; 3-4 packages (Mint/Ace/Jet/Slip) based on offensive personnel.", bold_prefix="Personnel: ")
    bullet("Strength = Field (\"protect space\").", bold_prefix="Strength System: ")
    bullet("Two-high shell.", bold_prefix="Shell: ")
    bullet("6-man (light box as default).", bold_prefix="Box: ")
    bullet("Front + Stunt + Blitz tag(s) + Coverage/Tags.", bold_prefix="Call Format: ")
    bullet("Signals with rules/adjustments per formation.", bold_prefix="Communication: ")

    heading2("Non-Negotiables")
    bullet("No explosives. Make them snap it again. Eliminate busts and cheap yards.")
    bullet("Aggressive with discipline. Attack protections and tendencies without losing leverage or penalties.")
    bullet("Box the run. Set hard edges, compress gaps, keep the ball in the help.")
    bullet("Relentless pursuit. 11 to the ball, correct angles, swarming tackles \u2014 finish every snap.")
    bullet("Win leverage. Outside vs perimeter/QB run; inside vs quick game/crossers.")
    bullet("Win situations. 1st down sets the series; 3rd down/red zone = get off the field.")
    bullet("Smart aggression. Pressure to take away something \u2014 when needed, call off the horses and force patience.")
    bullet("Communication standard. Simple, loud, early. Safeties are fixers. Alerts: BUMP / BANJO / EXCHANGE / UNDER.")

    heading2("Grading Non-Negotiables")
    bullet("No lost edges \u2014 contain player keeps QB and ball inside.")
    bullet("No uncovered gaps \u2014 if gap changes with stunt/pressure, adjust instantly.")
    bullet("No free releases \u2014 collision/deny access when called; communicate stack/bunch.")
    bullet("No missed tackles \u2014 leverage + near foot + wrap; eliminate yards after contact.")
    bullet("No penalties that extend drives.")
    bullet("No loafs \u2014 pursuit is mandatory; effort is a coaching point on film.")

    body("\"What We Want the Offense to Feel\": Every throw is contested. Every run is boxed and hit twice. "
         "Every series is harder than the last. Nothing is free.")

    doc.add_page_break()

    # =========================================================================
    # 2. POSITIONS & TECHNIQUES
    # =========================================================================
    heading1("2. POSITION DESIGNATIONS & TECHNIQUE CHART")

    heading2("Position Letters")
    table = doc.add_table(rows=12, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(table, ["Letter", "Name", "Role / Notes"])
    positions = [
        ("FC", "Field Corner", "Inside leverage, 6 yds (NINJA default)"),
        ("FS", "Field Safety", "2-high; NINJA field-side caller; 10-12 yds deep; align off #2"),
        ("B", "Bandit (Nickel)", "Field force/contain in base; hybrid DB/OLB; Field OLB in 3-down"),
        ("A", "Anchor (Field DE)", "Box edge; CAMP QB player; default field side"),
        ("T", "Tackle (Field DT)", "3-tech to field in Shade; default field side"),
        ("N", "Nose (Boundary DT)", "2i boundary in Shade; default boundary side"),
        ("E", "Edge (Boundary DE/LB)", "Boundary edge; Boundary OLB in 3-down"),
        ("M", "Mike (Field ILB)", "Open A gap (field) in base; field-side LB"),
        ("W", "Will (Boundary ILB)", "Open B gap (boundary) in base; boundary-side LB"),
        ("D", "Dawg (Boundary Safety)", "Boundary force/contain; NINJA boundary-side caller; 10-12 yds deep"),
        ("BC", "Boundary Corner", "Inside leverage, 6 yds (NINJA default)"),
    ]
    for i, (letter, name, role) in enumerate(positions):
        row = table.rows[i + 1]
        row.cells[0].text = letter
        row.cells[1].text = name
        row.cells[2].text = role

    doc.add_paragraph()
    bold_body("Default Sides: ", "A, T = field (unless called out). N, E = boundary (unless called out).")

    heading2("Technique Chart")
    ttable = doc.add_table(rows=13, cols=2)
    ttable.style = "Light Grid Accent 1"
    ttable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ttable, ["Technique", "Alignment"])
    techs = [
        ("0", "Head up on Center"),
        ("1", "Shade on Center"),
        ("2i", "Inside shade on Guard"),
        ("2", "Head up on Guard"),
        ("3", "Outside shade on Guard"),
        ("4i", "Inside shade on Tackle"),
        ("4", "Head up on Tackle"),
        ("5", "Outside shade on Tackle"),
        ("6i", "Inside shade on TE"),
        ("6", "Head up on TE"),
        ("7", "Outside shade on TE"),
        ("9", "Outside shade on Wing"),
    ]
    for i, (tech, align) in enumerate(techs):
        ttable.rows[i + 1].cells[0].text = tech
        ttable.rows[i + 1].cells[1].text = align

    doc.add_page_break()

    # =========================================================================
    # 3. 4-DOWN FRONTS
    # =========================================================================
    heading1("3. 4-DOWN FRONT CATALOG")

    heading2("Global Rules (4-Down, No Stunts)")
    bullet("M fits the open gap to the field side.")
    bullet("W fits the open gap to the boundary side.")
    bullet("Open gap = determined by interior DT technique: 3-tech closes B gap \u2192 open A; 2/2i/1 closes A \u2192 open B.")
    bullet("A and E = contain (box edges; set the edge, do not wrong-arm by default).")
    bullet("B = field force. D = boundary force.")
    bullet("2-back (20/21): default 6-man box. If inserting a 7th hat, B is first insert (before D).")
    bullet("Puller rule (Power/Counter/Wing-T): Box + overlap \u2014 edge stays outside; M/W overlap inside-out off pull path.")
    bullet("QB Run Rule (CAMP): DE (A) is the QB player (\"who we don't want running\").")

    heading2("TE SET \u2014 Default 4-Down Adjustment")
    body("Overrides ALL 4-down fronts EXCEPT GRIZZLY when there is a TE/Y-off surface.")
    bullet("Set front to TE: 3-tech to TE, 2i away from TE.")
    bullet("B inserts into the box.")
    bullet("End to TE aligns in a 7-technique and plays contain.")

    doc.add_paragraph()

    # Front table
    ftable = doc.add_table(rows=9, cols=5)
    ftable.style = "Light Grid Accent 1"
    ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ftable, ["Front", "A", "T", "N", "E"])
    front_data = [
        ("SHADE (base, set field)", "5", "3", "2i", "5"),
        ("UNDER (set boundary)", "5", "2i", "3", "5"),
        ("EYES (balanced)", "5", "2i", "2i", "5"),
        ("WIDE", "5", "3", "3", "5"),
        ("DEUCES", "5", "2", "2", "5"),
        ("GRIZZLY", "4i", "2i", "2i", "4i"),
        ("BOSS (bigs to field)", "5", "3 (F)", "1 (F)", "5"),
        ("BOSS UNDER (bigs to bnd)", "5", "1 (B)", "3 (B)", "5"),
    ]
    for i, (name, a, t, n, e) in enumerate(front_data):
        row = ftable.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = a
        row.cells[2].text = t
        row.cells[3].text = n
        row.cells[4].text = e

    doc.add_paragraph()

    # Fit table
    fit_table = doc.add_table(rows=9, cols=4)
    fit_table.style = "Light Grid Accent 1"
    fit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(fit_table, ["Front", "M Fit", "W Fit", "Special Notes"])
    fit_data = [
        ("SHADE", "Open A (field)", "Open B (boundary)", "Base vs spread/RPO; stable vs Air Raid"),
        ("UNDER", "Open B (field)", "Open A (boundary)", "Tendency breaker; pairs well with SLANT"),
        ("EYES", "Open B", "Open B", "Vs zone/duo; square interior"),
        ("WIDE", "Open A", "Open A", "Force bounce; vs B-gap heavy / gap schemes"),
        ("DEUCES", "Make T/N right", "Make T/N right", "Vs Wing-T / pullers; gap only if stunt creates it"),
        ("GRIZZLY", "10-tech; no gap (C to TE)", "\u2014 (OLB)", "B & W are OLBs; TE SET does NOT override"),
        ("BOSS", "A gap (boundary)", "B gap (boundary)", "Overload field bigs; pressure setup"),
        ("BOSS UNDER", "A gap (field)", "B gap (field)", "Overload boundary bigs"),
    ]
    for i, (name, m, w, notes) in enumerate(fit_data):
        row = fit_table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = m
        row.cells[2].text = w
        row.cells[3].text = notes

    doc.add_page_break()

    # =========================================================================
    # 4. 3-DOWN PACKAGES
    # =========================================================================
    heading1("4. 3-DOWN PACKAGES (MINT / ACE / JET / SLIP)")

    heading2("Common Rules (All 3-Down)")
    bullet("Field OLB = B (Bandit).")
    bullet("Boundary OLB = E (Edge).")
    bullet("T/N may swap based on personnel.")
    bullet("Slice rule (split zone / H-back kick): backside OLB has the slicer.")
    bullet("Strength declaration: Field (unless later TE/surface rule added).")

    heading2("ILB Fits (Locked)")
    bullet("Play-to (toward strength): STACK B \u2192 D (fit inside-out: B gap first, then spill/replace to D if it bounces).")
    bullet("Play-away (away from strength): SLOW PLAY A (patient shuffle, close backside A, look for cutback).")

    heading2("MINT (4i / 0 / 4i)")
    body("Signal: \"M\" on thigh.")
    bullet("A = 4i, T = 0 (nose), N = 4i. B/E = OLBs.")
    bullet("In man coverage (C0/C1): B is contain.", bold_prefix="Contain: ")
    bullet("In NINJA/VIKING (MINT only): M is QB contain on pass read; remains run-fit first on run.", bold_prefix="Contain: ")

    heading2("ACE (4 / 0 / 4)")
    body("Signal: 1 finger up (optional chest tap).")
    bullet("A = 4, T = 0, N = 4. B/E = OLBs.")
    bullet("A, T, N are ALL 2-gapping (read OL, control both adjacent gaps). No stunt.", bold_prefix="Core Rule: ")
    bullet("Primary contains are OLBs (B and E) \u2014 clean, since the down three are 2-gapping.", bold_prefix="Contain: ")
    bullet("Vs downhill/run-heavy looks when you want the front to absorb blocks and let ILBs run.", bold_prefix="When to Use: ")

    heading2("JET (5 / 0 / 5)")
    body("Signal: whoosh forward (hand flat, forward).")
    bullet("A = 5, T = 0, N = 5. B/E = OLBs.")
    bullet("T is 2-gapping. A and N are contain (C-gap edges).", bold_prefix="Core Rule: ")
    bullet("Vs perimeter run / option / teams trying to bounce; keeps edges firm.", bold_prefix="When to Use: ")

    heading2("SLIP (5 / 0 / 4i)")
    body("Signal: mime pushing a wall to the side (two hands).")
    bullet("A = 5, T = 0, N = 4i. B/E = OLBs.")
    bullet("1-gap penetrating. A and E are contain.", bold_prefix="Core Rule: ")
    bullet("3-down vs spread with NINJA or VIKING behind it. Takes B out of conflict \u2014 B plays coverage clean without force/contain responsibilities.", bold_prefix="When to Use: ")

    heading2("3-Down Contain Matrix")
    ctable = doc.add_table(rows=6, cols=2)
    ctable.style = "Light Grid Accent 1"
    ctable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ctable, ["Package + Coverage", "Contain"])
    cdata = [
        ("MINT + Man (C0/C1)", "B is contain"),
        ("MINT + NINJA/VIKING", "M is QB contain"),
        ("ACE/JET + Man (C0/C1)", "B is contain"),
        ("ACE/JET + NINJA/VIKING", "B is contain-first, then match on pass read"),
        ("SLIP + NINJA/VIKING", "A and E (B plays coverage clean)"),
    ]
    for i, (pkg, cont) in enumerate(cdata):
        ctable.rows[i + 1].cells[0].text = pkg
        ctable.rows[i + 1].cells[1].text = cont

    heading2("TE Adjustment (3-Down)")
    bullet("OLB-only adjust: bump B/E alignments to the TE/Y-off surface without rebuilding the front.")
    bullet("Keep Mint/Ace/Jet/Slip interior intact.")

    doc.add_page_break()

    # =========================================================================
    # 5. STUNT CATALOG
    # =========================================================================
    heading1("5. STUNT CATALOG")
    body("Stunts are callable across fronts. DL coaching point on all movement stunts: read run/pass; "
         "on run read, get heel depth (don't cross-face and create seams). Contain: edges play normal unless noted.")

    heading2("SLANT (to field)")
    bullet("4-down: T + N slant to field.")
    bullet("3-down: A + T + N slant to field.")
    bullet("Do NOT call with Shade (already set to field).", bold_prefix="Restriction: ")
    bullet("Under, BOSS Under.", bold_prefix="Best Fronts: ")
    bullet("Mike, Will, Bandit, sWarM, BooM, BoW.", bold_prefix="Best Pressures: ")

    heading2("ANGLE (to boundary)")
    bullet("Same movement as Slant, but to boundary.")
    bullet("Do NOT call with Under or BOSS Under.", bold_prefix="Restriction: ")
    bullet("Shade, BOSS.", bold_prefix="Best Fronts: ")
    bullet("Mike, Will, Bandit, sWarM, BooM, BoW.", bold_prefix="Best Pressures: ")

    heading2("PINCH")
    bullet("4-down: T + N pinch into A gaps (if already in A, play base).")
    bullet("3-down: A + N pinch into B gaps.")
    bullet("A and E play normal (no contain change).")
    bullet("Shade, Under, Wide, Deuces.", bold_prefix="Best Fronts: ")

    heading2("JACKS")
    bullet("4-down: T + N shoot B gaps (if already in B, play base).")
    bullet("3-down: A + N expand into C gaps (NOT pinch).")
    bullet("A and E are contain. In Mint, M and E are contain (base Mint rule).")
    bullet("Do NOT call with Wide.", bold_prefix="Restriction: ")
    bullet("Eyes, Shade, Under.", bold_prefix="Best Fronts: ")
    bullet("sWarM (primary).", bold_prefix="Best Pressure: ")

    heading2("SPLIT")
    bullet("With BOSS: any DL in a 1-tech crosses the center's face to the opposite A gap.")
    bullet("With GRIZZLY: field-side \"gap out\" one gap (A \u2192 C, T \u2192 B).")
    bullet("A and E are contain (4-down).")
    bullet("When run with BOSS, gives a weird look but plays like Shade.")
    bullet("Mike, sWarM.", bold_prefix="Best Pressures: ")

    heading2("ANCHOR ATTACK")
    bullet("A attacks OL to close one inside gap (usually paired with an outside blitz).")
    bullet("NOT paired with Cobra (Cobra is boundary \u2014 pair Cobra with Edge Attack instead).")

    heading2("EDGE ATTACK")
    bullet("E attacks OL to close one inside gap (usually paired with an outside blitz).")
    bullet("CAN be paired with Cobra.")

    heading2("CRASH")
    bullet("A + E shoot B gaps; T + N shoot A gaps.")
    bullet("B / M / W are contain.", bold_prefix="Contain: ")
    bullet("Do NOT call with Grizzly or Freebird.", bold_prefix="Restriction: ")
    bullet("Wide, Shade, Under.", bold_prefix="Best Fronts: ")
    bullet("BoW (primary).", bold_prefix="Best Pressure: ")

    heading2("ANCHOR RAVEN (A 5 \u2192 B)")
    bullet("A shoots B gap from a 5-technique.")
    bullet("Signal: point field \u2192 flap flap.")
    bullet("Not called with Hammer (Hammer is its own call). Not called with staB.")
    bullet("If paired with Bandit blitz, the call name is BANDIT RAVEN.")

    heading2("EDGE RAVEN (E 5 \u2192 B)")
    bullet("E shoots B gap from a 5-technique.")
    bullet("Signal: point boundary \u2192 flap flap.")

    doc.add_page_break()

    # =========================================================================
    # 6. BLITZ / PRESSURE TAGS
    # =========================================================================
    heading1("6. BLITZ / PRESSURE TAGS")
    body("Rule: All blitzes hit your gap responsibility unless noted.")

    heading2("Single Add-Ons")
    bullet("M blitzes to his gap.", bold_prefix="Mike \u2014 ")
    bullet("W blitzes to his gap.", bold_prefix="Will \u2014 ")
    bullet("B blitzes to his gap.", bold_prefix="Bandit \u2014 ")
    bullet("D blitzes to his gap.", bold_prefix="Dawg \u2014 ")

    heading2("Special Tags")
    bullet("Bandit one gap inside end of line (normally B or C gap).", bold_prefix="staB \u2014 ")
    bullet("MOF safety (per coverage) blitzes.", bold_prefix="Freebird \u2014 ")
    bullet("Boundary corner blitz. In 0-family: D takes #1 boundary, FS assumes D rules, B assumes FS rules, RB funneled by M/W.", bold_prefix="Cobra \u2014 ")

    heading2("Combination Pressures")
    btable = doc.add_table(rows=6, cols=2)
    btable.style = "Light Grid Accent 1"
    btable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(btable, ["Call", "Who Blitzes"])
    combos = [
        ("sWarM", "M + W"),
        ("BooM", "B + M"),
        ("BoW", "B + W"),
        ("MaD", "M + D"),
        ("Eat", "M + W + B"),
    ]
    for i, (call, who) in enumerate(combos):
        btable.rows[i + 1].cells[0].text = call
        btable.rows[i + 1].cells[1].text = who

    heading2("Packaged Pressures (Married to Stunts)")
    bullet("B edge blitz + Anchor Attack (A attacks OL closing one inside gap).", bold_prefix="Hammer \u2014 ")
    bullet("W edge blitz + Edge Attack (E attacks OL closing one inside gap).", bold_prefix="Shave \u2014 ")

    heading2("Weekly Called Pressures (Locked Friday Menu)")
    body("Mike, Will, Bandit, sWarM, BooM, BoW, Hammer, Eat.")
    body("Zeus and Z-family are featured separately \u2014 not part of the called pressure tags.")

    doc.add_page_break()

    # =========================================================================
    # 7. NINJA COVERAGE
    # =========================================================================
    heading1("7. COVERAGE: NINJA (COVER 7 FAMILY)")
    body("Coach calls \"NINJA.\" DBs communicate and execute the correct check based on formation. "
         "FS owns field-side checks. D owns boundary-side checks. Safeties are the fixers \u2014 eliminate busts, make corners right.")

    heading2("Alignment Rules")
    bullet("Inside leverage, 6 yards on #1.", bold_prefix="Corners (FC/BC): ")
    bullet("10-12 yards deep. Inside foot on #2 (~1 yard inside #2). Width: stay on top of seam/slot fade "
           "without overhanging the hash. If no #2 to your side: poach/post posture \u2014 hash-to-middle, "
           "depth 10-12, eyes on #3 threats and QB. Compressed/stack: tighten 1-2 steps, be ready to BANJO.",
           bold_prefix="Safeties (FS/D): ")
    bullet("Apex to field; ready to match out or play curl/flat depending on check.", bold_prefix="B (Bandit): ")

    heading2("NINJA vs 2x2 \u2014 MOD (Field) / CLAMP (Boundary)")
    heading3("Field Side: MOD (FS + FC + B)")
    bullet("Man-match rules on #1. Stay on top, deny explosives. If #1 goes shallow/under quickly, "
           "pass it if rules allow, look for next threat.", bold_prefix="FC: ")
    bullet("MOD safety rules: top-down on #2. If #2 vertical: match and stay on top. "
           "If #2 out/flat quickly: communicate and drive with control (don't open seams behind you).", bold_prefix="FS: ")
    bullet("Apex rules to #2/#3 threats. Eliminate quick RPO access. Carry seams long enough for safety help.", bold_prefix="B: ")
    body("Coaching point: MOD side must not give up seam/slot fade access.")

    heading3("Boundary Side: CLAMP (D + BC)")
    bullet("Clamp corner technique: be physical, deny easy release, protect outside leverage. "
           "If #1 vertical: stay on top. If #1 short/under: don't chase so far you lose help responsibilities.", bold_prefix="BC: ")
    bullet("Clamp safety: control #2 threats and help the corner. If #2 vertical: match and stay on top. "
           "If #2 out fast: drive with leverage and force throw to be perfect.", bold_prefix="D: ")
    body("Coaching point: CLAMP side stops cheap throws and wins 1-on-1 leverage without giving up the deep ball.")

    heading2("NINJA vs 3x1 \u2014 POACH")
    bullet("Trips to field: D is the poach safety.", bold_prefix="Poach Player: ")
    bullet("Trips to boundary: FS is the poach safety.")
    bullet("If #3 vertical: poach player takes it (no free seam).")
    bullet("If #3 shallow/under: poach player communicates and overlaps crossers.")
    bullet("Backside stays sound (no free go balls or glance RPO).")

    heading2("NINJA Motion Rule")
    bullet("BUMP is for MAN coverages ONLY (0/1 families), NOT for NINJA.", bold_prefix="Key: ")
    bullet("When motion creates or removes trips, the defense must re-check coverage:")
    bullet("2x2 \u2192 3x1 (trips created): call \"POACH! POACH!\"")
    bullet("3x1 \u2192 2x2 (trips removed): call \"MOD/CLAMP!\"")
    bullet("If motion does not change trips count: keep original check.")
    bullet("Nearest safety (fixer) calls the re-check; other safety echoes.")

    heading2("NINJA Special Formations")
    bullet("BANJO alert (safety called, not signaled in). "
           "Corner = first OUT; near safety = first IN; opposite safety = deeper fixer.", bold_prefix="Bunch/Stack/Tight Splits: ")
    bullet("EXCHANGE is automatic. Switch on contact; communicate early.", bold_prefix="Rub/Pick: ")
    bullet("Stay split-field, formation-based. DBs identify #3 threats. Safeties responsible for #3.", bold_prefix="Empty: ")
    bullet("Treat nub as #1 to that side. Safety communicates if nub becomes vertical immediately.", bold_prefix="Nub TE / Y-off: ")

    heading2("NINJA Communication Checklist")
    bullet("Coach: \"NINJA!\"")
    bullet("DBs: \"MOD!\" / \"CLAMP!\" (2x2) or \"POACH!\" (3x1)")
    bullet("Vs stack/bunch: \"BANJO!\"")
    bullet("Vs rub/pick: \"EXCHANGE!\"")
    bullet("Motion that changes formation: re-check (\"POACH\" or \"MOD/CLAMP\")")

    doc.add_page_break()

    # =========================================================================
    # 8. COVER 1
    # =========================================================================
    heading1("8. COVERAGE: COVER 1 FAMILY")

    heading2("CAMP Distribution (All 3 Variants)")
    bullet("FC/BC: man on #1 to your side.")
    bullet("Non-post safeties: #2 to their side, or #3 away if no #2 to their side.")

    heading2("RB Funnel / RAT Rules (All 3 Variants)")
    bullet("M and W funnel the RB.")
    bullet("RB releases to your side \u2192 you take RB.")
    bullet("RB releases away \u2192 you rush.")
    bullet("TAMPA tag: LB away from RB becomes the RAT.")

    heading2("OREGON (Post = FS)")
    bullet("FS: Post (MOF) \u2014 fixer; make the corner right.")
    bullet("FC/BC: man on #1 to your side.")
    bullet("D: man on #2 boundary, or #3 away if no #2 boundary.")
    bullet("B: man on #2 field, or #3 away if no #2 field.")
    bullet("M/W: RB funnel rules.")

    heading2("OKLAHOMA (Post = D)")
    bullet("D: Post (MOF) \u2014 fixer; make the corner right.")
    bullet("FC/BC: man on #1 to your side.")
    bullet("FS: man on #2 boundary, or #3 away if no #2 boundary.")
    bullet("B: man on #2 field, or #3 away if no #2 field.")
    bullet("M/W: RB funnel rules.")

    heading2("OHIO (Post = B)")
    bullet("B: Post (MOF) \u2014 fixer; make the corner right.")
    bullet("FC/BC: man on #1 to your side.")
    bullet("FS: man on #2 field, or #3 away if no #2 field.")
    bullet("D: man on #2 boundary, or #3 away if no #2 boundary.")
    bullet("M/W: RB funnel rules.")

    doc.add_page_break()

    # =========================================================================
    # 9. COVER 0 / Z-FAMILY
    # =========================================================================
    heading1("9. COVERAGE: COVER 0 / Z-FAMILY")
    body("All Z-calls are Cover 0 family. Default: man leverage + aggressive rush intent. No deep safety help.")

    heading2("CAMP DB Rules (0-Family Base Distribution)")
    bullet("FC: man on #1 field.")
    bullet("BC: man on #1 boundary.")
    bullet("FS: man on #2 field; if no #2 field, #3 away.")
    bullet("D: man on #2 boundary; if no #2 boundary, #3 away.")

    heading2("Universal Tags (All Z-Calls)")
    bullet("Converts specified 2nd-level players into droppers / inside help (RAT/middle runner).", bold_prefix="TAMPA: ")
    bullet("Tagged player spies the QB.", bold_prefix="SPY: ")

    heading2("ZEUS (Run-First / Pass-Read Delayed Pressure)")
    body("Intent: Run-first alignment that becomes Cover 0 pressure on pass, with built-in RB funnel and QB cage.")
    heading3("Rush on Pass (Always)")
    bullet("A & E: cage/contain rush lanes \u2014 do not run past QB depth; keep outside lanes; force QB to step up into interior push.")
    bullet("T & N: vertical push / collapse pocket.")
    bullet("M: primary add-on rusher (inside lane integrity).")
    heading3("B and W (Conditional \u2014 Run First, Then RB Check)")
    bullet("Run-first: fit rules still apply.")
    bullet("On pass read: RB to your side \u2192 take RB (man). RB away \u2192 rush. RB up the middle \u2192 W takes RB.")
    heading3("Zeus Call-Off")
    body("After showing Zeus 1-3 times, call same look and bail to NINJA or VIKING. "
         "Primary use of VIKING as a call-off from Zeus.")

    heading2("ZORRO")
    bullet("B has RB.", bold_prefix="RB: ")
    bullet("M/W: no pass responsibility unless tagged. TAMPA = M/W become droppers. SPY = tagged player spies QB.")

    heading2("ZUNNEL")
    bullet("M and W funnel RB (same funnel as Zeus). RB to your side \u2192 take him. RB away \u2192 rush. "
           "If TAMPA tagged: away player becomes RAT.", bold_prefix="RB: ")

    heading2("ZILL")
    bullet("W has RB (man).", bold_prefix="RB: ")

    heading2("ZIKE")
    bullet("M has RB (man).", bold_prefix="RB: ")

    heading2("Z-Family RB Player Reference (for UNDER Alert)")
    ztable = doc.add_table(rows=6, cols=2)
    ztable.style = "Light Grid Accent 1"
    ztable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(ztable, ["Z-Call", "RB Player"])
    zdata = [
        ("Zeus", "B or W (funnel)"),
        ("Zorro", "B"),
        ("Zunnel", "M or W (funnel)"),
        ("Zill", "W"),
        ("Zike", "M"),
    ]
    for i, (call, player) in enumerate(zdata):
        ztable.rows[i + 1].cells[0].text = call
        ztable.rows[i + 1].cells[1].text = player

    heading2("0-Family Motion Rules")
    bullet("Receiver motion (jet/orbit/across): Safeties travel (slide/bump). Underneath defenders do NOT slide/bump.")
    bullet("Safeties do NOT cross.")
    bullet("Trips reduce / 3x1 \u2192 2x2: Safeties travel and reset; underneath stays locked unless call assigns RB/empty.")
    bullet("RB motion out (empty): RB taken by nearest by call/personnel (typically B or W).")

    doc.add_page_break()

    # =========================================================================
    # 10. VIKING
    # =========================================================================
    heading1("10. COVERAGE: VIKING (COVER 3 FAMILY)")

    heading2("VIKING (Base) \u2014 Spot-Drop Cover 3 with Match Tags")
    body("Intent: safe, anti-explosive, low bust. Primary Zeus call-off. Also a standalone call.")

    heading3("Deep 3")
    bullet("FC = deep 1/3 (field).")
    bullet("FS = middle 1/3.")
    bullet("BC = deep 1/3 (boundary).")

    heading3("Underneath")
    bullet("B = curl/flat (field).")
    bullet("D = seam-curl-flat (boundary).")
    bullet("M = hook (eyes to QB; relate to #3).")
    bullet("W = hook.")

    heading3("Recommended Tags")
    bullet("Match #2 vertical rules.", bold_prefix="VIKING SEAM: ")
    bullet("Distribution adjustment vs 3x1.", bold_prefix="VIKING PUSH (Trips): ")
    bullet("Carry crosser to depth before passing.", bold_prefix="VIKING CROSS: ")
    bullet("Trigger/replace rules vs now/bubble.", bold_prefix="VIKING SCREEN: ")

    heading2("VIKING RIP / VIKING LIZ (Match Gear \u2014 True Match 3)")
    body("Intent: \"true match 3\" for Air Raid/RPO spacing weeks.")
    bullet("Single-side call (Rip or Liz) sets seam-match distribution rules.")
    bullet("Use cases: verts, sail/flood, switch releases, glance RPO weeks.")

    heading2("VIKING Fire Zone")
    body("Status: DEFERRED (\"forget the fire zone Viking for now\").")

    doc.add_page_break()

    # =========================================================================
    # 11. AUTOS & ALERTS
    # =========================================================================
    heading1("11. AUTOS & ALERTS")
    body("These automatic adjustments are built into the defense. They are triggered by offensive "
         "formations or actions and are NOT signaled in from the sideline (except where noted).")

    heading2("BUMP-BUMP (Motion \u2014 Man Coverages Only)")
    body("Applies to Cover 0 and Cover 1 families. NOT used in NINJA.")
    bullet("Safeties travel with motion. Underneath defenders do NOT slide/bump.")
    bullet("Safeties do NOT cross.")
    bullet("BUMP 1 (Safety \u2192 Safety): Traveling safety yells \"BUMP!\" to hand motion to the other safety because he will not cross.")
    bullet("BUMP 2 (Safety \u2192 Corner): If motion receiver continues and becomes new widest (#1) to that side, safety yells \"BUMP!\" to the corner.")
    bullet("Corner takes motion ONLY when he becomes the new widest (#1) to that side.")

    heading2("BANJO (Safety Alert/Check \u2014 NOT Signaled In)")
    bullet("Stack, tight/close splits, bunch-ish spacing (rub likelihood).", bold_prefix="Trigger: ")
    bullet("FS or D (their side).", bold_prefix="Authority: ")
    bullet("Corner = first OUT.")
    bullet("Near safety = first IN.")
    bullet("Opposite safety = deeper middle fixer (2nd in/2nd out/middle \u2014 makes it right).")

    heading2("EXCHANGE (Automatic \u2014 NOT a Called Tag)")
    bullet("Any rub/pick route concept (mesh, stack releases, traffic).", bold_prefix="Trigger: ")
    bullet("Exchange responsibilities to avoid chasing through wash. Switch on contact. Communicate early.")

    heading2("UNDER (On-Field Alert \u2014 Fixer-Driven)")
    bullet("A receiver runs UNDER inside to pick the defender responsible for the RB AND RB releases FAST OUT to that same side.", bold_prefix="Trigger: ")
    bullet("RB player takes the UNDER route (replaces the pick). The defender who had the UNDER takes the RB (swap).", bold_prefix="Rule: ")
    bullet("Fixer safety on that side calls \"UNDER! UNDER!\" (no numbers).", bold_prefix="Communication: ")
    bullet("RB player does not fight through traffic to chase RB \u2014 he replaces onto the under route immediately.", bold_prefix="Coaching Point: ")
    heading3("RB Player by Coverage Family")
    bullet("Cover 1: M or W (per funnel rules). B may be MOF safety or have a receiver \u2014 do not assume B is free.")
    bullet("Zeus: B or W (funnel rule).")
    bullet("Zorro: B.")
    bullet("Zunnel: M or W (funnel).")
    bullet("Zill: W.")
    bullet("Zike: M.")

    doc.add_page_break()

    # =========================================================================
    # 12. RUN-FIT RULES
    # =========================================================================
    heading1("12. RUN-FIT RULES")

    heading2("Defensive Run Goals")
    bullet("No explosives.")
    bullet("Box it (set hard edges).")
    bullet("Relentless pursuit (11 hats to the ball).")
    bullet("Force the ball to the help.")

    heading2("Base Run Structure (4-Down, Non-Grizzly)")
    bullet("Edges: A/E contain (box it \u2014 set the edge, do not wrong-arm by default).")
    bullet("Force: B = field force. D = boundary force.")
    bullet("ILBs (no stunt base): M fits open gap field. W fits open gap boundary.")
    bullet("Open gap determined by interior technique: 3-tech closes B \u2192 open A; 2/2i/1 closes A \u2192 open B.")

    heading2("Special Rules")
    bullet("Box + overlap: edge stays outside; M/W overlap inside-out off pull path.", bold_prefix="Pullers (Power/Counter/Wing-T): ")
    bullet("DE (A) is the QB player (CAMP rule).", bold_prefix="QB Run: ")
    bullet("Default 6-man box. If inserting a 7th hat, B is first insert (before D).", bold_prefix="2-Back (20/21): ")
    bullet("Backside OLB has the slicer.", bold_prefix="Slice (Split Zone / H-Back Kick): ")

    heading2("Front-Specific Run Notes")
    bullet("M/W make T and N right; only become true gap players when a stunt creates the gap.", bold_prefix="DEUCES: ")
    bullet("B and W are OLBs. M in a 10-tech, no gap by default. If TE surface: M has C gap to TE.", bold_prefix="GRIZZLY: ")
    bullet("3-tech to field + 1-tech to field. M = A to boundary. W = B to boundary.", bold_prefix="BOSS: ")
    bullet("Inverted. M = A to field. W = B to field.", bold_prefix="BOSS UNDER: ")

    heading2("3-Down ILB Fits")
    bullet("Play-to: STACK B \u2192 D (inside-out).")
    bullet("Play-away: SLOW PLAY A (hunt cutback).")

    doc.add_page_break()

    # =========================================================================
    # 13. CALL GRAMMAR, SIGNALS & WRISTBAND
    # =========================================================================
    heading1("13. CALL GRAMMAR, SIGNALS & WRISTBAND")

    heading2("Call Order")
    bold_body("Front + Stunt + Blitz tag(s) + Coverage / Tags")

    heading2("Example Calls")
    bullet("SHADE + SLANT + BoW + NINJA")
    bullet("GRIZZLY + PINCH + Mike + OREGON")
    bullet("BOSS + SPLIT + Cobra + ZEUS")
    bullet("ACE + ANGLE + Freebird + VIKING RIP")
    bullet("UNDER + SLANT + sWarM + NINJA")
    bullet("BOSS + SPLIT + Eat + VIKING")
    bullet("ZEUS (show) \u2192 call-off to NINJA / VIKING")

    heading2("Signal Set")
    heading3("3-Down Front Signals")
    bullet("\"M\" on thigh.", bold_prefix="MINT: ")
    bullet("1 finger up (optional chest tap).", bold_prefix="ACE: ")
    bullet("Whoosh forward (hand flat, forward).", bold_prefix="JET: ")
    bullet("Mime pushing a wall to the side (two hands).", bold_prefix="SLIP: ")

    heading3("Raven Stunt Signals")
    bullet("Point field \u2192 flap flap.", bold_prefix="Anchor Raven: ")
    bullet("Point boundary \u2192 flap flap.", bold_prefix="Edge Raven: ")

    heading3("Signal Grammar (Two-Part for DL/LBs)")
    body("Front sign, then Stunt sign.")

    heading2("Wristband Menu Structure")
    wtable = doc.add_table(rows=6, cols=2)
    wtable.style = "Light Grid Accent 1"
    wtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(wtable, ["Category (Code Range)", "Items"])
    wdata = [
        ("Fronts (01-20)", "01 Shade, 02 Under, 03 Eyes, 04 Wide, 05 Deuces, 06 Grizzly, 07 Boss, 08 Boss Under"),
        ("Stunts (21-35)", "21 Slant, 22 Angle, 23 Pinch, 24 Jacks, 25 Split, 26 Crash, 27 Anchor Raven, 28 Edge Raven"),
        ("Pressures (36-50)", "36 Mike, 37 Will, 38 Bandit, 39 sWarM, 40 BooM, 41 BoW, 42 Hammer, 43 Eat, 44 Zeus"),
        ("Coverages (51-65)", "51 Ninja, 52 Oregon, 53 Oklahoma, 54 Ohio, 55 Viking"),
        ("Tags/Autos (66-75)", "66 Tampa, 67 Spy (others as needed)"),
    ]
    for i, (cat, items) in enumerate(wdata):
        wtable.rows[i + 1].cells[0].text = cat
        wtable.rows[i + 1].cells[1].text = items

    doc.add_page_break()

    # =========================================================================
    # 14. DEFERRED ITEMS
    # =========================================================================
    heading1("14. DEFERRED / PENDING ITEMS")
    bullet("Fire zone version of VIKING \u2014 3 deep / 3 under discussed but explicitly deferred (\"forget the fire zone Viking for now\").")
    bullet("VIKING RIP/LIZ match-3 rules \u2014 referenced as deferred in documentation.")
    bullet("BOSS inside tech (N 2i vs 1 to field) \u2014 whether it is choose/stem/personnel not explicitly locked.")
    bullet("NINJA detail pages \u2014 exact position-by-position responsibilities for each route concept not fully authored.")
    bullet("3-4 TE/surface rule \u2014 OLB-only adjust recommended but not explicitly locked by DC.")
    bullet("Called pressure cards \u2014 full who-rushes/who-drops for each of the 8 weekly pressures not yet written.")
    bullet("Formation adjustment sheets \u2014 bunch/stack/empty/nub in detail form.")
    bullet("Situational packages \u2014 3rd down menu, red zone, backed up, sudden change, 2-minute, 4-minute.")
    bullet("Install plan \u2014 week-by-week order not yet authored.")
    bullet("Game-day call sheet, signal sheet, and wristband \u2014 not content-complete.")
    bullet("Diagrams \u2014 excluded (\"recreate no diagrams\").")

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run("\u2014 END OF PLAYBOOK \u2014")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY

    doc.add_paragraph()
    g = doc.add_paragraph()
    g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = g.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    output_path = "/home/ksc4130/src/defensive_playbook/River_Valley_Vikings_Defensive_Playbook.docx"
    doc.save(output_path)
    print(f"Playbook saved to: {output_path}")


if __name__ == "__main__":
    main()
